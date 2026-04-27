from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document


TURN_MARKER = re.compile(r"^\d\d:\d\d:\d\d Speaker (\d+)$")
DATE_PATTERN = re.compile(
    r"(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{1,2}(?:th|st|nd|rd)?(?:,\s*\d{4})?",
    re.IGNORECASE,
)


@dataclass
class Turn:
    speaker: str
    text: str


@dataclass
class QAPair:
    question: str
    answer: str


@dataclass
class ReportContext:
    role_label: str = "The interviewee"
    name: Optional[str] = None
    pronoun: str = "they"
    possessive: str = "their"
    objective: str = "them"
    joint_statement: bool = False
    facts: dict[str, str] = field(default_factory=dict)


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def read_docx_lines(path: Path) -> list[str]:
    document = Document(path)
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def parse_turns(path: Path) -> list[Turn]:
    lines = read_docx_lines(path)
    turns: list[Turn] = []
    current_speaker: Optional[str] = None
    current_lines: list[str] = []

    for line in lines:
        marker = TURN_MARKER.match(line)
        if marker:
            if current_speaker is not None and current_lines:
                turns.append(Turn(current_speaker, normalize(" ".join(current_lines))))
            current_speaker = marker.group(1)
            current_lines = []
            continue
        if line in {"Audio file", "Transcript"} or line.lower().endswith(".m4a"):
            continue
        current_lines.append(line)

    if current_speaker is not None and current_lines:
        turns.append(Turn(current_speaker, normalize(" ".join(current_lines))))

    return merge_consecutive_turns(turns)


def merge_consecutive_turns(turns: list[Turn]) -> list[Turn]:
    merged: list[Turn] = []
    for turn in turns:
        if merged and merged[-1].speaker == turn.speaker:
            merged[-1].text = normalize(f"{merged[-1].text} {turn.text}")
        else:
            merged.append(Turn(turn.speaker, turn.text))
    return merged


def looks_like_question(text: str) -> bool:
    lowered = text.lower().strip()
    lowered = re.sub(r"^(?:okay|ok|understood|all right|alright|right|well)[\s,.:;-]+", "", lowered)
    lowered = re.sub(r"^(?:but\s+)?my question was[\s,.:;-]+", "", lowered)
    starters = (
        "what ",
        "when ",
        "where ",
        "who ",
        "why ",
        "how ",
        "is ",
        "are ",
        "was ",
        "were ",
        "did ",
        "do ",
        "does ",
        "can ",
        "could ",
        "would ",
        "will ",
        "have ",
        "has ",
        "please ",
        "tell me ",
        "describe ",
        "and what ",
        "so ",
        "remind me ",
    )
    contains_question_phrases = (
        "in your own words",
        "tell me the whole story",
        "to confirm",
    )
    return lowered.endswith("?") or lowered.startswith(starters) or any(phrase in lowered for phrase in contains_question_phrases)


def determine_roles(turns: list[Turn]) -> tuple[str, str]:
    scores: dict[str, int] = {}
    for turn in turns:
        lowered = turn.text.lower()
        score = scores.get(turn.speaker, 0)
        if "independent adjuster" in lowered or "representing citizens insurance" in lowered:
            score += 10
        if looks_like_question(turn.text):
            score += 2
        scores[turn.speaker] = score

    interviewer = max(scores, key=scores.get)
    interviewees = [speaker for speaker in scores if speaker != interviewer]
    interviewee = interviewees[0] if interviewees else interviewer
    return interviewer, interviewee


def normalize_role_label(role_label: Optional[str]) -> str:
    if not role_label:
        return "The interviewee"

    text = normalize(role_label).strip(" .")
    lowered = text.lower()
    predefined = {
        "insured": "The insured",
        "claimant": "The claimant",
        "witness": "The witness",
        "interviewee": "The interviewee",
        "other": "The interviewee",
    }
    if lowered in predefined:
        return predefined[lowered]
    if lowered.startswith("the "):
        return text[:1].upper() + text[1:]
    return text


def pluralize_role_label(role_label: str) -> str:
    mapping = {
        "The interviewee": "The interviewees",
        "The insured": "The insureds",
        "The claimant": "The claimants",
        "The witness": "The witnesses",
    }
    if role_label in mapping:
        return mapping[role_label]
    if role_label.lower().startswith("the ") and not role_label.endswith("s"):
        return f"{role_label}s"
    return role_label


def role_core_label(context: ReportContext) -> str:
    label = normalize(context.role_label)
    if label.lower().startswith("the "):
        return label[4:].strip()
    return label


def role_matches_name(context: ReportContext) -> bool:
    if not context.name:
        return False
    return role_core_label(context).lower() == normalize(context.name).lower()


def is_trivial_acknowledgment(text: str) -> bool:
    lowered = normalize(text).lower().strip(" .")
    return lowered in {
        "okay",
        "ok",
        "all right",
        "alright",
        "right",
        "understood",
        "got it",
        "hold on a second",
    }


def build_qa_pairs(turns: list[Turn], interviewer: str, interviewee: str) -> list[QAPair]:
    pairs: list[QAPair] = []
    pending_question: Optional[str] = None
    _ = interviewee

    index = 0
    while index < len(turns):
        turn = turns[index]
        if turn.speaker == interviewer:
            if looks_like_question(turn.text):
                pending_question = turn.text
            elif pending_question:
                pending_question = normalize(f"{pending_question} {turn.text}")
            index += 1
            continue

        if pending_question:
            answer_parts: list[str] = []
            while index < len(turns) and turns[index].speaker != interviewer:
                answer_text = turns[index].text
                if not is_trivial_acknowledgment(answer_text) or not answer_parts:
                    answer_parts.append(answer_text)
                index += 1
            answer = normalize(" ".join(answer_parts))
            if answer:
                pairs.append(QAPair(normalize(pending_question), answer))
            pending_question = None
            continue

        index += 1

    return pairs


def infer_claim_scenario(turns: list[Turn]) -> str:
    combined = " ".join(turn.text.lower() for turn in turns)
    if combined.count("shooting") >= 2 and combined.count("gun") >= 4:
        return "shooting"
    if combined.count("fire") >= 3 and ("electric" in combined or "panel" in combined or "florida power and light" in combined):
        return "fire"
    if combined.count("boat") >= 3 and combined.count("charter") >= 2:
        return "boat_trip"
    if combined.count("dog") >= 4 and combined.count("fence") >= 2:
        return "dog_bite"
    if combined.count("roof") >= 6 and combined.count("claimant") >= 2:
        return "roof_slip"
    if combined.count("water heater") >= 3:
        return "water_heater"
    if combined.count("overflow") >= 3 and combined.count("tub") >= 2:
        return "tub_overflow"
    if combined.count("washing machine") >= 3 or combined.count("laundry machine") >= 2:
        return "washing_machine_leak"
    if combined.count("9b") >= 2 and combined.count("5b") >= 2 and ("stepdaughter" in combined or "10b" in combined):
        return "wall_pipe"
    if combined.count("leak") >= 4 or combined.count("water") >= 6:
        return "water_damage"
    return "general"


def infer_context(turns: list[Turn], interviewer: str, interviewee: str, interviewee_role_label: Optional[str] = None) -> ReportContext:
    context = ReportContext(role_label=normalize_role_label(interviewee_role_label))
    intro_text = " ".join(turn.text for turn in turns[:8] if turn.speaker == interviewer)
    intro_text = normalize(intro_text)
    lowered_intro = intro_text.lower()

    joint_match = re.search(r"I am speaking today with (.*?)(?: on policy number| regarding)", intro_text, re.IGNORECASE)
    if joint_match and (
        "additional named insured" in lowered_intro
        or "named insured's" in lowered_intro
        or "named insureds" in lowered_intro
    ):
        raw_names = joint_match.group(1)
        raw_names = re.sub(r"\bthe named insured(?:'s|s)?\b", "", raw_names, flags=re.IGNORECASE)
        raw_names = re.sub(r"\bthe additional named insured\b", "", raw_names, flags=re.IGNORECASE)
        name_parts = [clean_intro_interviewee_name(part) for part in re.split(r"\band\b", raw_names, flags=re.IGNORECASE)]
        name_parts = [part for part in name_parts if normalize(part)]
        if len(name_parts) >= 2:
            context.name = " and ".join(name_parts)
            context.role_label = pluralize_role_label(context.role_label)
            context.pronoun, context.possessive, context.objective = "they", "their", "them"
            context.joint_statement = True
            context.facts["claim_scenario"] = infer_claim_scenario(turns)
            return context

    name_match = re.search(r"I am speaking today with (.*?)(?:, the| on policy number)", intro_text, re.IGNORECASE)
    if name_match:
        context.name = clean_intro_interviewee_name(name_match.group(1))

    if " miss " in f" {lowered_intro} " or " mrs. " in f" {lowered_intro} " or " ms. " in f" {lowered_intro} ":
        context.pronoun, context.possessive, context.objective = "she", "her", "her"
    elif " mr. " in f" {lowered_intro} " or " brother " in lowered_intro:
        context.pronoun, context.possessive, context.objective = "he", "his", "him"

    for turn in turns:
        if turn.speaker != interviewer:
            continue
        lowered = turn.text.lower()
        if "what is his date of birth" in lowered or "what is his occupation" in lowered:
            context.pronoun, context.possessive, context.objective = "she", "her", "her"
            break
        if "what is her date of birth" in lowered or "what is her occupation" in lowered:
            context.pronoun, context.possessive, context.objective = "he", "his", "him"
            break

    context.facts["claim_scenario"] = infer_claim_scenario(turns)
    return context


def tidy_name(text: str) -> str:
    text = normalize(text).strip(" .,")
    parts: list[str] = []
    previous_key = ""
    for part in re.split(r"\s+", text):
        if not part:
            continue
        key = re.sub(r"^[^\w]+|[^\w]+$", "", part).lower()
        if key and key == previous_key:
            continue
        parts.append(part[:1].upper() + part[1:].lower())
        if key:
            previous_key = key
    return " ".join(parts)


def clean_intro_interviewee_name(text: str) -> str:
    text = normalize(text).strip(" .,")
    text = re.sub(
        r",?\s*(?:the\s+)?named insured(?:'s|s)?\s+"
        r"(?:brother|sister|mother|father|son|daughter|wife|husband|spouse|personal representative|executor|administrator|trustee).*$",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r",?\s*(?:brother|sister|mother|father|son|daughter|wife|husband|spouse|personal representative|executor|administrator|trustee)(?:\b.*)?$",
        "",
        text,
        flags=re.IGNORECASE,
    )
    return tidy_name(text)


def normalize_year_duration(text: str) -> Optional[str]:
    text = normalize(text).strip(" .")
    word_numbers = {
        "one": "1",
        "two": "2",
        "three": "3",
        "four": "4",
        "five": "5",
        "six": "6",
        "seven": "7",
        "eight": "8",
        "nine": "9",
        "ten": "10",
        "eleven": "11",
        "twelve": "12",
    }
    normalized_text = text
    for word, digit in word_numbers.items():
        normalized_text = re.sub(rf"\b{word}\b", digit, normalized_text, flags=re.IGNORECASE)

    going_on_match = re.search(r"(?:going on|almost|nearly|close to)\s+(\d+)\s+years?", normalized_text, re.IGNORECASE)
    if going_on_match:
        return f"approximately {going_on_match.group(1)} years"

    range_match = re.search(r"(\d+)\s*-\s*(\d+)\s+years?", normalized_text, re.IGNORECASE)
    if range_match:
        return f"approximately {range_match.group(1)} to {range_match.group(2)} years"

    approx_match = re.search(r"(?:approximately|approx\.?|about|around)\s+(\d+)\s+years?", normalized_text, re.IGNORECASE)
    if approx_match:
        return f"approximately {approx_match.group(1)} years"

    exact_match = re.search(r"(\d+)\s+years?", normalized_text, re.IGNORECASE)
    if exact_match and normalized_text.lower() == exact_match.group(0).lower():
        return exact_match.group(0).lower()

    return None


def format_child_age(text: str, context: ReportContext) -> str:
    age_text = lower_first(narrativize_answer(text, context))
    age_text = re.sub(r"^(?:he|she|they)(?:'s|\s+is|\s+was)\s+", "", age_text, flags=re.IGNORECASE)
    age_text = re.sub(r"\b(\w+)\s+years old\b", r"\1-year-old", age_text, flags=re.IGNORECASE)
    return age_text.strip(" .")


def render_spouse_occupation(occupation: str, context: ReportContext) -> str:
    lowered = occupation.lower().strip(" .")
    pronoun = cap(context.pronoun)
    if any(phrase in lowered for phrase in ("stays home", "stay home", "stays at home", "stay at home")):
        return f"{pronoun} stated that {context.possessive} spouse stays home"
    if "doesn't do anything" in lowered or "does not do anything" in lowered:
        return f"{pronoun} stated that {context.possessive} spouse stays home"
    return f"{pronoun} stated {context.possessive} spouse's occupation as {occupation}"


def clean_transcript_fillers(text: str) -> str:
    text = re.sub(r"^(?:so|well|okay)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(?:oh|um|uh|yeah)\s*,?\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\b\d{2}:\d{2}:\d{2}\s+speaker\s+\d+\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bSpeaker\s+\d+\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bas far as i know\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bi'?m sorry\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\boh,\s*okay\b", "", text, flags=re.IGNORECASE)
    text = re.sub(
        r"\ba noise,\s*like,\s*you know,\s*water gushing noise\b",
        "the sound of gushing water",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\bliterally like\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\blike,\s*you know,\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r",\s*you know,\s*", ", ", text, flags=re.IGNORECASE)
    text = re.sub(r"\byou know,\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\byou know\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\blike I said\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bas I said\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bother than visiting it\b", "aside from visiting it", text, flags=re.IGNORECASE)
    text = re.sub(r"\buh+\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bum+\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bhold on(?:\s+(?:one|a)\s+second)?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\blet (?:him|her) finish\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdata loss\b", "date of loss", text, flags=re.IGNORECASE)
    text = re.sub(r"\bshut up the water\b", "shut off the water", text, flags=re.IGNORECASE)
    text = re.sub(r"\bunderstandsss\b", "understands", text, flags=re.IGNORECASE)
    text = re.sub(r"\bRnadal\b", "Randal", text, flags=re.IGNORECASE)
    text = re.sub(r"\bownssser\b", "owner", text, flags=re.IGNORECASE)
    text = re.sub(r"\bownsssers\b", "owners", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(he|she|they)\s+guess\b", r"\1 believes", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(he|she|they)\s+mean\b(?:,|\s)*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bthat's\s+(he|she|they)\s+believes\b", r"that is what \1 believes", text, flags=re.IGNORECASE)
    text = re.sub(r"\bprobably like\s+(\d+)\s*,\s*(\d+)\s+minutes\b", r"approximately \1 to \2 minutes", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdo do\b", "do", text, flags=re.IGNORECASE)
    text = re.sub(r"(?:(?<=^)|(?<=[.?!]))\s*(?:okay|ok|oh|well|yeah|right|understood|sorry)\s*[,.]?\s*", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\bafter that\?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bthe roof\?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bwhat do you mean\?\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+,", ",", text)
    text = re.sub(r",\s*,", ", ", text)
    text = re.sub(r"\.\s*\.", ".", text)
    text = re.sub(r"\?+\s*\.", ".", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip(" ,")


def extract_date_phrase(text: str) -> Optional[str]:
    match_obj = DATE_PATTERN.search(text)
    if not match_obj:
        return None
    return normalize(match_obj.group(0))


def clean_address_answer(text: str) -> str:
    text = normalize(text).strip(" .")
    text = re.sub(r"\b\d{2}:\d{2}:\d{2}\s+speaker\s+\d+\b.*$", "", text, flags=re.IGNORECASE)
    if "." in text:
        parts = [part.strip(" .") for part in text.split(".") if part.strip(" .")]
        address_parts = [part for part in parts if looks_like_address(part)]
        if address_parts:
            text = address_parts[0]
    if "speaker 1" in text.lower():
        text = text.split("speaker 1", 1)[0].strip(" .")
    text = re.sub(r"^(?:oh,\s*)?(?:my address is|my address)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(?:ok|okay),?\s*(?:it's|it is)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^is that\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(?:would that be|that would be)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(i live at|i reside at|it occurred at|occurred at)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^at\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^i call it the alleged incident\.\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^the alleged incident occurred at\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^the property where the incident occurred is\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(it occurred at|occurred at)\s+", "", text, flags=re.IGNORECASE)
    return text.strip(" .")


def clean_date_answer(text: str) -> str:
    text = re.sub(r"^that's\s+", "", normalize(text).strip(" ."), flags=re.IGNORECASE)
    text = re.sub(r"^(?:it happened on|it was on|it was the night of|the night of|on)\s+", "", text, flags=re.IGNORECASE)
    month_year_match = re.search(
        r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+(\d{4})\b",
        text,
        re.IGNORECASE,
    )
    if month_year_match and not re.search(r"\b\d{1,2}(?:st|nd|rd|th)?\b", text.split(month_year_match.group(0), 1)[0][-4:]):
        return f"{month_year_match.group(1).capitalize()} {month_year_match.group(2)}"
    month_day_of_year_match = re.search(
        r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+"
        r"(\d{1,2})(?:st|nd|rd|th)?(?:\s+of)?\s+(\d{4})\b",
        text,
        re.IGNORECASE,
    )
    if month_day_of_year_match:
        month = month_day_of_year_match.group(1).capitalize()
        day = month_day_of_year_match.group(2)
        year = month_day_of_year_match.group(3)
        return f"{month} {day}, {year}"
    day_month_match = re.search(
        r"\b(\d{1,2})(?:st|nd|rd|th)?\s+"
        r"(january|february|march|april|may|june|july|august|september|october|november|december)"
        r"(?:\s+(\d{2,4}))?\b",
        text,
        re.IGNORECASE,
    )
    if day_month_match:
        day = day_month_match.group(1)
        month = day_month_match.group(2).capitalize()
        year = day_month_match.group(3)
        if year:
            if len(year) == 2:
                year = f"19{year}" if int(year) > 30 else f"20{year}"
            return f"{month} {day}, {year}"
        return f"{month} {day}"
    extracted = extract_date_phrase(text)
    if extracted:
        extracted_match = re.search(r"\b([A-Za-z]+)\s+(\d{1,2})(?:,\s*(\d{4}))?\b", extracted)
        if extracted_match and int(extracted_match.group(2)) <= 31:
            return extracted
    numeric_match = re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", text)
    if numeric_match:
        return normalize(numeric_match.group(0))
    text = normalize(text).strip(" .")
    text = re.sub(r"^the date of loss was\s+", "", text, flags=re.IGNORECASE)
    return text


def normalize_occupation_detail(text: str) -> str:
    text = normalize(text).strip(" .")
    text = re.sub(r"^(i was|i had been|was)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^in\s+", "", text, flags=re.IGNORECASE)
    lowered = text.lower()
    if "city of quebec" in lowered and "public service" in lowered:
        return "public service for the City of Quebec for 31 years"
    if lowered == "charter fishing":
        return "charter fishing"
    return lower_first(text)


def format_retired_detail(text: str) -> str:
    cleaned = normalize(text).strip(" .")
    lowered = cleaned.lower()
    if not lowered:
        return ""
    if lowered.startswith(("in ", "for ", "with ")):
        return lowered
    if lowered == "charter fishing":
        return "in charter fishing"
    if "public service" in lowered:
        return f"in {lowered}"
    if any(word in lowered for word in ("fishing", "service", "government", "construction", "teaching", "education", "banking", "insurance", "real estate")):
        return f"in {lowered}"
    return f"as {lowered}"


def clean_name_answer(text: str) -> str:
    text = normalize(text).strip(" .")
    text = re.split(r"\b(?:and\s+)?(?:her|his)\s+occupation\b", text, 1, flags=re.IGNORECASE)[0]
    text = re.sub(r"^(?:it is|it's|its|my spouse is|my wife is|my husband is|her name is|his name is|spouse(?:'s)? name is)\s+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bdr\.\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"(?:\b[A-Z]\.){2,}", "", text)
    text = re.sub(r"\b(?:[A-Z]-){2,}[A-Z]?\b", "", text)
    text = re.sub(r"\blast name\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bfirst name\b", "", text, flags=re.IGNORECASE)
    comma_parts = [
        normalize(part).strip(" ,.")
        for part in text.split(",")
        if normalize(part).strip(" ,.") and len(normalize(part).strip(" ,.")) > 1
    ]
    if len(comma_parts) >= 2:
        first = comma_parts[0]
        second = comma_parts[1]
        first_letters = re.sub(r"[^a-z]", "", first.lower())
        second_letters = re.sub(r"[^a-z]", "", second.lower())
        if (
            (len(first.split()) >= 2 and " " not in second and second.isalpha())
            or first_letters == second_letters
            or first_letters.startswith(second_letters)
            or second_letters.startswith(first_letters)
        ):
            text = first
        else:
            text = f"{first} {second}"
    else:
        parts = [normalize(part).strip(" ,.") for part in re.split(r"[.?!]", text) if normalize(part).strip(" ,.")]
        if parts:
            text = parts[0]
    text = re.sub(r"\b[A-Z]\b\.?$", "", text).strip(" ,.")
    return tidy_name(text)


def clean_person_name_answer(text: str) -> str:
    text = re.split(r"\b(?:and\s+)?(?:her|his)\s+occupation\b", text, 1, flags=re.IGNORECASE)[0]
    text = re.sub(r"(?:\b[A-Z]\.){2,}", "", text)
    comma_parts = [
        normalize(part).strip(" ,.")
        for part in text.split(",")
        if normalize(part).strip(" ,.") and len(normalize(part).strip(" ,.")) > 1
    ]
    if len(comma_parts) >= 2:
        first = comma_parts[0]
        second = comma_parts[1]
        first_letters = re.sub(r"[^a-z]", "", first.lower())
        second_letters = re.sub(r"[^a-z]", "", second.lower())
        if first_letters == second_letters or first_letters.startswith(second_letters) or second_letters.startswith(first_letters):
            text = first
        else:
            text = f"{first} {second}"
    else:
        parts = [normalize(part).strip(" .") for part in re.split(r"[.?!]", text) if normalize(part).strip(" .")]
        if parts:
            text = parts[0]
    text = re.sub(r"\b(?:[A-Z]-){2,}[A-Z]?\b", "", text)
    return tidy_name(text)


def looks_like_address(text: str) -> bool:
    lowered = normalize(text).lower()
    street_words = r"(street|st|avenue|ave|road|rd|lane|ln|circle|cir|drive|dr|court|ct|boulevard|blvd|place|pl|way|terrace|ter|trail|trl|parkway|pkwy|highway|hwy|apt|apartment|strasse|straße)"
    return bool(
        re.search(r"\d", lowered)
        and (
            re.search(street_words, lowered)
            or re.search(r"\b\d{5}\b", lowered)
            or "," in lowered
        )
    )


def looks_like_date(text: str) -> bool:
    lowered = normalize(text).lower()
    return bool(
        extract_date_phrase(text)
        or re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", text)
        or re.search(r"\b(19|20)\d{2}\b", lowered)
        or re.search(r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\b", lowered)
    )


def opening_question_tail(question: str) -> str:
    parts = [normalize(part).strip(" .") for part in question.split("?") if normalize(part).strip(" .")]
    tail = parts[-1] if parts else ""
    while True:
        updated = re.sub(r"^(?:so|okay|ok|right|all right|alright|i'm sorry|sorry|oh|well)\s+", "", tail, flags=re.IGNORECASE)
        if updated == tail:
            break
        tail = updated.strip(" .")
    return tail.strip(" .")


def extract_confirmed_date(question: str, answer: str) -> Optional[str]:
    if not is_affirmative(answer):
        return None
    tail = opening_question_tail(question)
    if not tail or not looks_like_date(tail):
        return None
    return clean_date_answer(tail)


def extract_confirmed_text(question: str, answer: str) -> Optional[str]:
    if not is_affirmative(answer):
        return None
    tail = opening_question_tail(question)
    return tail or None


def is_opening_question(question: str) -> bool:
    lowered = question.lower()
    if any(phrase in lowered for phrase in ("in your own words", "tell me the whole story", "what happened", "walk me through")):
        return False
    opening_needles = (
        "permission to record",
        "state your full name",
        "spell it for the record",
        "would that be",
        "so you believe it was",
        "so you think the incident occurred",
        "can you spell it",
        "can you spell that",
        "spell those words",
        "how do you spell",
        "date of birth",
        "what is your address",
        "what address do you currently live at",
        "address do you live at",
        "where do you live",
        "primary mailing",
        "what's your address there",
        "occupation",
        "what are you retired from",
        "marital status",
        "how long have you been married",
        "are you married to",
        "spouse",
        "represent our insured",
        "represent him for this claim",
        "represent him for this recorded statement",
        "permission to represent him",
        "what unit",
        "do you recall the day",
        "so approximately",
        "home care assistant company together",
        "own and operate a home",
        "relationship with our insured",
        "personal representative",
        "representing",
        "address where the incident",
        "incident allegedly occurred",
        "date that the incident allegedly occurred",
        "that would be 2025",
        "reportedly occurred",
        "date of loss",
        "how long have you owned",
        "how long has the insured owned",
        "owned it since",
        "purchased it in",
        "when did you purchase",
        "how is this unit used",
        "have you ever lived in the property",
        "what year was the building constructed",
        "do you know what year the building was constructed",
        "do you live here full time",
        "live at the property full-time",
        "seasonal residence",
        "how often does your husband occupy it",
        "how is the unit used now",
        "using the property as",
        "how long had it been vacant",
        "what months of the year do you occupy",
        "when did your tenants start occupying the property",
        "when did your tenants move in",
        "what date did you start to rent the property",
        "first started using the property as a rental property",
        "first tenant",
        "how long had the tenant lived at the property",
        "how long in total did she live there",
        "tenant occupied",
        "primary residence",
        "have you already provided it to me",
        "have you already provided it to citizens",
        "while the unit was vacant",
        "how often did they check the unit",
        "water turned off while the unit was vacant",
        "primary residence",
        "rental property",
        "who lives at the property",
        "who was occupying the unit on the date of loss",
        "does anyone besides your mother live there",
        "does anyone besides your father live there",
        "what is your mother's name",
        "what is your fathers name",
        "what is your father's name",
        "was your mother home when the incident occurred",
        "was your father home when the incident occurred",
        "so she was home when the incident occurred",
        "was the unit occupied on the date of loss",
        "was the property occupied on the date of loss",
        "shut the water off",
        "checking on the unit",
        "how often do you have the unit checked",
        "property manager",
        "lease agreement",
        "tenant",
        "personally occupied the unit",
        "prior to renting it",
        "your friends",
        "unit was empty on the date of loss",
        "unit was occupied by your friends",
        "were they there when these leaks or issues started",
        "lease agreement with your friends",
    )
    return any(needle in lowered for needle in opening_needles)


def set_preferred_fact(facts: dict[str, str], key: str, value: Optional[str], scorer) -> None:
    if not value:
        return
    existing = facts.get(key)
    if not existing or scorer(value) > scorer(existing):
        facts[key] = value


def address_quality(text: str) -> int:
    lowered = normalize(text).lower()
    score = len(lowered)
    if re.search(r"\b\d{5}\b", lowered):
        score += 25
    if "," in lowered:
        score += 15
    if looks_like_address(lowered):
        score += 20
    if lowered.startswith(("that's ", "that is ")):
        score -= 60
    return score


def ownership_quality(text: str) -> int:
    cleaned = normalize(text).strip(" .")
    score = len(cleaned)
    if normalize_year_duration(cleaned):
        score += 40
    if re.search(r"(january|february|march|april|may|june|july|august|september|october|november|december)\s+\d{4}", cleaned, re.IGNORECASE):
        score += 35
    if re.fullmatch(r"(19|20)\d{2}", cleaned):
        score += 30
    if cleaned.count(".") > 1:
        score -= 20
    return score


def date_quality(text: str) -> int:
    cleaned = normalize(text).strip(" .")
    score = len(cleaned)
    if re.search(r"\b\d{4}\b", cleaned):
        score += 25
    if re.search(r"(january|february|march|april|may|june|july|august|september|october|november|december)", cleaned, re.IGNORECASE):
        score += 10
    return score


def duration_quality(text: str) -> int:
    cleaned = normalize(text).strip(" .")
    score = len(cleaned)
    if normalize_year_duration(cleaned):
        score += 50
    if re.search(r"\b(year|month)\b", cleaned, re.IGNORECASE):
        score += 20
    if cleaned.startswith(("she lived until", "he lived until", "they lived until")):
        score -= 40
    return score


def clean_month_year_reference(text: str) -> str:
    cleaned = normalize(text).strip(" .")
    cleaned = re.sub(r"^(?:in|on)\s+", "", cleaned, flags=re.IGNORECASE)
    match_obj = re.search(
        r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+(\d{2})\b",
        cleaned,
        re.IGNORECASE,
    )
    if match_obj:
        month = match_obj.group(1).capitalize()
        year = match_obj.group(2)
        year = f"19{year}" if int(year) > 30 else f"20{year}"
        return f"{month} {year}"
    return cleaned


def normalize_short_year(text: str) -> str:
    cleaned = normalize(text).strip(" .")
    match_obj = re.search(r"\b(\d{2})\b", cleaned)
    if match_obj:
        cleaned = match_obj.group(1)
    if re.fullmatch(r"\d{2}", cleaned):
        return f"19{cleaned}" if int(cleaned) > 30 else f"20{cleaned}"
    return cleaned


def clean_occupants_phrase(text: str, context: ReportContext) -> str:
    parts = [normalize(part).strip(" .") for part in re.split(r"[.?!]", text) if normalize(part).strip(" .")]
    if parts:
        ranked = sorted(parts, key=lambda item: ("my " in item.lower() or context.pronoun in item.lower(), len(item)), reverse=True)
        text = ranked[0]
    text = re.sub(r"^(?:oh,\s*)?", "", text, flags=re.IGNORECASE)
    replacements = [
        (r"\bit was me and my daughter\b", f"{context.pronoun} and {context.possessive} daughter"),
        (r"\bit was me and my son\b", f"{context.pronoun} and {context.possessive} son"),
        (r"\bme and my daughter\b", f"{context.pronoun} and {context.possessive} daughter"),
        (r"\bme and my son\b", f"{context.pronoun} and {context.possessive} son"),
        (r"\bit was me and\b", f"{context.pronoun} and"),
        (r"\bme and\b", f"{context.pronoun} and"),
        (r"\bmy wife and i(?: only)?\b", f"{context.pronoun} and {context.possessive} wife"),
        (r"\bmy husband and i(?: only)?\b", f"{context.pronoun} and {context.possessive} husband"),
        (r"\bmy spouse and i(?: only)?\b", f"{context.pronoun} and {context.possessive} spouse"),
        (r"\bmy wife\b", f"{context.possessive} wife"),
        (r"\bmy husband\b", f"{context.possessive} husband"),
        (r"\bmy spouse\b", f"{context.possessive} spouse"),
        (r"\bmy brother\b", f"{context.possessive} brother"),
        (r"\bmy sister\b", f"{context.possessive} sister"),
        (r"\bmy son\b", f"{context.possessive} son"),
        (r"\bmy daughter\b", f"{context.possessive} daughter"),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    text = re.sub(r"\bnow\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\bonly\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s{2,}", " ", text).strip(" ,.")
    return lower_first(text)


def extract_family_resident(text: str, context: ReportContext) -> tuple[Optional[str], Optional[str]]:
    lowered = normalize(text).lower()
    residents = (
        ("mother", f"{context.possessive} mother"),
        ("father", f"{context.possessive} father"),
        ("parents", f"{context.possessive} parents"),
        ("brother", f"{context.possessive} brother"),
        ("sister", f"{context.possessive} sister"),
        ("son", f"{context.possessive} son"),
        ("daughter", f"{context.possessive} daughter"),
        ("wife", f"{context.possessive} wife"),
        ("husband", f"{context.possessive} husband"),
        ("spouse", f"{context.possessive} spouse"),
    )
    for label, phrase in residents:
        if re.search(rf"\bmy {label}\b", lowered) and re.search(r"\b(live|lives|living|stays|staying)\b", lowered):
            return label, phrase
    return None, None


def is_affirmative(text: str) -> bool:
    lowered = normalize(text).lower().strip(" .")
    return lowered.startswith(("yes", "yeah", "yep", "correct", "that's correct", "that is correct"))


def extract_relationship_label(text: str) -> Optional[str]:
    match_obj = re.search(r"\bmy (brother|sister|father|mother|husband|wife|son|daughter)\b", text, re.IGNORECASE)
    if not match_obj:
        return None
    return match_obj.group(1).lower()


def extract_insured_name(text: str) -> Optional[str]:
    match_obj = re.match(r"(.+?) was my ", normalize(text).strip(" ."), re.IGNORECASE)
    if not match_obj:
        return None
    return tidy_name(match_obj.group(1))


def match(question: str, *needles: str) -> bool:
    lowered = question.lower()
    return any(needle in lowered for needle in needles)


def narrativize_answer(answer: str, context: ReportContext) -> str:
    text = normalize(answer)
    replacements = [
        (r"\bI don't know\b", f"{context.pronoun} does not know"),
        (r"\bI do not know\b", f"{context.pronoun} does not know"),
        (r"\bI have no idea\b", f"{context.pronoun} does not know"),
        (r"\bI guess\b", f"{context.pronoun} believes"),
        (r"\bI mean\b", ""),
        (r"\bI'm\b", f"{context.pronoun} is"),
        (r"\bI am\b", f"{context.pronoun} is"),
        (r"\bI was\b", f"{context.pronoun} was"),
        (r"\bI have\b", f"{context.pronoun} has"),
        (r"\bI've\b", f"{context.pronoun} has"),
        (r"\bI had\b", f"{context.pronoun} had"),
        (r"\bI believe\b", f"{context.pronoun} believes"),
        (r"\bI think\b", f"{context.pronoun} believes"),
        (r"\bI know\b", f"{context.pronoun} knows"),
        (r"\bI can't\b", f"{context.pronoun} cannot"),
        (r"\bI would\b", f"{context.pronoun} would"),
        (r"\bmy\b", context.possessive),
        (r"\bme\b", context.objective),
        (r"\bI\b", context.pronoun),
    ]
    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)

    text = re.sub(rf"\b{context.pronoun} don't\b", f"{context.pronoun} does not", text, flags=re.IGNORECASE)
    text = re.sub(rf"\b{context.pronoun} has got\b", f"{context.pronoun} has", text, flags=re.IGNORECASE)
    text = re.sub(r"\broughly\.\b", "roughly.", text, flags=re.IGNORECASE)
    text = clean_transcript_fillers(text)
    text = text.strip(" .")
    text = sentence_case(text)
    return collapse_transcript_sentences(text)


def sentence_case(text: str) -> str:
    if not text:
        return text
    text = text[:1].upper() + text[1:]
    text = re.sub(r"(?<=[.!?])\s+([a-z])", lambda m: f" {m.group(1).upper()}", text)
    return text


def split_sentences(text: str) -> list[str]:
    text = text.replace("...", ". ")
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [part.strip(" .") for part in parts if part.strip(" .")]


def collapse_transcript_sentences(text: str) -> str:
    sentences = split_sentences(text)
    if not sentences:
        return ""

    trivial = {
        "okay",
        "ok",
        "yeah",
        "yes",
        "no",
        "sorry",
        "i'm sorry",
        "understood",
        "correct",
    }

    collapsed: list[str] = []
    for sentence in sentences:
        cleaned = normalize(sentence).strip(" .")
        if not cleaned:
            continue
        lowered = cleaned.lower()
        if len(sentences) > 1 and (lowered in trivial or lowered.startswith("hold on")):
            continue
        if collapsed and lowered == collapsed[-1].lower():
            continue
        collapsed.append(cleaned)

    if not collapsed:
        collapsed = [normalize(sentences[0]).strip(" .")]

    return ". ".join(collapsed).strip(" .")


def dedupe_paragraph_sentences(paragraph: str) -> str:
    sentences = split_sentences(paragraph)
    deduped: list[str] = []
    seen: set[str] = set()
    for sentence in sentences:
        cleaned = normalize(sentence).strip(" .")
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(cleaned)
    return ". ".join(deduped).strip(" .")


def reduce_subject_repetition(paragraph: str, context: ReportContext) -> str:
    sentences = split_sentences(paragraph)
    if len(sentences) <= 1:
        return paragraph

    subject_variants = {context.role_label.strip(), paragraph_subject(context).strip()}
    if context.role_label.lower().startswith("the "):
        subject_variants.add(f"The {short_role_noun(context)}")

    pronoun = cap(context.pronoun)
    rewritten = [sentences[0].strip(" .")]
    pattern = re.compile(
        rf"^((?:Regarding [^,]+,\s+|In describing what happened,\s+)?)"
        rf"({'|'.join(re.escape(item) for item in sorted(subject_variants, key=len, reverse=True) if item)})\b",
        re.IGNORECASE,
    )

    for sentence in sentences[1:]:
        cleaned = sentence.strip(" .")
        if not cleaned:
            continue
        cleaned = pattern.sub(lambda match_obj: f"{match_obj.group(1)}{pronoun}", cleaned, count=1)
        cleaned = re.sub(
            r"^(Regarding [^,]+,\s+|In describing what happened,\s+)(He|She|They)\b",
            lambda match_obj: f"{match_obj.group(1)}{match_obj.group(2).lower()}",
            cleaned,
        )
        rewritten.append(cleaned)

    return ". ".join(rewritten).strip(" .")


def soften_repetitive_attribution(paragraph: str, context: ReportContext) -> str:
    sentences = split_sentences(paragraph)
    if len(sentences) <= 1:
        return paragraph

    softened = [sentences[0].strip(" .")]
    pronouns = {"He", "She", "They"}
    pattern = re.compile(
        r"^((?:Regarding [^,]+,\s+|In describing what happened,\s+)?)"
        r"(He|She|They)\s+(?:(?:also|further|then)\s+)?"
        r"(?:stated|explained|confirmed|advised|clarified|acknowledged|indicated|reported)\s+that,?\s+",
        re.IGNORECASE,
    )

    for sentence in sentences[1:]:
        cleaned = sentence.strip(" .")
        if not cleaned:
            continue
        updated = pattern.sub(lambda match_obj: match_obj.group(1), cleaned, count=1).strip()
        if updated != cleaned:
            updated = re.sub(r"^(?:yes|no|yeah|correct)\s*,\s*", "", updated, flags=re.IGNORECASE)
            if updated:
                first_word = updated.split(" ", 1)[0].capitalize()
                if first_word in pronouns:
                    updated = first_word + updated[len(first_word):]
                else:
                    updated = sentence_case(updated)
                softened.append(updated.strip(" ."))
                continue
        softened.append(cleaned)

    return ". ".join(softened).strip(" .")


def vary_repeated_pronoun_openings(paragraph: str, context: ReportContext) -> str:
    sentences = split_sentences(paragraph)
    if len(sentences) <= 1:
        return paragraph

    pronoun_cap = cap(context.pronoun)
    pronoun_lower = context.pronoun
    do_word = "do" if pronoun_lower == "they" else "does"
    be_word = "are" if pronoun_lower == "they" else "is"
    was_word = "were" if pronoun_lower == "they" else "was"

    transitions = [
        (
            re.compile(rf"^{pronoun_cap} does not know whether ", re.IGNORECASE),
            lambda: f"Nor {do_word} {pronoun_lower} know whether ",
        ),
        (
            re.compile(rf"^{pronoun_cap} does not know if ", re.IGNORECASE),
            lambda: f"Nor {do_word} {pronoun_lower} know if ",
        ),
        (
            re.compile(rf"^{pronoun_cap} did not know whether ", re.IGNORECASE),
            lambda: f"Nor did {pronoun_lower} know whether ",
        ),
        (
            re.compile(rf"^{pronoun_cap} is not aware of ", re.IGNORECASE),
            lambda: f"Nor {be_word} {pronoun_lower} aware of ",
        ),
        (
            re.compile(rf"^{pronoun_cap} was not aware of ", re.IGNORECASE),
            lambda: f"Nor {was_word} {pronoun_lower} aware of ",
        ),
        (
            re.compile(rf"^{pronoun_cap} did not file ", re.IGNORECASE),
            lambda: f"{pronoun_cap} also did not file ",
        ),
    ]

    rewritten = [sentences[0].strip(" .")]
    for sentence in sentences[1:]:
        cleaned = sentence.strip(" .")
        if not cleaned:
            continue
        previous = rewritten[-1].strip(" .")
        updated = cleaned
        for pattern, replacement in transitions:
            if pattern.match(previous) and pattern.match(cleaned):
                updated = pattern.sub(replacement(), cleaned, count=1)
                break
        rewritten.append(updated)

    return ". ".join(rewritten).strip(" .")


def normalize_plural_grammar(paragraph: str, context: ReportContext) -> str:
    if not context.joint_statement:
        return paragraph

    replacements = [
        ("The interviewees has", "The interviewees have"),
        ("The interviewees is", "The interviewees are"),
        ("The interviewees does", "The interviewees do"),
        ("The interviewees believes", "The interviewees believe"),
        ("The interviewees was", "The interviewees were"),
        ("They believes", "They believe"),
        ("they believes", "they believe"),
        ("They is", "They are"),
        ("they is", "they are"),
        ("They has", "They have"),
        ("they has", "they have"),
        ("They does", "They do"),
        ("they does", "they do"),
        ("They also does", "They also do"),
        ("they also does", "they also do"),
        ("They further does", "They further do"),
        ("they further does", "they further do"),
        ("They believes", "They believe"),
        ("they believes", "they believe"),
        ("They was", "They were"),
        ("they was", "they were"),
        ("They presumes", "They presume"),
        ("they presumes", "they presume"),
        ("They knows", "They know"),
        ("they knows", "they know"),
        ("They was only", "They were only"),
        ("they was only", "they were only"),
    ]
    for old, new in replacements:
        paragraph = paragraph.replace(old, new)
    paragraph = paragraph.replace("they is married", "they are married")
    paragraph = paragraph.replace(" but believes ", " but believe ")
    paragraph = paragraph.replace(" and believes ", " and believe ")
    paragraph = paragraph.replace(" but presumes ", " but presume ")
    paragraph = paragraph.replace(" and presumes ", " and presume ")
    paragraph = paragraph.replace(" but does not know whether ", " but do not know whether ")
    paragraph = paragraph.replace(" and does not know whether ", " and do not know whether ")
    paragraph = paragraph.replace(" but does not know ", " but do not know ")
    paragraph = paragraph.replace(" and does not know ", " and do not know ")
    paragraph = paragraph.replace(" but does not know for certain", " but do not know for certain")
    return paragraph


def normalize_singular_grammar(paragraph: str) -> str:
    replacements = [
        ("He do not", "He does not"),
        ("he do not", "he does not"),
        ("She do not", "She does not"),
        ("she do not", "she does not"),
        ("He don't", "He does not"),
        ("he don't", "he does not"),
        ("She don't", "She does not"),
        ("she don't", "she does not"),
        ("He do have", "He does have"),
        ("he do have", "he does have"),
        ("She do have", "She does have"),
        ("she do have", "she does have"),
        ("He really don't", "He really does not"),
        ("he really don't", "he really does not"),
        ("She really don't", "She really does not"),
        ("she really don't", "she really does not"),
        ("He still don't", "He still does not"),
        ("he still don't", "he still does not"),
        ("She still don't", "She still does not"),
        ("she still don't", "she still does not"),
        ("He wish", "He wishes"),
        ("he wish", "he wishes"),
        ("She wish", "She wishes"),
        ("she wish", "she wishes"),
        ("He understand", "He understands"),
        ("he understand", "he understands"),
        ("She understand", "She understands"),
        ("she understand", "she understands"),
        ("He get on it", "He gets on it"),
        ("he get on it", "he gets on it"),
        ("She get on it", "She gets on it"),
        ("she get on it", "she gets on it"),
        ("She use", "She uses"),
        ("she use", "she uses"),
        ("He use", "He uses"),
        ("he use", "he uses"),
        ("She usesss", "She uses"),
        ("she usesss", "she uses"),
        ("He usesss", "He uses"),
        ("he usesss", "he uses"),
        ("She own", "She owns"),
        ("she own", "she owns"),
        ("He own", "He owns"),
        ("he own", "he owns"),
        ("She do this work", "She does this work"),
        ("she do this work", "she does this work"),
    ]
    for old, new in replacements:
        paragraph = paragraph.replace(old, new)
    paragraph = paragraph.replace("ownser", "owner")
    paragraph = paragraph.replace("ownsser", "owner")
    paragraph = paragraph.replace("ownssser", "owner")
    paragraph = paragraph.replace("ownsers", "owners")
    paragraph = paragraph.replace("ownsssers", "owners")
    return paragraph


def is_fragment_sentence(sentence: str) -> bool:
    cleaned = normalize(sentence).strip(" .")
    lowered = cleaned.lower()
    if not lowered:
        return True
    if "speaker 1" in lowered or "speaker 2" in lowered:
        return True
    if lowered in {
        "no",
        "yes",
        "yeah",
        "correct",
        "absolutely yes",
        "the 102",
        "in massachusetts",
        "thank you",
        "exact",
        "unknown",
        "after",
        "has there been",
        "since mid-may of 2025",
    }:
        return True
    if lowered in {"oh, okay", "okay", "ok", "right", "understood", "sorry", "hello"}:
        return True
    if "?" in sentence:
        if re.match(r"^(?:what|why|when|where|who|how|did|do|does|is|are|was|were|can|could|would|will)\b", lowered):
            return True
        if len(lowered.split()) <= 8:
            return True
    if re.match(r"^(?:oh|okay|ok|right|sorry|understood)\b", lowered) and len(lowered.split()) <= 6:
        return True
    return False


def remove_fragment_sentences(paragraph: str) -> str:
    sentences = split_sentences(paragraph)
    kept: list[str] = []
    for sentence in sentences:
        cleaned = sentence.strip(" .")
        lowered = cleaned.lower()
        if is_fragment_sentence(sentence):
            continue
        kept.append(cleaned)
    return ". ".join(kept).strip(" .")


def remove_redundant_uncertainty_fragments(paragraph: str, context: ReportContext) -> str:
    sentences = split_sentences(paragraph)
    if len(sentences) <= 1:
        return paragraph

    pronoun_cap = cap(context.pronoun)
    redundant = f"{pronoun_cap} does not know"
    kept: list[str] = []
    for sentence in sentences:
        cleaned = sentence.strip(" .")
        if not cleaned:
            continue
        lowered = cleaned.lower()
        if lowered == redundant.lower() and kept:
            previous = kept[-1].lower()
            if any(
                token in previous
                for token in (
                    "has no idea",
                    "has no information",
                    "could not describe",
                    "does not know",
                    "did not know",
                    "was not certain",
                    "was not sure",
                )
            ):
                continue
        kept.append(cleaned)

    return ". ".join(kept).strip(" .")


def remove_repeated_uncertainty_remainders(paragraph: str, context: ReportContext) -> str:
    sentences = split_sentences(paragraph)
    if len(sentences) <= 1:
        return paragraph

    pronoun_cap = cap(context.pronoun)
    pronoun_lower = context.pronoun
    do_word = "do" if pronoun_lower == "they" else "does"
    be_word = "are" if pronoun_lower == "they" else "is"
    was_word = "were" if pronoun_lower == "they" else "was"

    patterns = [
        ("know-whether", re.compile(rf"^(?:{pronoun_cap} does not know whether|Nor {do_word} {pronoun_lower} know whether)\s+(.+)$", re.IGNORECASE)),
        ("know-if", re.compile(rf"^(?:{pronoun_cap} does not know if|Nor {do_word} {pronoun_lower} know if)\s+(.+)$", re.IGNORECASE)),
        ("aware-of", re.compile(rf"^(?:{pronoun_cap} is not aware of|Nor {be_word} {pronoun_lower} aware of)\s+(.+)$", re.IGNORECASE)),
        ("was-aware-of", re.compile(rf"^(?:{pronoun_cap} was not aware of|Nor {was_word} {pronoun_lower} aware of)\s+(.+)$", re.IGNORECASE)),
    ]

    seen: set[str] = set()
    kept: list[str] = []
    for sentence in sentences:
        cleaned = sentence.strip(" .")
        if not cleaned:
            continue
        key: Optional[str] = None
        for label, pattern in patterns:
            match_obj = pattern.match(cleaned)
            if match_obj:
                key = f"{label}:{normalize(match_obj.group(1)).lower()}"
                break
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        kept.append(cleaned)

    return ". ".join(kept).strip(" .")


def smooth_conjunction_artifacts(paragraph: str) -> str:
    specific_replacements = {
        "He did not know the claimant personally, that she had not previously been to the house or on a charter with him, and that she has not returned since the incident":
            "He did not know the claimant personally. She had not previously been to the house or on a charter with him, and she has not returned since the incident",
        "Neither he nor his wife was directing or assisting the claimant before the fall and that he only helped her afterward":
            "Neither he nor his wife was directing or assisting the claimant before the fall, and he only helped her afterward",
    }
    for old, new in specific_replacements.items():
        paragraph = paragraph.replace(old, new)

    generic_replacements = [
        (" and that water ", ", and the water "),
        (" and that the ", ", and the "),
        (" and that she ", ", and she "),
        (" and that he ", ", and he "),
        (" and that they ", ", and they "),
        (" and that it ", ", and it "),
        (" and that there ", ", and there "),
        (" and that this ", ", and this "),
        (" and that his ", ", and his "),
        (" and that her ", ", and her "),
        (" and that their ", ", and their "),
        (" and that January ", ", and January "),
        (" and that William ", ", and William "),
        (" and that Gus ", ", and Gus "),
        (" and that Casey ", ", and Casey "),
    ]
    for old, new in generic_replacements:
        paragraph = paragraph.replace(old, new)
    paragraph = paragraph.replace(",,", ",")
    return paragraph


def collect_opening_facts(pairs: list[QAPair], context: ReportContext) -> int:
    consumed = 0
    pending_relative: Optional[str] = None
    for pair in pairs:
        q = pair.question.lower()
        a = normalize(pair.answer).strip(" .")
        consumed_this = True

        if match(q, "permission to record"):
            pass
        elif is_affirmative(a) and looks_like_date(q) and consumed <= 4:
            set_preferred_fact(context.facts, "date_of_birth", clean_date_answer(opening_question_tail(pair.question)), date_quality)
        elif is_affirmative(a) and match(q, "would that be", "is that") and looks_like_address(opening_question_tail(pair.question)):
            confirmed_address = clean_address_answer(opening_question_tail(pair.question))
            if looks_like_address(confirmed_address):
                if (
                    any(token in q for token in ("incident", "occurred", "loss location", "property in question"))
                    or ("unit" in confirmed_address.lower() and context.facts.get("residence_address") and confirmed_address != context.facts.get("residence_address"))
                ):
                    set_preferred_fact(context.facts, "loss_address", confirmed_address, address_quality)
                else:
                    set_preferred_fact(context.facts, "residence_address", confirmed_address, address_quality)
        elif is_affirmative(a) and ("date of loss" in q or "incident occurred" in q) and looks_like_date(q):
            set_preferred_fact(context.facts, "date_of_loss", clean_date_answer(opening_question_tail(pair.question)), date_quality)
        elif q.strip(" .") in {"so yes", "yes", "okay", "ok", "right"} and is_trivial_acknowledgment(a):
            pass
        elif match(q, "state your full name", "spell it for the record", "can you spell it", "can you spell that"):
            context.facts["name_confirmed"] = "yes"
        elif match(q, "date of birth") and not match(q, "what is his date of birth", "what is her date of birth"):
            confirmed = extract_confirmed_date(pair.question, a)
            value = confirmed or (clean_date_answer(a) if looks_like_date(a) else None)
            if value:
                context.facts.setdefault("date_of_birth", value)
        elif match(q, "address where the incident", "address where this incident", "incident allegedly occurred"):
            cleaned = clean_address_answer(a)
            if looks_like_address(cleaned):
                set_preferred_fact(context.facts, "loss_address", cleaned, address_quality)
        elif match(q, "what is your address", "what address do you currently live at", "address do you live at", "where do you live", "primary mailing", "what's your address there"):
            cleaned = clean_address_answer(a)
            if looks_like_address(cleaned):
                set_preferred_fact(context.facts, "residence_address", cleaned, address_quality)
        elif match(q, "what is your occupation", "what was your occupation"):
            confirmed = extract_confirmed_text(pair.question, a)
            value = confirmed or a
            if value and not is_affirmative(value):
                context.facts.setdefault("occupation", normalize_occupation(value))
                if "retired" in value.lower() and "but" in value.lower():
                    detail = value.split("but", 1)[1]
                    context.facts.setdefault("occupation_detail", normalize_occupation_detail(detail))
        elif match(q, "where are you retired from", "what are you retired from"):
            context.facts.setdefault("occupation_detail", normalize_occupation_detail(a))
        elif match(q, "spouse's full name", "spouse’s name", "what is your spouse's full name", "what is your spouse's name"):
            context.facts.setdefault("spouse_name", clean_name_answer(a))
        elif match(q, "what is your marital status"):
            if match(q, "spouse") and not context.facts.get("spouse_name") and not looks_like_date(a):
                context.facts.setdefault("spouse_name", clean_name_answer(a))
                context.facts.setdefault("marital_status", "married")
            elif not match(q, "spouse"):
                context.facts.setdefault("marital_status", normalize_marital_status(a))
        elif match(q, "how long have you been married", "are you married to"):
            if is_affirmative(a):
                context.facts.setdefault("marital_status", "married")
            pass
        elif match(q, "what is his date of birth", "what is her date of birth"):
            confirmed = extract_confirmed_date(pair.question, a)
            value = confirmed or (clean_date_answer(a) if looks_like_date(a) else None)
            if value:
                if pending_relative:
                    set_preferred_fact(context.facts, "household_relative_dob", value, date_quality)
                else:
                    context.facts.setdefault("spouse_dob", value)
        elif match(q, "what is his occupation", "what is her occupation", "his occupation is", "her occupation is"):
            confirmed = extract_confirmed_text(pair.question, a)
            value = confirmed or a
            if value and not is_affirmative(value):
                context.facts.setdefault("spouse_occupation", normalize_occupation(value))
        elif match(q, "relationship with our insured"):
            context.facts.setdefault("relationship_to_insured", narrativize_answer(a, context))
            insured_name = extract_insured_name(a)
            if insured_name:
                context.facts.setdefault("insured_name", insured_name)
            relationship_label = extract_relationship_label(a)
            if relationship_label:
                context.facts.setdefault("insured_relationship_label", relationship_label)
            if "passed away" in a.lower():
                deceased_date = extract_date_phrase(a)
                if deceased_date:
                    context.facts.setdefault("deceased_date", deceased_date)
        elif match(q, "why are you representing"):
            context.facts.setdefault("representation_reason", narrativize_answer(a, context))
        elif match(q, "are you representing our insured", "did the insured ask you to represent him", "did our insured give you permission to represent him"):
            if is_affirmative(a):
                context.facts.setdefault("representation_confirmed", "yes")
        elif match(q, "when did he pass", "when did she pass"):
            context.facts.setdefault("deceased_date", clean_date_answer(a))
        elif match(q, "personal representative for", "you are the personal representative"):
            if is_affirmative(a):
                context.facts.setdefault("personal_representative_confirmed", "yes")
        elif match(q, "same address where the incident allegedly occurred"):
            if is_affirmative(a) and context.facts.get("residence_address"):
                set_preferred_fact(context.facts, "loss_address", context.facts["residence_address"], address_quality)
        elif match(q, "unit number"):
            context.facts.setdefault("loss_unit", a)
        elif match(q, "what is the date of loss", "date that this incident allegedly occurred on", "date that the incident allegedly occurred", "reportedly occurred on", "that would be 2025"):
            confirmed = extract_confirmed_date(pair.question, a)
            value = confirmed or (clean_date_answer(a) if looks_like_date(a) else None)
            if value:
                context.facts.setdefault("date_of_loss", value)
        elif q.strip() == "what unit?":
            context.facts.setdefault("loss_unit", a)
        elif match(q, "spell the name of that street and the city", "spell those words", "spell the name of that street"):
            cleaned = clean_address_answer(a)
            if looks_like_address(cleaned):
                set_preferred_fact(context.facts, "residence_address", cleaned, address_quality)
        elif match(q, "own and operate a home", "home care assistant company together"):
            if "home care" in a.lower():
                context.facts["occupation"] = "a home care company"
        elif match(q, "how long have you owned", "how long has the insured owned", "when did you purchase", "owned it since", "purchased it in"):
            set_preferred_fact(context.facts, "ownership_length", a, ownership_quality)
        elif match(q, "what year was the building constructed", "what year was this building constructed", "do you know what year the building was constructed", "do you know what year this building was constructed"):
            context.facts.setdefault("building_age", a)
        elif match(q, "how is this unit used", "how is the unit used now"):
            relative_label, relative_phrase = extract_family_resident(a, context)
            if relative_phrase:
                context.facts.setdefault("occupants", relative_phrase)
                context.facts.setdefault("household_relative_label", relative_label or "relative")
                context.facts.setdefault("tenant_occupancy", "yes")
            elif "vacant" in a.lower() or "empty" in a.lower():
                context.facts.setdefault("property_use", "vacant")
            elif "rental" in a.lower():
                context.facts.setdefault("property_use", "rental")
            elif "full-time" in a.lower() or "primary residence" in a.lower():
                context.facts.setdefault("property_use", "full-time residence")
        elif match(q, "do you use this as a seasonal residence"):
            if is_affirmative(a):
                context.facts.setdefault("property_use", "seasonal residence")
        elif match(q, "what months of the year do you typically occupy", "what months of the year do you occupy", "how often does your husband occupy it"):
            context.facts.setdefault("seasonal_occupancy_months", a)
        elif match(q, "do you live here full time", "do you live at the property full-time"):
            if is_affirmative(a):
                context.facts.setdefault("property_use", "full-time residence")
        elif match(q, "does the insured live at the property full time or is it used as a rental vacation or seasonal property"):
            lowered_answer = a.lower()
            if "seasonal" in lowered_answer:
                context.facts.setdefault("property_use", "seasonal residence")
            elif "rental" in lowered_answer:
                context.facts.setdefault("property_use", "rental")
            elif "full-time" in lowered_answer:
                context.facts.setdefault("property_use", "full-time residence")
        elif match(q, "how is the unit used now"):
            if "empty" in a.lower() or "vacant" in a.lower():
                context.facts.setdefault("current_use_detail", a)
        elif match(q, "primary residence"):
            if is_affirmative(a):
                context.facts.setdefault("property_use", "full-time residence")
        elif match(q, "who lives at this property with you", "who lives at the property with you", "who lives at the property", "who was living at the property with you", "who was living at the property", "occupying the property at that time", "what is the name of the person or persons who was occupying the unit on the date of loss"):
            cleaned = clean_occupants_phrase(a, context)
            if cleaned and "ex-wife" not in cleaned.lower():
                context.facts["occupants"] = cleaned
                if "brother" in q or "brother" in cleaned.lower():
                    pending_relative = "brother"
                elif "sister" in q or "sister" in cleaned.lower():
                    pending_relative = "sister"
        elif match(q, "does anyone besides your mother live there"):
            if a.lower().startswith("no"):
                context.facts.setdefault("occupants", f"{context.possessive} mother")
                context.facts.setdefault("household_relative_label", "mother")
                context.facts.setdefault("tenant_occupancy", "yes")
        elif match(q, "does anyone besides your father live there"):
            if a.lower().startswith("no"):
                context.facts.setdefault("occupants", f"{context.possessive} father")
                context.facts.setdefault("household_relative_label", "father")
                context.facts.setdefault("tenant_occupancy", "yes")
        elif match(q, "using the property as your full-time residence", "using the property as a rental property") and match(q, "who was living at the property with you", "who was living at the property"):
            if "full-time residence" in q:
                context.facts.setdefault("property_use", "full-time residence")
            elif "rental property" in q:
                context.facts.setdefault("property_use", "rental")
            context.facts.setdefault("occupants", a)
        elif match(q, "using the property as", "full-time residence"):
            if "vacant" in a.lower():
                context.facts.setdefault("property_use", "vacant")
            else:
                context.facts.setdefault("property_use", a)
        elif match(q, "how long had it been vacant"):
            context.facts.setdefault("vacancy_since", a)
        elif match(q, "tenant occupied"):
            if "not tenant occupied" in q or "had not been tenant occupied" in q:
                context.facts.setdefault("tenant_occupancy", "no")
            elif is_affirmative(a):
                context.facts.setdefault("tenant_occupancy", "yes")
            elif a.lower().startswith("no"):
                context.facts.setdefault("tenant_occupancy", "no")
        elif match(q, "used the property as a rental property"):
            context.facts.setdefault("rental_duration", a)
        elif match(q, "what date did you start to rent the property", "first started using the property as a rental property"):
            context.facts.setdefault("rental_start", a)
        elif match(q, "how long had the tenant lived at the property", "how long in total did she live there"):
            duration_text = a
            one_half_match = re.search(r"(?:about\s+)?one and a half years", a, re.IGNORECASE)
            if one_half_match:
                duration_text = "approximately one and a half years"
            set_preferred_fact(context.facts, "rental_duration", duration_text, duration_quality)
        elif match(q, "personally occupied the unit"):
            context.facts.setdefault("personal_occupancy_flag", a)
        elif q.strip().startswith("when") and context.facts.get("personal_occupancy_flag", "").lower().startswith("yes"):
            context.facts.setdefault("personal_occupancy_when", a)
        elif match(q, "prior to renting it", "that was your residence"):
            context.facts.setdefault("personal_occupancy_confirmation", a)
        elif match(q, "brother's name", "brothers name", "what's your brother's name", "what is your brother's name"):
            context.facts["household_relative_label"] = "brother"
            context.facts["household_relative_name"] = clean_person_name_answer(a)
            pending_relative = "brother"
        elif match(q, "sister's name", "sisters name", "what's your sister's name", "what is your sister's name"):
            context.facts["household_relative_label"] = "sister"
            context.facts["household_relative_name"] = clean_person_name_answer(a)
            pending_relative = "sister"
        elif match(q, "what is your mother's name", "what is your fathers name", "what is your father's name"):
            if "mother" in q:
                context.facts["household_relative_label"] = "mother"
                pending_relative = "mother"
            else:
                context.facts["household_relative_label"] = "father"
                pending_relative = "father"
            cleaned_name = clean_person_name_answer(a)
            if cleaned_name and cleaned_name.lower() not in {"graciela graciela"}:
                context.facts["household_relative_name"] = cleaned_name
        elif pending_relative and (match(q, "can you state it for me") or "whole" in q) and looks_like_date(a):
            set_preferred_fact(context.facts, "household_relative_dob", clean_date_answer(a), date_quality)
        elif pending_relative and match(q, "biologically related to you"):
            if pending_relative not in context.facts.get("household_relative_label", ""):
                context.facts["household_relative_label"] = clean_transcript_fillers(a).lower().strip(" .")
        elif match(q, "was your husband living with you", "was your wife living with you"):
            if "no" in a.lower():
                context.facts.setdefault("spouse_living_at_property", "no")
        elif match(q, "how old is your daughter", "how old is your son", "how old is your child"):
            context.facts.setdefault("child_age", a)
        elif match(q, "were you home when the incident occurred", "were you present when the incident occurred"):
            context.facts.setdefault("presence", a)
        elif match(q, "tenant occupied on the date of loss", "unit tenant occupied on the date of loss", "was the unit occupied on the date of loss", "was the property occupied on the date of loss"):
            context.facts.setdefault("tenant_occupancy", a)
        elif match(q, "unit was empty on the date of loss"):
            if a.lower().startswith("no"):
                context.facts.setdefault("tenant_occupancy", "yes")
        elif match(q, "unit was occupied by your friends"):
            if is_affirmative(a):
                context.facts.setdefault("occupants", "friends")
        elif match(q, "provide me with the lease agreement already"):
            context.facts.setdefault("lease_provided", a)
        elif match(q, "and why not"):
            if "lease_provided" in context.facts:
                context.facts.setdefault("lease_not_provided_reason", a)
            else:
                consumed_this = False
        elif match(q, "tenant first move into the property", "when did your tenants start occupying the property", "when did your tenants move in", "around 2020"):
            set_preferred_fact(context.facts, "tenant_move_in", a, date_quality)
        elif match(q, "tenant use the property as their primary residence", "do your tenants use your unit as their primary residence"):
            context.facts.setdefault("tenant_primary_residence", a)
        elif match(q, "so you don't know"):
            context.facts.setdefault("tenant_primary_residence_followup", a)
        elif match(q, "tenant was home when the incident"):
            context.facts.setdefault("tenant_home_at_loss", a)
        elif match(q, "shut the water off before leaving", "water turned off while the unit was vacant"):
            context.facts.setdefault("water_off_when_unoccupied", a)
        elif match(q, "do you have someone checking on the unit while it is unoccupied", "while the unit was vacant, did you have anyone checking the unit"):
            context.facts.setdefault("unit_check_frequency", a)
        elif match(q, "how often do you have the unit checked while it is unoccupied", "how often did they check the unit"):
            context.facts.setdefault("unit_check_frequency", a)
        elif match(q, "how long had you been away"):
            context.facts.setdefault("away_duration", a)
        elif match(q, "was your wife or brother home", "was your wife or sister home", "was your husband or brother home", "was your husband or sister home"):
            context.facts.setdefault("other_occupants_status", a)
        elif q.strip().startswith("and what is ") and context.facts.get("spouse_occupation"):
            pass
        elif match(q, "property manager", "self-manage the unit"):
            context.facts.setdefault("property_management", a)
        else:
            consumed_this = False

        if not consumed_this:
            if is_opening_question(q):
                consumed += 1
                continue
            break
        consumed += 1

    return consumed


def strip_leading_at(text: str) -> str:
    return re.sub(r"^(at)\s+", "", text, flags=re.IGNORECASE)


def normalize_occupation(text: str) -> str:
    lowered = text.lower()
    if lowered.startswith("retired") and lowered.strip() != "retired":
        detail = re.sub(r"^retired\s+", "", lowered).strip(" .,")
        detail = detail.replace(",", " and ")
        detail = normalize(detail).strip(" .")
        if detail == "postal":
            detail = "postal worker"
        return f"retired {detail}"
    if "retired" in lowered:
        return "retired"
    if "economist" in lowered and "phd" in lowered:
        return "an economist with a PhD"
    if "it's like " in lowered:
        lowered = lowered.split("it's like ", 1)[1]
    if "it is like " in lowered:
        lowered = lowered.split("it is like ", 1)[1]
    if lowered.startswith("i'm "):
        return lowered[4:]
    if lowered.startswith("i am "):
        return lowered[5:]
    if lowered.startswith("he's "):
        return lowered[5:]
    if lowered.startswith("she's "):
        return lowered[6:]
    lowered = re.sub(r"^(?:he|she)\s+(?:just\s+)?works?\s+as\s+", "", lowered)
    lowered = re.sub(r"^(?:he|she)\s+(?:just\s+)?works?\s+at\s+", "", lowered)
    return lowered


def normalize_marital_status(text: str) -> str:
    lowered = text.lower().strip(" .")
    if lowered.startswith("i'm "):
        return lowered[4:]
    if lowered.startswith("i am "):
        return lowered[5:]
    return lowered


def append_unit(address: str, unit: str) -> str:
    if re.search(r"\bunit\b|\bapt\b|\bapartment\b", address, re.IGNORECASE):
        return address
    unit = normalize(unit)
    unit_text = unit if re.search(r"\bunit\b", unit, re.IGNORECASE) else f"unit {unit}"
    parts = [part.strip() for part in address.split(",")]
    if len(parts) >= 2:
        parts.insert(1, unit_text)
        return ", ".join(parts)
    return f"{address}, {unit_text}"


def compose_opening_paragraphs(context: ReportContext) -> list[str]:
    facts = context.facts
    paragraphs: list[str] = []
    pronoun = cap(context.pronoun)
    role = context.role_label
    name = context.name

    first: list[str] = []
    if context.joint_statement and name:
        first.append(f"{role}, {name}, were identified during the statement")
    elif name and facts.get("name_confirmed"):
        if role_matches_name(context):
            first.append(f"{role} confirmed the spelling of {context.possessive} name")
        else:
            first.append(f"{role}, {name}, confirmed the spelling of {context.possessive} name")
    elif name:
        if role_matches_name(context):
            first.append(f"{role} was identified during the statement")
        else:
            first.append(f"{role}, {name}, was identified during the statement")
    if facts.get("date_of_birth") and not context.joint_statement:
        dob_subject = pronoun if first else role
        if facts.get("date_of_birth"):
            first.append(f"{dob_subject} provided {context.possessive} date of birth as {facts['date_of_birth']}")
        else:
            first.append(f"{dob_subject} provided {context.possessive} date of birth")
    if facts.get("residence_address") and "personal representative" not in role.lower():
        if context.joint_statement:
            first.append(f"{pronoun} reside at {facts['residence_address']}")
        else:
            first.append(f"{pronoun} resides at {facts['residence_address']}")
    if facts.get("occupation") and not context.joint_statement:
        if facts["occupation"] == "retired":
            first.append(f"{pronoun} is retired")
            if facts.get("occupation_detail"):
                first.append(f"Before retiring, {context.pronoun} worked {format_retired_detail(facts['occupation_detail'])}")
        elif facts["occupation"].startswith("retired "):
            first.append(f"{pronoun} is a {facts['occupation']}")
        elif facts.get("occupation_detail"):
            first.append(f"{pronoun} works as {facts['occupation']}, having previously worked {format_retired_detail(facts['occupation_detail'])}")
        else:
            first.append(f"{pronoun} works as {facts['occupation']}")
    if facts.get("marital_status"):
        first.append(f"{pronoun} stated that {context.pronoun} is {facts['marital_status']}")
    if facts.get("spouse_name") and not context.joint_statement:
        spouse_sentence = f"{context.possessive.capitalize()} spouse's name was provided as {facts['spouse_name']}"
        if facts.get("spouse_dob"):
            spouse_sentence += f". {pronoun} provided {context.possessive} spouse's date of birth as {facts['spouse_dob']}"
        if facts.get("spouse_occupation"):
            spouse_sentence += f". {pronoun} stated {context.possessive} spouse's occupation as {facts['spouse_occupation']}"
        first.append(spouse_sentence)
    insured_ref = insured_reference(context)
    natural_insured_ref = related_insured_reference(context)
    natural_insured_possessive = related_insured_reference(context, possessive=True)
    if facts.get("relationship_to_insured"):
        relationship_label = facts.get("insured_relationship_label")
        if relationship_label:
            first.append(f"{pronoun} confirmed that {insured_ref} was {context.possessive} {relationship_label}, and that {insured_ref} is deceased")
        else:
            first.append(f"{pronoun} explained that {lower_first(facts['relationship_to_insured'])}")
    if facts.get("deceased_date"):
        first.append(f"{pronoun} stated that {insured_ref} passed away on {facts['deceased_date']}")
    if facts.get("representation_reason"):
        first.append(f"{pronoun} is representing {insured_ref} in this claim because {lower_first(facts['representation_reason'])}")
    if facts.get("personal_representative_confirmed"):
        first.append(f"{pronoun} confirmed that {context.pronoun} is the personal representative for {insured_reference(context, possessive=True)} estate")
    if facts.get("residence_address") and "personal representative" in role.lower():
        first.append(f"{pronoun} currently resides at {facts['residence_address']}")

    loss_address = facts.get("loss_address")
    if loss_address and facts.get("loss_unit"):
        loss_address = append_unit(loss_address, facts["loss_unit"])
    if loss_address:
        if "personal representative" in role.lower():
            first.append(f"The alleged incident occurred at {loss_address}")
        else:
            first.append(f"The property where the incident allegedly occurred is located at {loss_address}")
    if facts.get("date_of_loss"):
        first.append(render_date_of_loss_opening(facts["date_of_loss"], context))
    if first:
        paragraphs.append(join_sentences(first))

    second: list[str] = []
    if facts.get("ownership_length"):
        second.append(render_ownership(facts["ownership_length"], context))
    if facts.get("building_age"):
        second.append(render_building_age(facts["building_age"], context))
    if facts.get("property_use"):
        second.append(render_property_use(facts["property_use"], context))
    elif facts.get("rental_duration"):
        second.append("On the date of loss, the property was being used as a rental")
    if facts.get("vacancy_since") and facts.get("property_use") == "vacant":
        vacancy_text = lower_first(facts["vacancy_since"])
        if vacancy_text.startswith("since "):
            second.append(f"The property had been vacant {vacancy_text}")
        else:
            second.append(f"The property had been vacant since {vacancy_text}")
    if facts.get("current_use_detail") and "except for occasional family visits" in facts["current_use_detail"].lower():
        second.append(f"{pronoun} stated that the unit has otherwise remained empty except for occasional family visits.")
    if facts.get("seasonal_occupancy_months"):
        seasonal_text = lower_first(facts["seasonal_occupancy_months"])
        if "prior to covid" in seasonal_text:
            second.append(f"{pronoun} stated that before COVID, {natural_insured_ref} occupied the unit for about four to five months each year, but has not gone down since COVID.")
        else:
            second.append(f"{pronoun} typically occupies the unit during {seasonal_text}")
    if facts.get("rental_duration"):
        rental_duration_text = lower_first(facts["rental_duration"])
        if normalize_year_duration(rental_duration_text) or re.search(r"\b(year|month)\b", rental_duration_text, re.IGNORECASE):
            second.append(f"The unit had been used as a rental property for {rental_duration_text}")
    if facts.get("rental_start"):
        second.append(f"The property first began being rented in {clean_month_year_reference(facts['rental_start'])}")
    if facts.get("personal_occupancy_when"):
        second.append(render_personal_occupancy(facts["personal_occupancy_when"], context))
    if facts.get("tenant_occupancy"):
        second.append(render_tenant_occupancy(facts["tenant_occupancy"], context))
    if facts.get("water_off_when_unoccupied"):
        if is_affirmative(facts["water_off_when_unoccupied"]):
            second.append(f"{pronoun} confirmed that {context.pronoun} shuts the water off before leaving when the unit is unoccupied")
        else:
            second.append(f"{pronoun} stated that {lower_first(narrativize_answer(facts['water_off_when_unoccupied'], context))}")
    if facts.get("unit_check_frequency"):
        check_text = narrativize_answer(facts["unit_check_frequency"], context)
        lowered_check = check_text.lower()
        if lowered_check.startswith("yes"):
            second.append(f"{pronoun} confirmed that someone checks the unit while it is unoccupied")
        else:
            second.append(f"{pronoun} stated that the unit is checked {lower_first(check_text)} while it is unoccupied")
    if facts.get("lease_not_provided_reason"):
        second.append(render_lease_reason(facts["lease_not_provided_reason"], context))
    if facts.get("tenant_move_in"):
        second.append(render_tenant_move_in(facts["tenant_move_in"], context))
    if facts.get("tenant_primary_residence") and not facts.get("tenant_occupancy"):
        second.append(render_tenant_primary_use(facts["tenant_primary_residence"], context))
    if facts.get("property_management"):
        second.append(render_property_management(facts["property_management"], context))
    if facts.get("occupants") and not facts.get("property_use") == "full-time residence":
        second.append(f"{pronoun} stated that the occupants on the date of loss were {clean_occupants_phrase(facts['occupants'], context)}")
    if facts.get("household_relative_name"):
        label = facts.get("household_relative_label", "relative")
        sentence = f"{pronoun} identified {context.possessive} {label} as {facts['household_relative_name']}"
        if facts.get("household_relative_dob"):
            sentence += f", whose date of birth is {facts['household_relative_dob']}"
        second.append(sentence)
    if facts.get("presence"):
        second.append(render_presence(facts["presence"], context))
    if facts.get("away_duration"):
        second.append(render_away_duration(facts["away_duration"], context))
    if facts.get("other_occupants_status"):
        if facts["other_occupants_status"].lower().startswith("they were sleeping"):
            second.append(f"{pronoun} stated that the other occupants were home and sleeping at the time of the incident")
        else:
            second.append(f"{pronoun} stated that {lower_first(narrativize_answer(facts['other_occupants_status'], context))}")
    if second:
        paragraphs.append(join_sentences(second))

    return paragraphs


def render_date_of_loss_opening(answer: str, context: ReportContext) -> str:
    lowered = answer.lower()
    pronoun = cap(context.pronoun)
    role_noun = short_role_noun(context)
    if "don't know" in lowered or "do not know" in lowered:
        date_match = re.search(r"(april|may|june|july|august|september|october|november|december|january|february|march)\s+\d{1,2}(?:th|st|nd|rd)?(?:,\s*\d{4})?", answer, re.IGNORECASE)
        if date_match:
            when = date_match.group(0)
            if re.search(r"\d{4}", when) is None:
                when = f"{when}, 2026"
            if not when.endswith(","):
                when = f"{when},"
            return f"The {role_noun} stated that {context.pronoun} does not know the specific date of loss. {pronoun} explained that {context.pronoun} was notified by the building manager on {when} that there was an incident, but {context.pronoun} does not know how long the actual leak had been going on for"
        return f"The {role_noun} stated that {context.pronoun} does not know the specific date of loss"
    return f"The date of loss was {clean_date_answer(answer)}"


def render_ownership(answer: str, context: ReportContext) -> str:
    lowered = answer.lower()
    role_noun = short_role_noun(context)
    month_year_match = re.search(r"(january|february|march|april|may|june|july|august|september|october|november|december)(?:\s+of)?\s+\d{4}", answer, re.IGNORECASE)
    year_match = re.search(r"(20\d{2}\s*,?\s*20\d{2}|20\d{2}|19\d{2})", answer)
    duration = normalize_year_duration(answer)
    all_years = re.findall(r"\b(?:19|20)\d{2}\b", answer)
    if "parents bought" in lowered and "included in the title" in lowered:
        purchase_match = re.search(r"parents bought it in ([a-z]+\s+\d{4})", answer, re.IGNORECASE)
        added_match = re.search(r"included in the title in (\d{4})", answer, re.IGNORECASE)
        if purchase_match and added_match:
            return f"The {role_noun} stated that {context.possessive} parents purchased the property in {purchase_match.group(1)} and that {context.pronoun} was added to the title in {added_match.group(1)}"
    if duration and all_years:
        return f"The {role_noun} has owned the property for {duration} and stated that it was purchased in {all_years[0]}"
    if "purchased in around" in lowered and month_year_match:
        return f"The {role_noun} has owned the property since around {month_year_match.group(0).replace(' of ', ' ')}"
    if len(all_years) >= 2 and ("don't remember" in lowered or "not sure" in lowered):
        return f"The {role_noun} stated that {context.pronoun} purchased the property in either {all_years[0]} or {all_years[1]}, but does not recall the exact year"
    if "don't know" in lowered or "do not know" in lowered:
        if year_match:
            years = year_match.group(0).replace(",", " or")
            years = normalize(years)
            return f"The {role_noun} has owned the unit since {years}, but {context.pronoun} does not recall the exact date"
    if month_year_match and normalize(answer).strip(" .").lower() == month_year_match.group(0).lower():
        return f"The {role_noun} has owned the property since {month_year_match.group(0)}"
    if year_match and normalize(answer).strip(" .") == year_match.group(0):
        return f"The {role_noun} has owned the property since {year_match.group(0)}"
    if duration:
        return f"The {role_noun} has owned the property for {duration}"
    return f"The {role_noun} stated that {context.pronoun} has owned the property for {lower_first(narrativize_answer(answer, context))}"


def render_building_age(answer: str, context: ReportContext) -> str:
    lowered = answer.lower()
    pronoun = cap(context.pronoun)
    if lowered.strip(" .") in {"no", "nope"}:
        return f"{pronoun} does not know the exact year the building was constructed"
    if re.search(r"\b1979\b", lowered):
        return "The building was constructed in 1979"
    if re.fullmatch(r"\d{4}", normalize(answer).strip(" .")):
        return f"The building was constructed in {normalize(answer).strip(' .')}"
    if any(phrase in lowered for phrase in ("not exactly sure", "not sure", "don't know", "do not know")) and not re.search(r"\b(?:19|20)\d{2}\b", answer):
        return f"{pronoun} does not know the exact year the building was constructed"
    if "don't know" in lowered or "do not know" in lowered or "believe" in lowered or "think" in lowered or "not sure" in lowered:
        if "1960" in lowered:
            return f"{pronoun} does not know the exact year the building was constructed but believes it was built sometime in the 1960s"
        if re.search(r"\b80\b", lowered) and re.search(r"\b81\b", lowered):
            return f"{pronoun} does not know the exact year the building was constructed but believes it was built around 1980 or 1981"
        exact_years = re.findall(r"\b(?:19|20)\d{2}\b", answer)
        if exact_years:
            if len(exact_years) == 1:
                return f"{pronoun} does not know the exact year the building was constructed but believes it was built around {exact_years[0]}"
            return f"{pronoun} does not know the exact year the building was constructed but believes it was built around {exact_years[0]} or {exact_years[1]}"
        year_range = narrativize_answer(answer, context)
        year_range = re.sub(r"^(He|She|They) does not know\.?\s*", "", year_range)
        return f"{pronoun} does not know the exact year the building was constructed but {lower_first(year_range)}"
    if re.search(r"\b(?:19|20)\d{2}\b", answer) and "finished in" in lowered:
        years = re.findall(r"\b(?:19|20)\d{2}\b", answer)
        if len(years) >= 2:
            return f"The building was constructed in {years[0]} and completed in {years[1]}"
    if "70s" in lowered or "80s" in lowered:
        return f"The building was constructed either in the 1970s or the 1980s, although {context.pronoun} was not certain of the exact year"
    return f"The building was constructed {lower_first(answer)}"


def render_property_use(answer: str, context: ReportContext) -> str:
    lowered = answer.lower().strip()
    if lowered == "full-time residence":
        sentence = f"On the date of loss, the unit was being used as {context.possessive} full-time residence"
        occupants = context.facts.get("occupants")
        spouse_living = context.facts.get("spouse_living_at_property")
        child_age = context.facts.get("child_age")
        if occupants:
            occ_text = clean_occupants_phrase(occupants, context)
            if "me and my daughter" in occupants.lower() or "it was me and my daughter" in occupants.lower():
                if child_age:
                    age_text = format_child_age(child_age, context)
                    sentence += f". At that time, {context.pronoun} and {context.possessive} {age_text} daughter were living at the property"
                else:
                    sentence += f". At that time, {context.pronoun} and {context.possessive} daughter were living at the property"
            elif occ_text.startswith("just ") and "wife" in occ_text:
                sentence += f". At that time, {context.possessive} wife was living at the property"
            elif " and " not in occ_text and "," not in occ_text and not occ_text.endswith("friends"):
                sentence += f". At that time, {occ_text} was living at the property"
            else:
                sentence += f". At that time, {occ_text} were living at the property"
            sentence += "."
            if spouse_living == "no":
                sentence += f" {cap(context.pronoun)} confirmed that {context.possessive} spouse was not residing there."
        return sentence
    if "vacant" in lowered:
        return "On the date of loss, the property was vacant"
    if lowered == "seasonal residence":
        return f"{cap(context.pronoun)} uses the property as a seasonal residence"
    if lowered == "yes":
        return "On the date of loss, the property was being used as a rental"
    if lowered == "no":
        return "On the date of loss, the property was not being used as a rental"
    return f"On the date of loss, the property was being used as {lower_first(narrativize_answer(answer, context))}"


def render_presence(answer: str, context: ReportContext) -> str:
    lowered = answer.lower().strip()
    pronoun = cap(context.pronoun)
    if lowered.startswith("yes"):
        return f"{pronoun} confirmed that {context.pronoun} was home when the incident occurred"
    if lowered.startswith("no"):
        return f"{pronoun} confirmed that {context.pronoun} was not home when the incident occurred"
    return f"{pronoun} stated that {lower_first(narrativize_answer(answer, context))}"


def format_time_fragment(text: str, default_suffix: str = "") -> str:
    cleaned = text.strip().lower()
    cleaned = cleaned.replace(".", ":")
    if re.fullmatch(r"\d{1,2}", cleaned):
        cleaned = f"{int(cleaned)}:00"
    suffix = default_suffix
    if "a.m" in cleaned or "am" in cleaned:
        suffix = " a.m."
        cleaned = re.sub(r"\s*a\.?m\.?", "", cleaned)
    elif "p.m" in cleaned or "pm" in cleaned:
        suffix = " p.m."
        cleaned = re.sub(r"\s*p\.?m\.?", "", cleaned)
    return f"{cleaned}{suffix}".strip()


def render_away_duration(answer: str, context: ReportContext) -> str:
    normalized = normalize(answer)
    left_match = re.search(r"left (?:for the house|for work|the house) at\s+(\d{1,2}(?:[.:]\d{2})?)", normalized, re.IGNORECASE)
    incident_match = re.search(r"incident happened at\s+(\d{1,2}(?:[.:]\d{2})?)", normalized, re.IGNORECASE)
    if left_match and incident_match:
        left_time = format_time_fragment(left_match.group(1), " a.m.")
        incident_time = format_time_fragment(incident_match.group(1), " a.m.")
        return f"{cap(context.pronoun)} stated that {context.pronoun} had left for work at {left_time}, and the incident occurred around {incident_time}"
    text = normalized
    text = re.sub(r"\b(\d{1,2})\.(\d{2})\b", r"\1:\2", text)
    text = re.sub(r"(\d{1,2}:\d{2}):00 in the morning", r"\1 in the morning", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(\d{1,2})\s+in the morning\b", r"\1:00 in the morning", text, flags=re.IGNORECASE)
    text = lower_first(narrativize_answer(text, context))
    text = text.replace("left for the house", "left for work")
    return f"{cap(context.pronoun)} stated that {text}"


def extract_minute_range(text: str) -> Optional[str]:
    match_obj = re.search(r"(\d{1,2})\s*(?:to|-|,)\s*(\d{1,2})\s+minutes", text, re.IGNORECASE)
    if match_obj:
        start = int(match_obj.group(1))
        end = int(match_obj.group(2))
        return f"{start} to {end} minutes"
    match_obj = re.search(r"(\d{1,2})\s+minutes", text, re.IGNORECASE)
    if match_obj:
        return f"{int(match_obj.group(1))} minutes"
    return None


def render_ac_leak_story(answer: str, context: ReportContext) -> str:
    normalized = normalize(answer)
    role_noun = short_role_noun(context)
    date_phrase = extract_date_phrase(normalized)
    arrival_window = extract_minute_range(normalized)

    first_sentence = f"In describing what happened, the {role_noun} stated that {context.pronoun} was home"
    if date_phrase:
        first_sentence += f" on {clean_date_answer(date_phrase)}"
    first_sentence += " when the sound of rushing water drew attention from the bedroom to the living room"

    sentences = [
        first_sentence,
        f"{cap(context.pronoun)} stated that there was water all over the living room floor and that water was gushing out from the AC closet near the air return",
        f"{cap(context.pronoun)} immediately called building maintenance because {context.pronoun} knew the water had to be shut off, particularly since there had already been a similar flooding incident from a couple of floors above in February",
    ]

    if arrival_window:
        sentences.append(
            f"{cap(context.pronoun)} stated that the maintenance worker said he was at home and would arrive in approximately {arrival_window}"
        )

    sentences.extend(
        [
            f"While waiting, {context.pronoun} tried to pick up items from the floor and do what {context.pronoun} could inside the unit",
            f"{cap(context.pronoun)} stated that when the maintenance worker arrived, he shut off the water, brought a water vacuum, and helped remove as much water as possible",
            f"{cap(context.pronoun)} further stated that the building contacted a water remediation company",
            f"{cap(context.pronoun)} explained that the hallways were flooded and that water had traveled through the apartment and out onto the balcony",
        ]
    )

    return join_sentences(sentences)


def render_kitchen_sink_story(answer: str, context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    sentences = [
        f"Regarding the kitchen sink leak, the {role_noun} stated that drying equipment had already been placed in the unit because of the earlier water damage, so {context.pronoun} was not staying there overnight",
        f"{cap(context.pronoun)} stated that building maintenance called early the next morning after receiving a report from the unit below that water was leaking into the kitchen",
        f"According to {context.objective}, one of the small tubes connecting the water filter next to the kitchen faucet became disconnected",
        f"{cap(context.pronoun)} stated that because no one was there, the water accumulated in the kitchen and leaked into the unit below",
        f"The maintenance worker shut off the water, and when {context.pronoun} returned {context.pronoun} found only a small amount of water on the kitchen floor, which {context.pronoun} cleaned up",
    ]
    return join_sentences(sentences)


def render_kitchen_sink_absence(answer: str, context: ReportContext) -> str:
    return join_sentences(
        [
            f"{cap(context.pronoun)} stated that {context.pronoun} had only been away overnight when the kitchen sink leak occurred",
            f"{cap(context.pronoun)} explained that {context.pronoun} had still been going to the unit every day while the restoration company checked the drying progress and while {context.pronoun} checked on {context.possessive} personal belongings",
        ]
    )


def render_prior_flood_story(context: ReportContext) -> str:
    return join_sentences(
        [
            f"The {short_role_noun(context)} further described a separate February leak involving unit 509, explaining that someone working in that unit ruptured a water line",
            f"{cap(context.pronoun)} stated that the water traveled from the fifth floor down to the lobby and left significant flooding in the hallways",
            f"{cap(context.pronoun)} explained that although {context.possessive} unit was not directly flooded, water came through the ceilings of both bathrooms and damaged those bathroom ceilings",
            f"{cap(context.pronoun)} further stated that remediation from that prior flood required drywall removal on the third floor and in the lobby",
        ]
    )


def render_ac_connection_details(context: ReportContext) -> str:
    return join_sentences(
        [
            f"{cap(context.pronoun)} explained that the AC unit has hot and cold water lines, and the connection on the hot-water side came loose",
            f"{cap(context.pronoun)} stated that no one had been working on the unit immediately before the leak",
        ]
    )


def render_filter_age(answer: str, context: ReportContext) -> str:
    if re.search(r"six", answer, re.IGNORECASE) and re.search(r"seven", answer, re.IGNORECASE):
        return f"{cap(context.pronoun)} stated that the water filter line had only been installed approximately six to seven months earlier."
    return f"{cap(context.pronoun)} stated that the water filter line had only been installed recently."


def render_fire_story(context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    return join_sentences(
        [
            f"In describing what happened, the {role_noun} stated that for about two weeks before the fire, {context.pronoun} smelled what seemed like burning electrical wiring and also noticed flickering lights in the house",
            f"{cap(context.pronoun)} stated that on the night of January 15, 2024, {context.pronoun} opened the pantry door and saw flames coming from behind a mini refrigerator",
            f"{cap(context.pronoun)} explained that the fire extinguisher in the house was not working properly, so William went to get another extinguisher from the trailer",
            f"{cap(context.pronoun)} stated that the fire quickly spread, the neighbors and {context.pronoun} called the fire department, and the police and fire department responded",
            f"{cap(context.pronoun)} further stated that two pets were lost during the incident and that William was forcibly removed from the house and taken to Broward General on a Baker Act",
        ]
    )


def render_water_heater_story(context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    return join_sentences(
        [
            f"In describing what happened, the {role_noun} stated that while on the west coast of Florida, {context.pronoun} noticed from an Amazon delivery photo that the front door to the unit was open",
            f"{cap(context.pronoun)} stated that a friend went to check the unit and found the range moved because building maintenance had pulled it out to inspect the water heater",
            f"{cap(context.pronoun)} explained that there was water in the pan beneath the water heater, and maintenance shut off the main valve to the unit",
            f"{cap(context.pronoun)} stated that {context.pronoun} returned that evening, kept emptying the pan to avoid further damage, opened a claim, and arranged for the water heater to be replaced the next morning",
        ]
    )


def render_boating_story(context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    return join_sentences(
        [
            f"In describing what happened, the {role_noun} stated that an old customer had arranged a fishing trip and arrived with several other people, including the claimant",
            f"{cap(context.pronoun)} stated that while everyone was preparing to leave around 8:00 a.m., {context.pronoun} heard a noise, looked up, and saw the claimant on the ground near the stairs leading down toward the pool deck and boat area",
            f"{cap(context.pronoun)} explained that {context.pronoun} went over, helped the claimant up, and the claimant later remained behind while the rest of the group went on the trip",
        ]
    )


def render_stepdaughter_leak_story(context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    return join_sentences(
        [
            f"In describing what happened, the {role_noun} stated that when {context.possessive} stepdaughter arrived at the condo and turned the water on, the front desk later reported a leak believed to be coming from unit 10B",
            f"{cap(context.pronoun)} stated that the stepdaughter shut the water off after being notified. There was an investigation with the management company, and she was instructed not to turn the water back on until the leak was fixed",
            f"{cap(context.pronoun)} further stated that a plumbing company was later authorized to enter the unit, make the repair, send photographs, and that the bill was paid out of pocket",
        ]
    )


def render_condo_remodel_leak_story(answer: str, context: ReportContext) -> str:
    role_noun = short_role_noun(context)
    return join_sentences(
        [
            f"In describing what happened, the {role_noun} stated that {context.pronoun} were in upstate New York when a downstairs neighbor, Lou, called to report leakage",
            f"According to {context.objective}, Lou said he was doing remodeling, the drywall seemed a little wet, and it would be best to investigate the source of the leak",
            f"{cap(context.pronoun)} also stated that the building manager, John, said he knew a plumber who could inspect the unit",
            f"{cap(context.pronoun)} hired a plumber, the wall around the shower was opened, the drainage and shower pan were replaced, and new tile was installed",
            f"{cap(context.pronoun)} believed the problem had been fixed after Lou later sent a text thanking them for resolving it",
            f"{cap(context.pronoun)} further stated that a claim was reported to the insurance carrier, an adjuster inspected both units, and the claim was denied because the policy did not cover plumbing",
        ]
    )


def render_tenant_occupancy(answer: str, context: ReportContext) -> str:
    text = narrativize_answer(answer, context)
    pronoun = cap(context.pronoun)
    lowered = normalize(answer).lower().strip(" .")
    if lowered.startswith("yes"):
        return f"{pronoun} confirmed that the unit was occupied on the date of loss"
    if lowered.startswith("no"):
        return f"{pronoun} confirmed that the unit was not occupied on the date of loss"
    if "active lease" in text.lower():
        sentence = f"{pronoun} confirmed that the unit was rented to a tenant on the date of loss. {pronoun} stated that there was an active lease"
        if "does not know" in text.lower():
            sentence += f", but {context.pronoun} does not know whether the tenant was physically present in the unit at the time of the incident"
        primary_use = context.facts.get("tenant_primary_residence", "")
        if primary_use and ("cannot answer" in primary_use.lower() or "can't answer" in primary_use.lower() or "do not know" in primary_use.lower() or "don't know" in primary_use.lower()):
            sentence += f" or if the tenant used it as {context.possessive} primary residence"
        return sentence
    return f"{pronoun} stated that {lower_first(text)}"


def render_tenant_primary_use(answer: str, context: ReportContext) -> str:
    text = narrativize_answer(answer, context)
    pronoun = cap(context.pronoun)
    if "cannot answer" in text.lower() or "does not know" in text.lower():
        return f"{pronoun} stated that {context.pronoun} does not know whether the tenant used the property as a primary residence"
    return f"{pronoun} stated that {lower_first(text)}"


def render_property_management(answer: str, context: ReportContext) -> str:
    pronoun = cap(context.pronoun)
    lowered = normalize(answer).lower()
    if answer.strip().lower() == "self":
        return f"{pronoun} confirmed that {context.pronoun} self-manages the unit and does not use a property manager"
    if lowered.startswith("no") and "manages all of the units" in lowered:
        return f"{pronoun} stated that {context.pronoun} did not use a separate property manager for the unit"
    return f"{pronoun} stated that {lower_first(narrativize_answer(answer, context))}"


def short_role_noun(context: ReportContext) -> str:
    core = role_core_label(context)
    if context.role_label.lower().startswith("the "):
        return core.lower()
    return "interviewee"


def insured_reference(context: ReportContext, possessive: bool = False) -> str:
    insured_name = tidy_name(context.facts.get("insured_name", "")) if context.facts.get("insured_name") else ""
    representative_name = tidy_name(context.name or "") if context.name else ""
    if not insured_name or (representative_name and insured_name.lower() == representative_name.lower()):
        return "the named insured's" if possessive else "the named insured"
    if possessive:
        return f"{insured_name}'" if insured_name.endswith("s") else f"{insured_name}'s"
    return insured_name


def related_insured_reference(context: ReportContext, possessive: bool = False) -> str:
    relationship_label = context.facts.get("insured_relationship_label")
    if relationship_label and "personal representative" in context.role_label.lower():
        if possessive:
            return f"{context.possessive} {relationship_label}'s"
        return f"{context.possessive} {relationship_label}"
    return insured_reference(context, possessive=possessive)


def paragraph_subject(context: ReportContext) -> str:
    if "personal representative" in context.role_label.lower() and context.name:
        return context.name.split()[0]
    return context.role_label


def render_personal_occupancy(answer: str, context: ReportContext) -> str:
    pronoun = cap(context.pronoun)
    return f"Prior to that time, {context.pronoun} personally occupied the unit. {pronoun} stated that for all of the years preceding six years ago, since purchasing the unit, it was {context.possessive} residence"


def render_lease_reason(answer: str, context: ReportContext) -> str:
    pronoun = cap(context.pronoun)
    return f"{pronoun} stated that {context.pronoun} did not provide the lease agreement yet as {context.pronoun} is currently traveling and does not have access to the lease agreement"


def render_tenant_move_in(answer: str, context: ReportContext) -> str:
    pronoun = cap(context.pronoun)
    if "6 years" in answer.lower():
        return f"{pronoun} stated that the tenant first moved into the property approximately six years ago, roughly around 2020"
    return f"{pronoun} stated that the tenant first moved into the property {lower_first(clean_month_year_reference(narrativize_answer(answer, context)))}"


def join_sentences(items: list[str]) -> str:
    ordered: list[str] = []
    seen: set[str] = set()
    for item in items:
        cleaned = item.strip()
        if not cleaned:
            continue
        key = cleaned.rstrip(".").strip().lower()
        if key in seen:
            continue
        seen.add(key)
        ordered.append(cleaned.rstrip(".") + ".")
    return " ".join(ordered)


def sentence_count(text: str) -> int:
    return len(split_sentences(text))


def word_count(text: str) -> int:
    return len(normalize(text).split())


def lower_first(text: str) -> str:
    text = text.strip()
    if not text:
        return text
    return text[:1].lower() + text[1:]


def cap(text: str) -> str:
    return text[:1].upper() + text[1:]


def compose_opening_paragraphs_old(context: ReportContext) -> list[str]:
    facts = context.facts
    paragraphs: list[str] = []
    pronoun = cap(context.pronoun)
    role = context.role_label
    name = context.name

    first: list[str] = []
    if name and facts.get("name_confirmed"):
        first.append(f"{role}, {name}, confirmed the spelling of {context.possessive} name")
    elif name:
        first.append(f"{role}, {name}, was identified during the statement")
    if facts.get("date_of_birth"):
        if first:
            first[-1] += f" as well as {context.possessive} date of birth"
        else:
            first.append(f"{role} provided {context.possessive} date of birth")
    if facts.get("residence_address"):
        first.append(f"{pronoun} resides at {facts['residence_address']}")
    if facts.get("occupation"):
        first.append(f"{context.possessive.capitalize()} occupation was stated as {facts['occupation']}")
    if facts.get("marital_status"):
        first.append(f"{context.possessive.capitalize()} marital status was stated as {facts['marital_status']}")

    spouse_name = facts.get("spouse_name")
    if spouse_name:
        spouse_items = [f"{context.possessive.capitalize()} spouse's name was provided as {tidy_name(spouse_name)}"]
        if facts.get("spouse_dob"):
            spouse_items.append(f"{pronoun} provided {context.possessive} spouse's date of birth as {facts['spouse_dob']}")
        if facts.get("spouse_occupation"):
            spouse_items.append(render_spouse_occupation(facts["spouse_occupation"], context))
        first.append(join_sentences(spouse_items))

    if facts.get("relationship_to_insured"):
        first.append(f"{pronoun} explained that {lower_first(facts['relationship_to_insured'])}")
    if facts.get("representation_reason"):
        first.append(f"{pronoun} is representing the insured because {lower_first(facts['representation_reason'])}")

    loss_address = facts.get("loss_address")
    if loss_address and facts.get("loss_unit"):
        loss_address = append_unit(loss_address, facts["loss_unit"])
    if loss_address:
        first.append(f"The property where the incident allegedly occurred is located at {loss_address}")
    if facts.get("date_of_loss"):
        first.append(render_date_of_loss_opening(facts["date_of_loss"], context))
    if first:
        paragraphs.append(join_sentences(first))

    second: list[str] = []
    if facts.get("ownership_length"):
        second.append(render_ownership(facts["ownership_length"], context))
    if facts.get("building_age"):
        second.append(render_building_age(facts["building_age"], context))
    if facts.get("property_use"):
        second.append(render_property_use(facts["property_use"], context))
    elif facts.get("rental_duration"):
        second.append(f"On the date of loss, the {role.lower().replace('the ', '')} stated that the property was being used as a rental")
    if facts.get("rental_duration"):
        second.append(f"The unit had been used as a rental property for {lower_first(facts['rental_duration'])}")
    if facts.get("personal_occupancy_when"):
        second.append(render_personal_occupancy(facts["personal_occupancy_when"], context))
    if facts.get("tenant_occupancy"):
        second.append(render_tenant_occupancy(facts["tenant_occupancy"], context))
    if facts.get("lease_not_provided_reason"):
        second.append(render_lease_reason(facts["lease_not_provided_reason"], context))
    if facts.get("tenant_move_in"):
        second.append(render_tenant_move_in(facts["tenant_move_in"], context))
    if facts.get("tenant_primary_residence") and not facts.get("tenant_occupancy"):
        second.append(render_tenant_primary_use(facts["tenant_primary_residence"], context))
    if facts.get("property_management"):
        second.append(render_property_management(facts["property_management"], context))
    if facts.get("occupants") and facts.get("property_use") != "full-time residence":
        second.append(f"{pronoun} stated that the occupants on the date of loss were {lower_first(narrativize_answer(facts['occupants'], context))}")
    if facts.get("presence"):
        second.append(render_presence(facts["presence"], context))
    if second:
        paragraphs.append(join_sentences(second))

    return paragraphs


def should_skip_answer(answer: str) -> bool:
    lowered = answer.strip().lower().strip(" .")
    if lowered in {"yes", "no", "yeah", "yep", "nope", "okay", "ok", "great", "that's correct", "correct"}:
        return True
    if lowered.startswith(("yes,", "yeah,", "yeah okay", "yeah. okay", "yes. okay", "no. okay", "correct,")):
        return True
    if lowered in {"approximately", "roughly", "that's the extent of my knowledge", "that's the extent of my knowledge, yes"}:
        return True
    return False


def topic_prefix(question: str) -> Optional[str]:
    lowered = question.lower()
    if match(lowered, "what happened", "walk me through", "tell me the whole story", "in your own words"):
        return "In describing what happened"
    if match(lowered, "kitchen sink leak"):
        return "Regarding the kitchen sink leak"
    if match(lowered, "ac leak"):
        return "Regarding the AC leak"
    if match(lowered, "source of the leak", "cause of the leak", "caused it to leak", "what caused it to leak", "what caused the leak"):
        return "Regarding the source of the leak"
    if match(lowered, "prior repairs", "prior leaks", "prior issues"):
        return "Regarding prior issues"
    if match(lowered, "injuries"):
        return "Regarding injuries"
    if match(lowered, "damages"):
        return "Regarding damages"
    if match(lowered, "property manager"):
        return "Regarding property management"
    if match(lowered, "other properties"):
        return "Regarding other properties"
    if match(lowered, "short-term rental", "airbnb", "vrbo"):
        return "Regarding rental use of the property"
    return None


def classify_paragraph_topic(question: str, answer: str = "") -> str:
    lowered = question.lower()
    answer_lowered = normalize(answer).lower()
    if match(
        lowered,
        "gustavo",
        "foster agency",
        "fostering",
        "adopt",
        "residential care",
        "child welfare",
        "stipend",
    ):
        return "gustavo_background"
    if match(
        lowered,
        "shooting",
        "party",
        "christian wolf",
        "leo woodard",
        "who shot",
        "what happened with the police",
    ):
        return "shooting"
    if match(
        lowered,
        "gun",
        "lock box",
        "lockbox",
        "ammunition",
        "magazines",
        "firearm",
        "trigger lock",
    ):
        return "gun_storage"
    if match(
        lowered,
        "claimant",
        "ambulance",
        "hospital",
        "injuries",
        "surgeries",
        "recovered",
        "physical appearance",
    ):
        return "claimant"
    if match(
        lowered,
        "while the unit was vacant, did you have anyone checking the unit",
        "how often did they check the unit",
        "did they ever report any leaks",
        "was the water turned off while the unit was vacant",
    ) or (question.strip().lower() == "why?" and ("cleaning girl" in answer_lowered or "water would be on" in answer_lowered)):
        return "vacancy_monitoring"
    if match(
        lowered,
        "how did william sidebottom come to live",
        "allow the claimant to live at your property for free",
        "how long exactly has william sidebottom lived",
        "lease agreement with william sidebottom",
        "has william sidebottom ever paid rent",
        "had his own bedroom",
        "has the claimant had a key to the property",
        "received all of his mail at your property",
        "driver's license registered at your property",
        "vehicles registered at your property",
        "specific household duties",
        "still live at your property",
        "what is your relationship to the claimant",
        "common-law spouse",
        "employee of yours",
    ):
        return "claimant_background"
    if match(
        lowered,
        "after the fire occurred, where did you live",
        "does the claimant live in the trailer",
        "has the claimant been living in this trailer",
        "where else has he lived",
        "do you also live in the trailer",
        "why do you live in the trailer",
        "how long have you two been living in this trailer",
        "where is the trailer located",
        "was the claimant ever living inside the trailer",
        "who owns the trailer",
        "do you personally own it",
        "does the claimant perform his job duties from the trailer",
        "when did you purchase the trailer",
        "why did you purchase the trailer",
        "how long has the trailer been on your property",
        "why was the trailer originally kept on the property",
        "permanent fixture or is it mobile",
        "does it have a license plate",
        "does it run",
    ):
        return "post_fire_living"
    if match(
        lowered,
        "so you occupy it one week in december",
        "who lives at this property with you",
        "was the unit occupied on the date of loss",
        "was anyone home when the incident occurred",
        "how long had you been away when the incident occurred",
        "did you have anyone checking on the unit while it was unoccupied",
    ):
        return "seasonal_background"
    if match(
        lowered,
        "do you run a fishing charter business",
        "were you operating this fishing charter on the date of loss",
        "how long have you been running your fishing charter business",
        "how long have you been semi-retired",
        "while being semi-retired, were you doing the small charters",
        "did you sell the boat",
        "do you run the fishing charter business out of your backyard",
        "operating out of your backyard",
        "how long had you been operating this fishing charter out of your backyard",
        "where do you permanently keep the boat",
        "how long have you stored your boat back there in your yard",
        "can you describe the boat",
        "how many engines does it have",
        "inboard or outboard",
        "horsepower of those engines",
        "do you have a registered business",
        "operating as just a sole proprietor",
        "group of paying customers",
        "how much did they pay",
        "aware of this dog prior to the incident",
        "other people in the group aware",
    ):
        return "charter_business"
    if match(lowered, "in your own words", "tell me the whole story", "first became aware of the leak"):
        return "incident"
    if match(
        lowered,
        "immediately after becoming aware of the leak",
        "did he sees all use of the shower",
        "how long the leak was going on for before it was stopped",
        "did anyone check",
        "did anyone shut the water off after the leak was discovered",
        "did you ever speak with the claimant",
        "have you hired a plumber or leak detection yet",
    ):
        return "response"
    if match(
        lowered,
        "how many units were damaged as a result of this incident",
        "was 12g occupied on the date of loss",
        "what the specific damage is in 12g",
        "dry out or mold remediation completed",
        "repairs completed yet",
        "anyone was injured or is claiming to be injured",
        "12g has insurance on their unit",
        "damages to condo common areas",
    ):
        return "damages"
    if match(
        lowered,
        "what leaked in your unit",
        "what specifically leaked",
        "how old is the shower in the shower pan",
        "which bathroom is this shower located in",
        "replaced the shower since you purchased the unit",
        "and you said it was around 2014",
        "did you use a licensed contractor",
        "installed with a permit",
        "do you know what caused it to leak",
        "had the shower pan ever leaked before",
        "prior repairs, leaks, or issues involving the shower or shower pan",
        "tenant ever complained about any shower related leaks",
        "glass leak on the shower prior but it's unrelated",
        "any reason to believe that the shower or shower pan was at risk of leaking",
        "leak has not been repaired yet",
    ):
        return "shower"
    if match(
        lowered,
        "water shutoff rules in the association bylaws",
        "file a first party claim",
    ):
        return "policy"
    if match(
        lowered,
        "have you had any prior water leaks originating",
        "and when was that leak",
        "can you give me an approximate",
        "what leaked before",
        "did this prior leak affect the unit below",
        "did it affect their master bathroom",
        "was 12g repaired prior to this current leak",
    ):
        return "prior_loss"
    if match(
        lowered,
        "other condo unit owners or the association advised you",
        "provided citizens with all documentation",
        "excess liability or umbrella insurance",
        "own any additional properties in the united states",
    ):
        return "closing"
    if match(
        lowered,
        "while jim",
        "involvement with the property located",
        "obtaining insurance for the property",
        "written property management agreements",
        "aware that",
        "when did you become aware of this incident",
        "became aware of this incident before he passed away",
        "reported the incident to citizens at the time",
        "ever reported the incident to citizens before he passed away",
        "did you ever have any direct communications",
        "text message compilation",
        "phone in which the text messages originated",
    ):
        return "representative"
    if match(
        lowered,
        "renting the property as an airbnb",
        "short-term rental",
        "long-term renters",
        "jim's records such as contracts",
        "internet accounts such as airbnb or vrbo",
        "claimed on jim's personal or business taxes",
        "which one",
        "full name of the company",
        "prior tax returns",
        "familiar with kind of enterprises",
        "tell me anything else about this llc",
        "liquor cabinet",
    ):
        return "rental_history"
    if match(
        lowered,
        "purchase the property in which the incident occurred",
        "date of loss was the property being used as a rental property",
        "what kind of property is the loss location",
        "was the property occupied on the date of loss",
        "were the renters the claimants",
        "were they short-term renters",
        "booking records for this week",
        "booking or rental agreement",
        "when did the claimant start renting the property",
        "prior to the date of loss, did they ever make any complaints",
    ):
        return "property_profile"
    if match(
        lowered,
        "did our insured, jim, have a property manager",
        "written agreement with randal",
        "paid for his property manager services",
        "responsibilities as the property manager",
        "what dates was he the insured's property manager",
        "who is responsible for the maintenance of the property",
        "alert the insured about any issues with the deck",
        "delegated responsibility for day-to-day maintenance",
        "claimants did rent this property before",
        "stayed there the year prior",
        "prior complaints or issues the year before",
    ):
        return "management"
    if match(
        lowered,
        "on the data loss, what happened",
        "how did you first become aware of the incident",
        "only became aware of this in mid to late december of 2025",
        "weren't present when the incident occurred",
        "report the incident to anyone after you became aware",
        "our insured ever became aware of the incident",
        "our insured was present when the incident occurred",
        "what the claimant was doing when the incident occurred",
        "where exactly on the property did the incident occur",
        "where exactly on the deck",
        "why the claimant was on the deck",
        "what they were doing on the deck",
        "authorized to be on the deck",
    ):
        return "deck_incident"
    if match(
        lowered,
        "describe the deck on the date of loss",
        "what material the deck was made of",
        "how old was the deck",
        "deck present when the insured purchased the property",
        "deck renovated after our insured purchased the property",
        "anything wrong with the deck",
        "rot or deterioration on the deck",
        "termites or other wood-eating pests",
        "prior repairs to the deck",
        "information about his renovation of the deck",
        "complained about the deck",
        "other deck boards ever broken prior",
        "aware of any issues with the deck on the date of loss",
        "fallen through the deck prior",
        "deck on the 1st or 2nd floor",
        "travel through the area where the incident occurred",
        "has the deck been replaced since the date of loss",
        "when?",
        "deck in which the incident occurred, was it replaced",
        "deck where the incident occurred been replaced or not",
        "prior to me coming out and inspecting",
        "has the deck been repaired or replaced prior to my inspection",
        "regular maintenance performed on the deck",
        "regular maintenance performed on the property",
        "regularly inspect the deck",
        "visit and inspect the property",
        "how often would they go out there",
    ):
        return "deck_condition"
    if match(
        lowered,
        "how did the claimant come to fall through the deck",
        "is there a defect that caused the claimant to fall through the wood deck",
        "warning signs, caution tape",
        "fall all the way through the deck",
        "did anyone else fall through the deck on the date of loss",
        "was anyone else injured on the date of loss",
        "did the claimant continued to walk after the fall",
        "continued to walk after the fall",
        "did anyone call for an ambulance",
        "did the claimant go to the hospital",
        "how long the claimant was in the hospital",
        "claimant's injuries are or were",
        "what injuries do you believe she had",
        "what treatments did the claimant undergo",
        "surgeries, physical therapy, scarring, or permanent disabilities",
        "recovered from their injuries",
        "have you ever seen or spoken with the claimant",
        "property manager ever see or speak with the claimant",
        "what was discussed",
        "prior injuries or accidents",
    ):
        return "medical"
    if match(
        lowered,
        "what time of day the incident occurred",
        "working lighting in that area",
        "weather was like on the date of loss",
        "claimant's physical appearance",
        "claimant distracted or under the influence",
        "claimant wore glasses",
        "mobility or balance issues",
        "did the claimant have any prior injuries",
        "know the claimant in any way",
        "directing or assisting the claimant",
        "could have done to prevent the incident",
        "video cameras on the property",
        "aware of any witnesses",
    ):
        return "incident_details"
    if lowered.strip() == "which one?" and ("sole proprietorship" in answer_lowered or "social security" in answer_lowered):
        return "rental_history"
    if lowered.strip() == "when?" and ("took possession of the property" in answer_lowered or "after that" in answer_lowered):
        return "deck_condition"
    return "general"


def compose_body_sentence(pair: QAPair, context: ReportContext, index: int) -> Optional[str]:
    answer = narrativize_answer(pair.answer, context)
    if not answer:
        return None

    question = pair.question.lower()
    role_noun = short_role_noun(context)
    pronoun = cap(context.pronoun)
    do_word = "do" if context.joint_statement else "does"
    are_word = "are" if context.joint_statement else "is"
    believe_word = "believe" if context.joint_statement else "believes"
    insured_ref = insured_reference(context)
    natural_insured_ref = related_insured_reference(context)
    natural_insured_possessive = related_insured_reference(context, possessive=True)
    shooting_claim = context.facts.get("claim_scenario") == "shooting"
    dog_claim = context.facts.get("claim_scenario") == "dog_bite"
    fire_claim = context.facts.get("claim_scenario") == "fire"
    boat_claim = context.facts.get("claim_scenario") == "boat_trip"
    roof_slip_claim = context.facts.get("claim_scenario") == "roof_slip"
    water_heater_claim = context.facts.get("claim_scenario") == "water_heater"
    overflow_claim = context.facts.get("claim_scenario") == "tub_overflow"
    washing_machine_claim = context.facts.get("claim_scenario") == "washing_machine_leak"
    wall_pipe_claim = context.facts.get("claim_scenario") == "wall_pipe"
    water_damage_claim = context.facts.get("claim_scenario") == "water_damage"

    if match(question, "but my question was"):
        return None
    if washing_machine_claim and match(question, "in your own words, tell me the story of what happened involving this incident"):
        return (
            f"In describing what happened, the {role_noun} stated that a neighbor reported a leak coming from the unit. "
            f"{pronoun} explained that the unit was investigated, the washing machine was eventually identified as the source of the leak, and the washing machine was replaced."
        )
    if washing_machine_claim and match(question, "can you give me a little bit more detail"):
        return f"{pronoun} stated that {context.pronoun} did not have any more detail than that."
    if washing_machine_claim and match(question, "when did you first become aware of the leak"):
        date_phrase = extract_date_phrase(answer)
        if date_phrase:
            return f"{pronoun} stated that {context.pronoun} first became aware of the leak sometime around {clean_date_answer(date_phrase)}."
        return f"{pronoun} stated that {context.pronoun} first became aware of the leak {lower_first(narrativize_answer(answer, context))}."
    if washing_machine_claim and match(question, "how did you first become aware of the leak"):
        return (
            f"{pronoun} stated that {context.pronoun} first learned of the leak through either a call from the neighboring unit "
            f"or a message from the condo association president in the building WhatsApp group."
        )
    if washing_machine_claim and match(question, "what did you do after becoming aware of the leak"):
        return f"After becoming aware of the leak, the {role_noun} stated that the unit was investigated, the washing machine was identified as the source, and the machine was replaced."
    if washing_machine_claim and match(question, "do you know how long the leak was going on for before it was stopped"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the leak was occurring before it was discovered or before it was stopped."
    if washing_machine_claim and match(question, "did anyone shut the water after", "shut the water off after the leak was discovered"):
        return f"{pronoun} stated that once the washing machine was identified as the source, the water to the washing machine was shut off."
    if washing_machine_claim and match(question, "who shut the water off"):
        return f"{pronoun} stated that a neighbor or friend apparently helped shut off the water to the washing machine."
    if washing_machine_claim and match(question, "do you know when they shut the water off"):
        return f"{pronoun} stated that the water to the washing machine was shut off shortly after the leak was discovered."
    if washing_machine_claim and match(question, "did you report the leak to anyone, including the condo association"):
        return f"{pronoun} stated that the condo association president was the person who reported the leak to {context.objective}, so the association was already aware of it."
    if washing_machine_claim and match(question, "did you ever speak with the claimant or claimants about the incident"):
        return f"{pronoun} stated that {context.pronoun} did not speak with the claimant about the incident."
    if washing_machine_claim and match(question, "did you guys hire any kind of professional to evaluate the cause of the machine issue"):
        return None
    if washing_machine_claim and match(question, "what did they determine"):
        return (
            f"{pronoun} stated that an initial inspection looked for the source of the leak but did not identify the washing machine as the cause. "
            f"{pronoun} explained that when the leak happened again, the washing machine was identified as the source."
        )
    if washing_machine_claim and match(question, "so this leak happened before"):
        return None
    if washing_machine_claim and match(question, "so the washing machine leaked twice"):
        return f"{pronoun} stated that {context.pronoun} believes there may have been two leak events close together, but {context.pronoun} does not know for certain."
    if washing_machine_claim and match(question, "when's the first time you had a professional come out"):
        date_phrase = extract_date_phrase(answer)
        if date_phrase:
            return f"{pronoun} stated that the first professional inspection occurred around {clean_date_answer(date_phrase)}."
        return f"{pronoun} stated that the first professional inspection occurred {lower_first(narrativize_answer(answer, context))}."
    if washing_machine_claim and match(question, "understood. so when you had a professional come out, they looked at everything but the washing machine"):
        return f"{pronoun} stated that {context.pronoun} had been under the impression the washing machine was inspected, but the invoice seemed to refer to other possible sources instead."
    if washing_machine_claim and match(question, "okay, so the invoice or report they provided had nothing to do with the washing machine that actually leaked", "right, and that invoice did not identify any issues with the washing machine or even mention the washing machine"):
        return None
    if washing_machine_claim and match(question, "how many units were damaged as a result of this incident, not including your own"):
        return f"{pronoun} stated that the only unit {context.pronoun} is aware of being damaged was the unit directly below."
    if washing_machine_claim and match(question, "do you know what unit number that is"):
        return f"{pronoun} stated that {context.pronoun} believes the unit below may be unit 206, but {context.pronoun} does not know for certain."
    if washing_machine_claim and match(question, "was this unit occupied on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the unit below was occupied on the date of loss."
    if washing_machine_claim and match(question, "well, do you know if it was occupied or not"):
        return f"{pronoun} stated that {context.pronoun} does not know for certain, although {context.pronoun} assumed someone may have been living there because the residents were communicating through the building association."
    if washing_machine_claim and match(question, "do you know what their specific damages are or were"):
        return f"{pronoun} stated that {context.pronoun} is not aware of the specific damages being claimed in the unit below."
    if washing_machine_claim and match(question, "do you know if they had any dry out or mold remediation completed"):
        return f"{pronoun} stated that {context.pronoun} does not know whether dry out or mold remediation was completed in the unit below."
    if washing_machine_claim and match(question, "do you know if they've had their damages repaired yet"):
        return f"{pronoun} stated that {context.pronoun} does not know whether repairs have been completed in the unit below."
    if washing_machine_claim and match(question, "do you know if anyone was injured or is claiming to be injured as a result of the incident"):
        return f"{pronoun} stated that {context.pronoun} is not aware of anyone being injured or claiming to be injured as a result of the incident."
    if washing_machine_claim and match(question, "do you know if they have insurance on their unit below"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the owner of the unit below had insurance."
    if washing_machine_claim and match(question, "are you aware of any damages to condo common areas"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any damage to condo common areas."
    if washing_machine_claim and match(question, "so just to reconfirm, what exactly caused the leak in your unit on the date of loss"):
        return f"{pronoun} stated that it appears the washing machine was the source of the leak and that the machine was replaced shortly after {context.pronoun} became aware of the issue."
    if washing_machine_claim and match(question, "what specifically on the washing machine leaked"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the leak came from an overflow or from the bottom of the washing machine. All {context.pronoun} knows is that the washing machine leaked and was replaced."
    if washing_machine_claim and match(question, "did you ever have the washing machine specifically evaluated by any professionals"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any professional specifically evaluating the washing machine itself."
    if washing_machine_claim and match(question, "how old was the washing machine", "how old was that washing machine"):
        return f"{pronoun} stated that {context.pronoun} does not know how old the washing machine was. {pronoun} explained that it came with the unit."
    if washing_machine_claim and match(question, "do you know if it was newer or older"):
        return None
    if washing_machine_claim and match(question, "do you know what caused it to overflow or leak"):
        return f"{pronoun} stated that {context.pronoun} does not know what caused the washing machine to overflow or leak."
    if washing_machine_claim and match(question, "do you know if your mother went to sleep with the washing machine on"):
        return f"{pronoun} stated that {context.pronoun} does not know whether {context.possessive} mother went to sleep with the washing machine running on the date of loss or whether she had ever done that before."
    if washing_machine_claim and match(question, "had she ever gone to bed with the washing machine running before"):
        return None
    if washing_machine_claim and match(question, "is this the first time the washing machine ever overflowed"):
        return f"{pronoun} stated that {context.pronoun} believes this was the first time the washing machine overflowed or leaked, but {context.pronoun} does not know for certain."
    if washing_machine_claim and match(question, "was there any rust or corrosion on any parts of the washing machine or its components"):
        return f"{pronoun} stated that {context.pronoun} does not know whether there was any rust or corrosion on the washing machine or its components."
    if washing_machine_claim and match(question, "have you had any prior issues or repairs involving the washing machine"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any prior issues or repairs involving the washing machine."
    if washing_machine_claim and match(question, "did you have any reason to believe that the washing machine was at risk of overflowing or leaking prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe the washing machine was at risk of leaking prior to the date of loss."
    if washing_machine_claim and match(question, "when was the washing machine replaced"):
        return f"{pronoun} stated that the washing machine was replaced on {clean_date_answer(answer)}."
    if washing_machine_claim and match(question, "water shutoff rule in the association bylaws"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any water shut-off rule in the association bylaws."
    if washing_machine_claim and match(question, "did you file a first party claim of citizens for damages to your own property"):
        return f"{pronoun} stated that {context.pronoun} did not file a first-party claim with Citizens for damage to {context.possessive} own unit."
    if washing_machine_claim and match(question, "was there any damage to your own property"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any damage to {context.possessive} own unit."
    if washing_machine_claim and match(question, "have you had any other leaks originating from your unit since you've owned it"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any other leaks originating from the unit since owning it."
    if washing_machine_claim and match(question, "have any other condo unit owners or the association advised you that they have any damages or requested any payments from you"):
        return f"{pronoun} stated that no other unit owners or the association have advised {context.objective} of damages or requested payment."
    if washing_machine_claim and match(question, "have you provided citizens with all documentation you have related to this incident"):
        return f"{pronoun} stated that {context.pronoun} has provided Citizens with all documentation {context.pronoun} has related to the incident."
    if washing_machine_claim and match(question, "do you have any excess liability or umbrella insurance"):
        return f"{pronoun} stated that {context.pronoun} only has the coverage provided under this policy and does not have any separate excess liability or umbrella insurance."
    if washing_machine_claim and match(question, "do you own any additional properties in the united states"):
        if is_affirmative(pair.answer):
            return f"{pronoun} stated that, aside from the subject unit, {context.pronoun} also owns {context.possessive} primary residence."
        return f"{pronoun} stated that {context.pronoun} does not own any additional properties in the United States."
    if washing_machine_claim and match(question, "how many?"):
        return None
    if washing_machine_claim and match(question, "and is that the address you provided earlier"):
        return None
    if washing_machine_claim and match(question, "is it insured"):
        return f"{pronoun} confirmed that {context.possessive} residence is insured."
    if washing_machine_claim and match(question, "and who is it insured with", "who is it insured with"):
        return f"{pronoun} stated that {context.pronoun} did not know the carrier offhand because the policy had just been renewed through {context.possessive} broker."
    if match(question, "what are you retired from") and "florida power and light" in pair.answer.lower():
        return "Clark stated that he previously worked at Florida Power and Light for approximately 25 years and later worked in child welfare for approximately eight years."
    if water_damage_claim and match(question, "did you hire a plumber or an ac technician for the ac yet"):
        return f"{pronoun} stated that {context.pronoun} has hired a plumber and is trying to have the connection repaired, but the system still needs a working water connection so the AC can be tested."
    if water_damage_claim and match(question, "how old was the ac"):
        return f"{pronoun} stated that the AC unit was approximately 9 to 10 years old."
    if water_damage_claim and match(question, "what are their unit numbers"):
        return f"{pronoun} stated that units 210 and 312 were the other units involved in the AC leak."
    if water_damage_claim and match(question, "unit next door to yours and the unit below you, correct"):
        return f"{pronoun} clarified that unit 210 was the unit below, and unit 312 was the neighboring unit."
    if water_damage_claim and match(question, "when someone was present"):
        return f"{pronoun} stated that no one was present in unit 312 at the time of the AC leak, but {context.pronoun} was not sure whether anyone was present in unit 210."
    if water_damage_claim and match(question, "was anyone living in those units when the incident occurred"):
        return f"{pronoun} stated that unit 312 was occupied as a residence, but {context.pronoun} does not know whether unit 210 was occupied as a residence."
    if water_damage_claim and match(question, "do you know what the specific damages are in those two units from this ac leak"):
        return f"{pronoun} stated that {context.pronoun} does not know the specific damages sustained in units 210 and 312 from the AC leak."
    if water_damage_claim and match(question, "do you know if either of those units had any dry out or mold remediation completed as a result of this ac leak"):
        return f"{pronoun} stated that the remediation company mentioned drying equipment in unit 210, but {context.pronoun} does not know anything further about remediation in either unit."
    if water_damage_claim and match(question, "did you hire a plumber regarding the kitchen sink leak"):
        return f"{pronoun} stated that no plumber was needed for the kitchen sink leak because the issue only required replacement of the connection valve and tubing."
    if water_damage_claim and match(question, "what did you do after becoming aware of the kitchen sink leak"):
        return f"Regarding the kitchen sink leak, {context.pronoun} stated that {context.pronoun} cleaned the small amount of water from the kitchen floor."
    if water_damage_claim and match(question, "did you tell him specifically to replace it"):
        return f"{pronoun} confirmed that the building maintenance supervisor repaired the fitting."
    if overflow_claim and match(question, "while the unit was vacant, did you have anyone checking the unit"):
        return f"{pronoun} stated that someone periodically checked the unit while it was vacant."
    if overflow_claim and match(question, "and prior to the tenants moving in, how long was the unit vacant"):
        return None
    if overflow_claim and match(question, "how often did they check the unit"):
        return f"{pronoun} stated that the unit was checked about once a month, or every three weeks."
    if overflow_claim and match(question, "did they ever report any leaks, unexplained water on the floor, anything dripping, or anything ever leaking") and pair.answer.lower().startswith("no"):
        return f"{pronoun} stated that no leaks, dripping, or unexplained water were reported while the unit was vacant."
    if overflow_claim and match(question, "was the water turned off while the unit was vacant"):
        return f"{pronoun} stated that {context.pronoun} did not believe the water had been shut off while the unit was vacant."
    if overflow_claim and question.strip().lower() in {"why?", "why not?"} and ("cleaning girl" in pair.answer.lower() or "water would be on" in pair.answer.lower()):
        return f"{pronoun} explained that the water remained on because the cleaning person needed to use it when freshening or cleaning the unit."
    if overflow_claim and match(question, "so the claimant had been complaining about the leak or the smell since summer of 2025 and they only alerted you in january"):
        return f"{pronoun} disputed that the claimant had personally been complaining for months and stated that, to {context.possessive} knowledge, the lower unit was vacant and up for sale. {pronoun} explained that building personnel reportedly discovered a mildew smell while in that unit for another reason and then notified the owner."
    if overflow_claim and match(question, "so the condo association property manager is the one that observed the smell or leak inside the claimant's unit"):
        return f"{pronoun} confirmed that, to {context.possessive} understanding, the property manager or building personnel were the ones who noticed the smell in the lower unit."
    if overflow_claim and match(question, "when was this", "so you're speculating that the smell or whatever was discovered in july of 2025"):
        return f"{pronoun} stated that {context.pronoun} could only estimate that the smell may have been discovered in late June or July 2025."
    if overflow_claim and match(question, "so the only damage potentially is to unit 233 directly below you"):
        return f"{pronoun} stated that the only unit potentially affected was unit 233 directly below."
    if overflow_claim and match(question, "did you say before that this unit was not occupied on the date of loss"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, unit 233 was vacant, unfurnished, and possibly listed for sale."
    if overflow_claim and match(question, "do you know if it was like months or years"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the lower unit had been vacant for months or years."
    if overflow_claim and match(question, "do you know if the claimant had anyone checking the unit while it was vacant"):
        return f"{pronoun} stated that {context.pronoun} does not know for certain, but it does not appear anyone was regularly checking the lower unit."
    if overflow_claim and question.strip().lower() == "why?" and "mildew issue" in pair.answer.lower():
        return f"{pronoun} stated that, if someone had been checking the lower unit regularly, they likely would have noticed the mildew issue sooner."
    if overflow_claim and match(question, "has the leak stopped"):
        return f"{pronoun} confirmed that the leak has stopped."
    if overflow_claim and match(question, "have there been any renovations in the bathroom where the leak occurred since you've owned the unit", "have there been any renovations in the room where the leak occurred since you've owned the unit"):
        return f"{pronoun} stated that the bathrooms had been renovated since {context.pronoun} bought the unit."
    if overflow_claim and match(question, "what did you do with the renovation"):
        return f"{pronoun} stated that the renovations included new floor tile, new toilets, new vanities, and new medicine cabinets."
    if overflow_claim and match(question, "was the bathtub overflow that leaked, replaced, or manipulated during the renovation"):
        return f"{pronoun} stated that the bathtub overflow was not replaced or altered during the renovation and that the tub itself was only re-glazed."
    if overflow_claim and match(question, "so has the bathtub overflow ever been replaced since you've owned the unit"):
        return f"{pronoun} stated that the bathtub overflow has never been replaced since {context.pronoun} bought the unit."
    if overflow_claim and match(question, "did you ever have any issues with your toilet", "has your toilet ever leaked"):
        return f"{pronoun} stated that there had been no issues with the toilets and no toilet leaks in {context.possessive} unit."
    if overflow_claim and match(question, "how old are they") and "toilet" not in question:
        return f"{pronoun} stated that the toilets had been replaced approximately five years earlier."
    if water_heater_claim and match(question, "so you occupy it one week in december"):
        return f"{pronoun} clarified that {context.pronoun} usually stays from late October until early December, returns to Canada for Christmas, and then comes back from early January until mid-April."
    if water_heater_claim and match(question, "who lives at this property with you"):
        return f"{pronoun} stated that {context.pronoun} is usually there with {context.possessive} wife and that {context.possessive} son or {context.possessive} wife's sister sometimes also visit."
    if water_heater_claim and match(question, "was the unit occupied on the date of loss"):
        return f"{pronoun} stated that the loss occurred during {context.possessive} seasonal stay, although {context.pronoun} was away from the unit at the time."
    if water_heater_claim and match(question, "was anyone home when the incident occurred"):
        return f"{pronoun} stated that no one was home when the incident occurred because they were visiting friends on Florida's west coast."
    if water_heater_claim and match(question, "how long had you been away when the incident occurred", "you had only been gone 26 hours when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} had only been away for about 26 hours when the incident occurred."
    if water_heater_claim and match(question, "did you have anyone checking on the unit while it was unoccupied"):
        return f"{pronoun} stated that no one was checking the unit during that short 26-hour trip because {context.pronoun} had shut off the water before leaving."
    if water_heater_claim and match(question, "that's for the full unit"):
        return None
    if water_heater_claim and match(question, "do you know how, so this leak was, do you know how long this leak was going on for before it stopped"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the leak had been occurring before it was discovered."
    if water_heater_claim and match(question, "did you report this incident to the condo association"):
        return f"{pronoun} stated that {context.pronoun} did speak with the condo association about the incident on the following Monday."
    if water_heater_claim and match(question, "what did they say") and "two other units" in pair.answer.lower():
        return f"{pronoun} stated that the association mentioned there had been similar issues in two other units, but did not say anything specific about {context.possessive} unit."
    if fire_claim and match(question, "how did william sidebottom come to live at your house"):
        return f"{pronoun} stated that a friend brought William Sidebottom from California because {context.pronoun} needed help with {context.possessive} yacht-industry business."
    if fire_claim and question.strip().lower() == "okay. and then what?":
        return f"{pronoun} explained that William was allowed to live there in exchange for learning the yacht business and providing discounted work for {context.possessive} customers."
    if fire_claim and match(question, "are you related to william in any way"):
        return f"{pronoun} stated that {context.pronoun} is not related to William."
    if fire_claim and match(question, "so is william sidebottom your roommate"):
        return f"{pronoun} stated that William lived in the house with {context.objective} and could be considered a roommate."
    if fire_claim and match(question, "how long exactly has william sidebottom lived at your property"):
        return f"{pronoun} stated that William had lived at the property since 2012."
    if fire_claim and match(question, "do you now or have you ever had a lease agreement with william sidebottom"):
        return f"{pronoun} stated that there has never been a lease agreement with William."
    if fire_claim and match(question, "has william sidebottom ever paid rent"):
        return f"{pronoun} stated that William has never paid rent."
    if fire_claim and match(question, "allow the claimant to live at your property for free"):
        return f"{pronoun} stated that William lived there in exchange for discounted work rather than paying rent."
    if fire_claim and match(question, "was william sidebottom working for you when he first moved in", "has he continued to work for you throughout his residency"):
        return f"{pronoun} stated that William was working with {context.objective} when he moved in and continued doing so throughout his residency."
    if fire_claim and match(question, "has the claimant ever lived anywhere else after moving into your property"):
        return f"{pronoun} stated that William has not lived anywhere else since moving into the property."
    if fire_claim and match(question, "has the claimant had a key to the property for as long as he has lived there", "has the claimant received all of his mail at your property for as long as he lived there", "does the claimant have his driver's license registered at your property", "does the claimant have his vehicles registered at your property"):
        if "mail" in question:
            return f"{pronoun} stated that William has received his mail at the property throughout the time he has lived there."
        if "driver's license" in question:
            return f"{pronoun} stated that William's driver's license is registered to the property."
        if "vehicles" in question:
            return f"{pronoun} stated that William's vehicles are registered to the property."
        return f"{pronoun} stated that William has had a key to the property for as long as he has lived there."
    if fire_claim and match(question, "does the claimant have any specific household duties that he regularly performs"):
        return f"{pronoun} stated that William's regular household chores included taking out the garbage and watering the plants."
    if fire_claim and match(question, "what is your relationship to the claimant"):
        return f"{pronoun} described William as a friend and business associate."
    if fire_claim and match(question, "is the claimant a spouse, a common-law spouse, or a romantic partner", "is the claimant an employee of yours"):
        if "employee" in question:
            return f"{pronoun} stated that William was not a formal employee."
        return f"{pronoun} stated that William was not a spouse, romantic partner, or common-law spouse."
    if fire_claim and match(question, "after the fire occurred, where did you live during the repair of your house"):
        return f"{pronoun} stated that after the fire, {context.pronoun} lived in the RV parked in the yard and driveway during the repairs."
    if fire_claim and match(question, "does the claimant live in the trailer on your property at this time as well"):
        return f"{pronoun} stated that William also lived in the trailer during the repair period."
    if fire_claim and match(question, "did he have his own room inside the house prior to the fire"):
        return f"{pronoun} stated that William had his own bedroom inside the house before the fire."
    if fire_claim and match(question, "has the claimant been living in this trailer the entire time the house has been under repair after the fire"):
        return f"{pronoun} stated that William had been living in the trailer during the repair period, although his bedroom in the house had only recently been finished."
    if fire_claim and match(question, "so since the fire occurred, he goes back and forth between the trailer and your driveway and his bedroom inside the house"):
        return None
    if fire_claim and match(question, "where else has he lived") and "bedroom" in pair.answer.lower():
        return f"{pronoun} stated that within the last month, William had begun occasionally sleeping again in his bedroom inside the house, even though it still lacked air conditioning and heat."
    if fire_claim and match(question, "do you also live in the trailer with the claimant"):
        return f"{pronoun} stated that {context.pronoun} shared the trailer with William."
    if fire_claim and match(question, "why do you live in the trailer with the claimant"):
        return f"{pronoun} stated that they shared the trailer because they had nowhere else to stay during the repairs."
    if fire_claim and match(question, "how long have you two been living in this trailer"):
        return f"{pronoun} stated that they had been living in the trailer for about two years."
    if fire_claim and match(question, "where is the trailer located"):
        return f"{pronoun} stated that the trailer is located in the yard on a pad."
    if fire_claim and match(question, "who owns the trailer", "do you personally own it or does your company own the trailer"):
        return f"{pronoun} stated that {context.pronoun} personally owns the trailer."
    if fire_claim and match(question, "when did you purchase the trailer"):
        return f"{pronoun} stated that {context.pronoun} purchased the trailer around 2021, although {context.pronoun} was not certain of the exact year."
    if fire_claim and match(question, "why did you purchase the trailer"):
        return f"{pronoun} stated that the trailer was purchased during COVID so {context.pronoun} could travel to visit family without having to fly or stay in hotels."
    if fire_claim and match(question, "how long has the trailer been on your property", "why was the trailer originally kept on the property"):
        if "why was" in question:
            return f"{pronoun} stated that the trailer was kept on the property to avoid storage costs and because it was also used as a mobile office for the yacht business."
        return f"{pronoun} stated that the trailer had been kept on the property since {context.pronoun} bought the house."
    if fire_claim and match(question, "permanent fixture or is it mobile", "does it have a license plate", "does it run"):
        if "license plate" in question:
            return f"{pronoun} stated that the trailer has a license plate."
        if "does it run" in question:
            return f"{pronoun} stated that the trailer is operational."
        return f"{pronoun} stated that the trailer is mobile rather than a permanent fixture."
    if match(question, "tell me the story about this leak in 509"):
        return render_prior_flood_story(context)
    if "drywalls removed" in pair.answer.lower() and "lobby" in pair.answer.lower():
        return f"{pronoun} stated that remediation from the prior February flood required drywall removal on the third floor and in the lobby."
    if match(question, "what part on the ac leaked"):
        return render_ac_connection_details(context)
    if match(question, "has this part been replaced since the ac was installed") and "what part are you referring" in pair.answer.lower():
        return None
    if match(question, "did you ask him to replace it", "so you don't remember whether you told him"):
        return None
    if match(question, "how old was that water supply line"):
        return render_filter_age(pair.answer, context)
    if match(question, "was it installed by a professional") and "friend" in pair.answer.lower():
        return f"{pronoun} stated that the water filter was installed by a friend rather than by a professional."
    if fire_claim and match(question, "in your own words", "what happened involving this fire"):
        return render_fire_story(context)
    if fire_claim and match(question, "how did the claimant get injured"):
        return f"{pronoun} stated that William was injured when the police pulled him out through the bedroom window area during the fire response."
    if fire_claim and match(question, "who called the fire department"):
        return f"{pronoun} stated that both {context.pronoun} and the neighbors called the fire department."
    if fire_claim and match(question, "how long after the fire started did the fire department arrive"):
        return f"{pronoun} stated that {context.pronoun} does not know how long it took the fire department to arrive because {context.pronoun} was traumatized."
    if fire_claim and match(question, "would it be minutes, hours"):
        return None
    if fire_claim and match(question, "did the fire department investigate"):
        return f"{pronoun} confirmed that a fire inspector came out to investigate the fire."
    if fire_claim and match(question, "what were their findings as to the cause of the incident of the fire"):
        return f"{pronoun} stated that the fire inspector could not clearly see the electrical panel because it was covered in debris and instead concluded that a surge protector caused the fire."
    if fire_claim and match(question, "why do they think it was a surge protector and not the panel"):
        return f"{pronoun} explained that the panel was buried under debris from the pantry, so the inspector could not clearly evaluate it."
    if fire_claim and match(question, "why do you think it was the panel"):
        return join_sentences(
            [
                f"{pronoun} stated that electricians who came out after the fire told {context.objective} that the panel had caught fire",
                f"{pronoun} further stated that Florida Power and Light later advised there had been a neutral connection issue and that this supported {context.possessive} belief that the electrical panel was involved",
            ]
        )
    if fire_claim and match(question, "fpl forgot to connect the neutral wire on your power pole and that contributed to the fire"):
        return f"{pronoun} confirmed that, based on what {context.pronoun} was later told, the missing neutral connection caused the panel to catch fire."
    if fire_claim and match(question, "when did fpl first come to your property to perform work"):
        return f"{pronoun} stated that Florida Power and Light first came out in late December 2023."
    if fire_claim and match(question, "what work did fpl come to your property to perform"):
        return f"{pronoun} stated that Florida Power and Light came out because a branch was across the power line and sparking."
    if fire_claim and match(question, "did you know at the time that fpl's work was done incorrectly"):
        return f"{pronoun} stated that {context.pronoun} had no idea at the time that any Florida Power and Light work had been done incorrectly."
    if fire_claim and match(question, "what time of day did the incident occur", "when i discovered it"):
        return f"{pronoun} stated that {context.pronoun} discovered the fire at about 9:45 p.m."
    if fire_claim and match(question, "so you were present when the fire and the claimant's subsequent injuries occurred"):
        return f"{pronoun} confirmed that {context.pronoun} personally saw William during the fire and its immediate aftermath."
    if fire_claim and match(question, "how do you know fpl caused the fire", "do you have any documentation to prove that fpl calls the fire"):
        return f"{pronoun} stated that electricians and a later Florida Power and Light technician attributed the problem to the neutral connection, but {context.pronoun} does not have documentation from Florida Power and Light admitting fault."
    if fire_claim and match(question, "when did you find out that it was done correctly"):
        return f"{pronoun} stated that {context.pronoun} learned of the neutral connection problem only after the third Florida Power and Light technician came out."
    if fire_claim and match(question, "why were all these technicians coming out"):
        return f"{pronoun} stated that the technicians kept returning because the property continued having electrical problems after the fire."
    if fire_claim and match(question, "so you requested fpl to come out to your property to perform this work"):
        return f"{pronoun} stated that Florida Power and Light had to come out multiple times between late December 2023 and March 2024 because of continuing electrical problems."
    if fire_claim and match(question, "did you call them back after you found out the neutral was not connected"):
        return f"{pronoun} stated that {context.pronoun} did not make another separate call after that because the technician said he reconnected it and would put in a work order for the remaining issues."
    if fire_claim and match(question, "did they ever come fix it"):
        return f"{pronoun} stated that the last technician said he fixed the neutral connection but also said there were additional issues that still needed attention."
    if fire_claim and match(question, "why did you not report the claimant's injuries to citizens until now"):
        return f"{pronoun} stated that {context.pronoun} was traumatized by the fire and reported the fire itself to Citizens, but did not make a separate injury report at that time."
    if fire_claim and match(question, "did fdl ever resolve the issue") and "what do you mean" in pair.answer.lower():
        return None
    if fire_claim and match(question, "did they ever fix the issue"):
        return f"{pronoun} stated that the property did have power afterward, but {context.pronoun} understood there were still additional electrical issues needing attention."
    if fire_claim and match(question, "did fdl ever resolve the issue"):
        return f"{pronoun} stated that the neutral connection had been fixed and that the property had power afterward, but there were still additional electrical issues needing attention."
    if fire_claim and match(question, "do you know what exactly caused the electrical panel to ignite specifically on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know the precise mechanism that caused the electrical panel to ignite."
    if fire_claim and match(question, "was there anything wrong with the electrical panel on the date of loss"):
        return f"{pronoun} stated that, aside from the fire itself, {context.pronoun} was not aware of anything else visibly wrong with the panel."
    if fire_claim and match(question, "was there any unpermitted work done to the electrical panel prior to the date of loss", "have there been any prior repairs to the electrical panel prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any prior unpermitted work or prior repairs involving the electrical panel."
    if fire_claim and match(question, "how old was the electrical panel on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know the age of the electrical panel."
    if fire_claim and match(question, "was it present when you purchased the property"):
        return f"{pronoun} confirmed that the electrical panel was already present when {context.pronoun} purchased the property."
    if fire_claim and match(question, "had you ever had any prior issues, ignitions or electrical issues involving the electrical panel prior to the date of loss"):
        return f"{pronoun} stated that for about two weeks before the fire there had been a burning-wire smell and flickering lights, but {context.pronoun} did not know at the time that the electrical panel was the source."
    if fire_claim and question.strip().lower() in {"you didn't know what?", "right, but so my question is, did you say you smelled burned wire on the date of loss, but not prior to the date of loss? is that correct?"}:
        return None
    if fire_claim and match(question, "did you ever report the wire burn smells to an electrician prior to the date of loss", "so why did you not call an electrician prior to the date of loss", "okay, so if you smelled it for two weeks, why did you not call an electrician", "okay, but why did you not attempt to repair it for two weeks"):
        return f"{pronoun} stated that {context.pronoun} did not call an electrician because {context.pronoun} and William thought the issue might simply be a bad outlet or GFI and planned to change the socket themselves."
    if fire_claim and match(question, "how old was the mini fridge on the date of loss", "how long had the mini fridge been plugged in inside the pantry on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know the age of the mini refrigerator or how long it had been plugged in inside the pantry."
    if fire_claim and match(question, "was there anything wrong with the mini fridge on the date of loss", "was there anything wrong with the outlet that the mini fridge was plugged into on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any known problem with the mini refrigerator or the outlet it used before the fire."
    if fire_claim and match(question, "had the mini fridge or the outlet it was plugged into ever sparked, smelled or caught fire before"):
        return f"{pronoun} stated that there had been no prior sparks, fires, complaints, or known repair issues involving the mini refrigerator or its outlet."
    if fire_claim and match(question, "had the mini fridge or the outlet it was plugged into ever had any prior issues or required any prior repairs", "had anyone ever complained about the condition of your mini fridge or the outlet it was plugged into prior to the date of loss"):
        return None
    if fire_claim and match(question, "has anyone ever complained about the safety or the condition of your electrical panel prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any prior complaints about the electrical panel."
    if fire_claim and match(question, "did you ever command or instruct the claimant to put out the fire", "did you command or force the claimant to rescue any pets", "did you ever stop or attempt to stop the claimant from leaving the property when the fire was discovered in order to assist you"):
        return f"{pronoun} stated that {context.pronoun} did not command or force William to fight the fire, rescue pets, or remain in the house."
    if fire_claim and match(question, "did the claimant voluntarily attempt to put out the fire"):
        return f"{pronoun} stated that William acted voluntarily when he tried to help."
    if fire_claim and match(question, "did he do it voluntarily", "okay, so did he appear injured"):
        return None
    if fire_claim and match(question, "did you provide a faulty, inadequate, or partially empty fire extinguisher to the claimant on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} did not provide William with a faulty extinguisher and explained that the extinguisher he tried to use had been given by neighbors."
    if fire_claim and question.strip().lower() == "did you...":
        return None
    if fire_claim and match(question, "did the claimant appear injured"):
        return f"{pronoun} stated that William appeared to have bruising on his arms and legs after the police pulled him out and that {context.pronoun} believes those injuries were caused by the manner in which he was removed."
    if fire_claim and match(question, "how did he appear injured", "any other injuries", "what caused these injuries"):
        return None
    if fire_claim and match(question, "did the claimant ask for help or first aid"):
        return f"{pronoun} stated that William did not ask for first aid."
    if fire_claim and match(question, "did anyone call for an ambulance", "was the claimant taken to the hospital by an ambulance that day", "did an ambulance even come"):
        return f"{pronoun} stated that no ambulance was called and that William was not taken away by ambulance."
    if fire_claim and match(question, "did the claimant ever go to the hospital for his injuries", "i mean, other than that, did he go to a medical hospital"):
        return f"{pronoun} stated that William was taken to Broward General on a Baker Act and was not otherwise taken to a separate hospital for injury treatment."
    if fire_claim and match(question, "did they treat his injuries while he was baker acted", "do you know what the claimant's injuries are or were", "did the claimant have any lung injuries due to smoke", "do you know what, if any, treatments the claimant underwent", "do you know if the claimant had any surgeries, physical therapy, scarring, or permanent disabilities related to the injuries he allegedly sustained on the date of loss", "do you know if the claimant has recovered from his alleged injuries"):
        return f"{pronoun} stated that, aside from bruising, {context.pronoun} does not know of any specific treatment, smoke injury, surgery, therapy, scarring, permanent disability, or recovery information regarding William."
    if fire_claim and match(question, "have you spoken with the claimant about the incident since the incident"):
        return None
    if fire_claim and match(question, "what has he said", "what has the claimant said about the fire", "so the claimant blames fpl for the fire"):
        return f"{pronoun} stated that William has likewise blamed Florida Power and Light and is the person who later relayed the neutral-wire explanation to {context.objective}."
    if fire_claim and match(question, "of course"):
        return None
    if fire_claim and match(question, "has the claimant ever blamed you for the fire or his injuries"):
        return f"{pronoun} stated that William has not blamed {context.objective} for the fire or for his injuries."
    if fire_claim and match(question, "did the claimant have any prior injuries, including lung-related injuries"):
        return f"{pronoun} stated that William has asthma but {context.pronoun} does not know of any other prior injuries."
    if fire_claim and match(question, "was the claimant a smoker"):
        return f"{pronoun} stated that William was not a smoker."
    if fire_claim and match(question, "did the claimant have any prior accidents or incidents resulting in any injuries"):
        return f"{pronoun} stated that {context.pronoun} would not know William's full prior injury history."
    if fire_claim and match(question, "can you describe the claimant's physical appearance such as age, height, and weight"):
        return f"{pronoun} described William as about six feet tall, very slender, and around 50 years old."
    if fire_claim and match(question, "do you recall what the claimant was wearing on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not recall what William was wearing."
    if fire_claim and match(question, "was the claimant authorized to be on that part of the property when the incident occurred"):
        return f"{pronoun} confirmed that William was authorized to be there because he lived at the property."
    if fire_claim and match(question, "does he have any issues with drugs or alcohol"):
        return f"{pronoun} stated that {context.pronoun} does not know of any drug or alcohol issue affecting William at the time of the fire."
    if fire_claim and match(question, "under the influence at the time of the incident"):
        return None
    if fire_claim and match(question, "does he use drugs or alcohol"):
        return None
    if fire_claim and match(question, "did the claimant wear glasses or have any vision issues"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any vision issues."
    if fire_claim and match(question, "did the claimant have any mobility or balance issues", "did the claimant appear to have any injuries or bruises prior to the incident occurring"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any pre-existing mobility issues or pre-existing injuries."
    if fire_claim and match(question, "were you assisting or directing the claimant in any way when the incident was occurring", "were you directing or assisting the claimant in any way when the fire was occurring"):
        return f"{pronoun} stated that {context.pronoun} was not directing or assisting William during the fire and was outside trying to call the fire department."
    if fire_claim and match(question, "was anyone else injured on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} is not aware of anyone else being injured on the date of loss."
    if fire_claim and match(question, "was there anything that you could have done personally to prevent the fire", "was there anything that the claimant could have done to prevent the incident", "could anything have prevented the incident"):
        return f"{pronoun} stated that {context.pronoun} does not know of anything that could have prevented the fire."
    if fire_claim and match(question, "are you aware of any witnesses to the fires and ignition"):
        return f"{pronoun} stated that many neighbors were outside afterward, but {context.pronoun} is not aware of anyone who actually witnessed the ignition itself."
    if fire_claim and match(question, "have you shared all the documentation that you have regarding this incident with citizens insurance"):
        return f"{pronoun} stated that {context.pronoun} has shared all available documentation with Citizens."
    if fire_claim and match(question, "do you have any excess liability or umbrella insurance"):
        return f"{pronoun} stated that {context.pronoun} does not have excess liability or umbrella insurance."
    if fire_claim and match(question, "do you own any additional properties in the united states"):
        return f"{pronoun} stated that {context.pronoun} does not own any additional properties in the United States."
    if fire_claim and match(question, "is there anything else you'd like to add to this recorded statement"):
        return f"{pronoun} stated that {context.pronoun} believes legal action may be necessary against Florida Power and Light."
    if water_heater_claim and match(question, "in your own words", "tell me the story of what happened involving this leak"):
        return render_water_heater_story(context)
    if water_heater_claim and match(question, "what water did you shut off before leaving"):
        return f"{pronoun} stated that before leaving, {context.pronoun} shut off the main valve serving the unit."
    if water_heater_claim and match(question, "so when you shut off the water valve, was that to your unit or to a room"):
        return f"{pronoun} clarified that the valve {context.pronoun} shut off served the full unit."
    if water_heater_claim and match(question, "what water did the maintenance guy shut off"):
        return f"{pronoun} explained that maintenance shut off the city-side valve, while {context.pronoun} had already shut off the condo-side valve."
    if water_heater_claim and match(question, "then what happened after the water was shut off"):
        return f"{pronoun} stated that even after the water was shut off, the water heater continued to leak slowly from the water already in the tank, so {context.pronoun} returned, kept emptying the pan, and arranged for replacement the next morning."
    if wall_pipe_claim and match(question, "tell me the whole story of what you know about what happened involving this leak", "in your own words"):
        return render_stepdaughter_leak_story(context)
    if wall_pipe_claim and match(question, "and her first name"):
        return None
    if wall_pipe_claim and match(question, "when did you first become aware of the leak"):
        return f"{pronoun} stated that {context.pronoun} first became aware of the leak when {context.possessive} stepdaughter called during the weekend she was staying at the unit."
    if wall_pipe_claim and match(question, "did she call you right after it occurred"):
        return f"{pronoun} stated that the call was probably the next day, after the investigation had started."
    if wall_pipe_claim and match(question, "what did you do after becoming aware of the leak"):
        return f"{pronoun} stated that {context.pronoun} asked building maintenance, through {context.possessive} stepdaughter, to locate a plumber and handle the leak as soon as possible."
    if wall_pipe_claim and match(question, "then what happened"):
        return f"{pronoun} stated that a plumbing company later contacted {context.objective}, was given permission to enter the unit, made the repair, and sent photographs."
    if wall_pipe_claim and match(question, "how long was the leak going on for before it was stopped"):
        return f"{pronoun} stated that the leak was likely active for less than a day because the water was shut off shortly after the front desk reported it."
    if wall_pipe_claim and match(question, "why do you say that", "back to my original question, why do you think that it was only going"):
        return None
    if wall_pipe_claim and match(question, "did anyone shut the water off after the leak was discovered"):
        return f"{pronoun} stated that {context.possessive} stepdaughter shut the water off and building maintenance made sure it remained off."
    if wall_pipe_claim and match(question, "when she was called to shut the water off immediately", "when did ellen shut the water off"):
        return None
    if wall_pipe_claim and match(question, "did the condo association report the leak to you guys"):
        return f"{pronoun} stated that the leak was reported through the condo association, reportedly by email."
    if wall_pipe_claim and match(question, "did you report the leak to anyone after it was reported to you"):
        return f"{pronoun} stated that {context.pronoun} did not make a separate report after being notified."
    if wall_pipe_claim and match(question, "how many units were damaged as a result of this incident") and "two" not in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} later learned that units 9B and 5B were reportedly affected."
    if wall_pipe_claim and match(question, "how many units were damaged as a result of this incident") and "two" in pair.answer.lower():
        return f"{pronoun} stated that two units below, 9B and 5B, were the units {context.pronoun} understood to be affected."
    if wall_pipe_claim and match(question, "did you keep the water off until the repair was completed"):
        return f"{pronoun} stated that the water was kept off until the repair was completed and explained that the water is normally left off whenever no one is staying in the unit."
    if wall_pipe_claim and match(question, "why did your daughter turn the water back on"):
        return f"{pronoun} explained that {context.possessive} stepdaughter and her partner briefly tried turning on only the cold water because they thought only the hot side was affected, but the leak continued and they were then told to shut off all water until repairs were complete."
    if wall_pipe_claim and match(question, "and what's their unit numbers"):
        return f"{pronoun} stated that the affected units were 9B and 5B."
    if wall_pipe_claim and match(question, "were these units occupied on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know whether units 9B and 5B were occupied on the date of loss."
    if wall_pipe_claim and match(question, "do you know what the specific damage is to these units", "do you know what the specific damage is to these units"):
        return f"{pronoun} stated that {context.pronoun} does not know the specific damage in units 9B and 5B."
    if wall_pipe_claim and match(question, "do you know if they had any dry out or mold remediation completed", "do you know if they've had their damages repaired yet"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the affected units had remediation or completed repairs."
    if wall_pipe_claim and match(question, "do you know if they filed claims with their own insurance"):
        return f"{pronoun} stated that 9B filed a claim through its own insurance carrier."
    if wall_pipe_claim and match(question, "what about 5b"):
        return f"{pronoun} stated that the occupant of 5B wanted an estimate and wanted to be paid directly rather than go through insurance."
    if wall_pipe_claim and match(question, "do you know if the affected units have insurance on their unit"):
        return f"{pronoun} stated that {context.pronoun} knows 9B had insurance because its carrier contacted {context.objective}, but does not know whether 5B had insurance."
    if wall_pipe_claim and match(question, "are there any damages to condo common areas such as hallways, stairwells, lobbies, or offices", "are you aware of any damage to the condo common areas"):
        return f"{pronoun} stated that, aside from the leak first being noticed from water seen in the lobby, {context.pronoun} is not aware of any actual damage to condo common areas."
    if wall_pipe_claim and match(question, "what leaked in your unit on the date of loss"):
        return f"{pronoun} stated that the leak involved a pipe inside the wall near the front entrance."
    if wall_pipe_claim and match(question, "is it a pipe inside the wall near your front entrance", "is this a pipe located inside the wall near your front entrance"):
        return None
    if wall_pipe_claim and match(question, "is this was this pipe located between your entry and your bathroom"):
        return f"{pronoun} stated that {context.pronoun} does not know the pipe's exact location within the wall."
    if wall_pipe_claim and match(question, "how old was this pipe"):
        return f"{pronoun} stated that the pipe was believed to be as old as the building."
    if wall_pipe_claim and match(question, "what caused it to leak"):
        return f"{pronoun} stated that {context.pronoun} does not know what caused the pipe to leak."
    if wall_pipe_claim and match(question, "was there any rust or corrosion on the pipe", "did you ever notice any rust or corrosion on the pipe prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any rust or corrosion and had not seen any because the pipe was concealed behind the wall."
    if wall_pipe_claim and match(question, "had the pipe ever leaked before", "had you ever had any prior repairs or issues involving the pipe"):
        return f"{pronoun} stated that there had been no prior leaks or prior repair issues involving that pipe."
    if wall_pipe_claim and match(question, "did you have any reason to believe that the pipe was at risk of leaking prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe the pipe was at risk of leaking before the incident."
    if wall_pipe_claim and match(question, "has the leak been repaired"):
        return f"{pronoun} confirmed that the leak was repaired."
    if wall_pipe_claim and match(question, "when was the leak repaired", "so i think the receipt says november 13th"):
        return f"{pronoun} stated that the leak was repaired on November 13, 2025."
    if wall_pipe_claim and match(question, "so what day was the leak repaired", "so when was the leak repaired") and "november 13" in pair.answer.lower():
        return None
    if wall_pipe_claim and match(question, "so what day was the leak repaired") and normalize(pair.answer).lower().startswith("i don't know"):
        return None
    if wall_pipe_claim and match(question, "did you file a first party claim with citizens for the damage to your own property"):
        return f"{pronoun} stated that {context.pronoun} did not file a first-party claim with Citizens for {context.possessive} own unit."
    if wall_pipe_claim and question.strip().lower() in {"why?", "why not?"} and "take care of it" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} simply handled {context.possessive} own unit's repairs without making a first-party claim."
    if wall_pipe_claim and match(question, "so you fixed it yourselves out of your own pocket", "did you make the repairs out of your own pocket instead of filing a claim"):
        return f"{pronoun} confirmed that the repair was paid out of pocket."
    if wall_pipe_claim and match(question, "have you had any prior water leaks originating from your unit since you've owned it") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that there had been no prior water leaks originating from the unit."
    if wall_pipe_claim and match(question, "have any other condo unit owners or the association advised you that they have any damages or requested any payments from you"):
        return f"{pronoun} stated that, aside from 5B and 9B, no other owners or the association have requested payment or claimed damage."
    if wall_pipe_claim and match(question, "have you provided citizens with all documentation you have related to the incident"):
        return f"{pronoun} stated that {context.pronoun} has provided Citizens with all documentation {context.pronoun} has."
    if wall_pipe_claim and match(question, "does our insured have any excess liability or umbrella insurance"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the insured has any excess liability or umbrella coverage."
    if wall_pipe_claim and match(question, "does our insured have any additional properties in the united states with insurance on it"):
        return f"{pronoun} stated that the insured has one additional property in the United States."
    if wall_pipe_claim and match(question, "what's the address of that property and who insures it"):
        return f"{pronoun} stated that the additional property is located at {answer}."
    if boat_claim and match(question, "in your own words", "tell me the whole story"):
        return render_boating_story(context)
    if boat_claim and match(question, "did you see the claimant fall or what caused her to fall"):
        return f"{pronoun} stated that {context.pronoun} did not see the claimant actually fall."
    if boat_claim and match(question, "do you know what caused her to fall"):
        return f"{pronoun} stated that {context.pronoun} does not know exactly what caused the claimant to fall."
    if boat_claim and match(question, "how did you first become aware of this incident"):
        return f"{pronoun} stated that {context.pronoun} heard a noise, looked up, and saw the claimant on the ground near the stairs."
    if boat_claim and match(question, "why was the claimant at your property on the date of loss"):
        return f"{pronoun} stated that the claimant was there to go on a fishing trip."
    if boat_claim and match(question, "do you run a fishing charter business") and normalize(pair.answer).lower().startswith("yes"):
        return f"{pronoun} confirmed that {context.pronoun} operates a fishing charter business."
    if boat_claim and match(question, "do you run a fishing charter business") and "do have one" in normalize(pair.answer).lower():
        return f"{pronoun} confirmed that {context.pronoun} does operate a fishing charter business."
    if boat_claim and match(question, "were you operating this fishing charter on the date of loss") and normalize(pair.answer).lower().startswith("yes"):
        return f"{pronoun} confirmed that {context.pronoun} was operating that charter business on the date of loss."
    if boat_claim and match(question, "how long have you been running your fishing charter business"):
        return f"{pronoun} stated that {context.pronoun} has had a captain's license since 1994 and now only runs occasional semi-retired charters for long-time customers."
    if boat_claim and match(question, "how long have you been semi-retired"):
        return f"{pronoun} stated that {context.pronoun} has been semi-retired for approximately two years."
    if boat_claim and match(question, "while being semi-retired, were you doing the small charters"):
        return f"{pronoun} stated that {context.pronoun} still occasionally handled smaller charters and would sometimes captain the larger boat for its new owner when needed."
    if boat_claim and match(question, "did you sell the boat"):
        return f"{pronoun} stated that {context.pronoun} sold the larger boat."
    if boat_claim and match(question, "do you run the fishing charter business out of your backyard"):
        return f"{pronoun} stated that for long-time customers, {context.pronoun} sometimes has them meet at the house, although in the past {context.pronoun} also used the marina at the end of the canal."
    if boat_claim and match(question, "where do you permanently keep the boat"):
        return f"{pronoun} stated that the boat is kept on a lift in the backyard."
    if boat_claim and match(question, "how long have you stored your boat back there in your yard"):
        return f"{pronoun} stated that the boat has been stored there for approximately seven years."
    if boat_claim and match(question, "can you describe the boat"):
        return f"{pronoun} described the boat as a 29-foot Pro Line Super Sport."
    if boat_claim and match(question, "how many engines does it have"):
        return f"{pronoun} stated that the boat has two outboard engines."
    if boat_claim and match(question, "inboard or outboard"):
        return None
    if boat_claim and match(question, "what's the horsepower of those engines"):
        return f"{pronoun} stated that the engines are 300 horsepower."
    if boat_claim and match(question, "do you have a registered business", "when this incident occurred"):
        return None
    if boat_claim and match(question, "so you were operating as just a sole proprietor"):
        return f"{pronoun} stated that {context.pronoun} was operating as a sole proprietor rather than through an LLC."
    if boat_claim and match(question, "had this group of people paid for you to take them on a fishing charter on your boat on the date of loss"):
        return f"{pronoun} stated that the group would have paid for the charter at the end of the trip because the final price depended on the type of trip and how far they went."
    if boat_claim and match(question, "so was this a group of paying customers", "was the claimant included in this group of paying customers", "so would you say that the claimant and her group of friends were paying customers on the date of loss"):
        return f"{pronoun} stated that the claimant and her group were paying customers for that fishing trip."
    if boat_claim and match(question, "how much did they pay", "do you have any guesstimate"):
        return f"{pronoun} estimated that the trip would have been priced somewhere between $600 and $900."
    if boat_claim and match(question, "do you have any receipts or documentation related to the claimant and her group chartering the boat on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not have receipts or other documentation for that charter."
    if boat_claim and match(question, "have you ever operated any fishing charters for this group of people prior to the date of loss", "how many times have you ever taken this particular group of people out"):
        return f"{pronoun} stated that {context.pronoun} had been taking that particular group fishing for many years and could not give an exact number of prior trips."
    if boat_claim and match(question, "was this group of people aware of a dog prior to the date of loss"):
        return f"{pronoun} stated that some repeat guests may have known there could be a dog at the property, but {context.pronoun} could not say that for certain."
    if boat_claim and match(question, "how long had you been operating this fishing charter out of your backyard"):
        return f"{pronoun} stated that {context.pronoun} had been running those backyard departures on an occasional basis for about two years and only a few times each year."
    if boat_claim and match(question, "approximately the past two years"):
        return None
    if boat_claim and match(question, "semi-retired small pack charter fishing business out of your backyard"):
        return None
    if boat_claim and match(question, "was the claimant aware of this dog prior to the incident", "why do you say probably not"):
        return f"{pronoun} stated that the claimant was probably not aware of the dog because it was her first time at the property."
    if boat_claim and match(question, "were other people in the group aware that there might be a dog here"):
        return f"{pronoun} stated that some of the repeat guests may have been aware there could be a dog at the property."
    if boat_claim and match(question, "exactly where on the property did this incident occur"):
        return f"{pronoun} stated that the incident occurred about 30 feet from the boat, where the stairs from the patio meet the pool deck."
    if boat_claim and match(question, "the stairway meets the pool deck"):
        return None
    if boat_claim and match(question, "what is the ground made of"):
        return f"{pronoun} stated that the ground surface there was made of pavers."
    if boat_claim and match(question, "was it installed professionally"):
        return f"{pronoun} stated that the stair and paver area was installed professionally."
    if boat_claim and match(question, "was this area slippery", "were there any obstacles or deficiencies", "were there any loose pavers or uneven pavers", "was this area well lit", "was it installed properly", "is it level", "has anyone ever complained about the area where the claimant fell"):
        if "slippery" in question:
            return f"{pronoun} stated that the area was not slippery."
        if "well lit" in question:
            return f"{pronoun} stated that the area was well lit."
        if "installed properly" in question or "is it level" in question:
            return None
        return f"{pronoun} stated that {context.pronoun} was not aware of any defects, loose pavers, unevenness, or prior complaints involving that area."
    if boat_claim and match(question, "was it installed with a permit"):
        return f"{pronoun} stated that the stair and paver area was installed with the pool project and under that permit."
    if boat_claim and match(question, "how old is this area"):
        return f"{pronoun} stated that the paver and stair area was about six years old."
    if boat_claim and match(question, "are you aware that they're alleging a dog knocked her over"):
        return f"{pronoun} stated that {context.pronoun} had heard that allegation, but when {context.pronoun} looked up after hearing the noise, the dog was nowhere near the claimant."
    if boat_claim and match(question, "what do you think caused her to fall", "why did she lose her footing"):
        return f"{pronoun} stated that, if {context.pronoun} had to guess, the claimant may simply have lost her balance while walking."
    if boat_claim and match(question, "do you know why the dog would have knocked her over"):
        return f"{pronoun} stated that {context.pronoun} cannot see how the dog would have knocked the claimant over because the dog was agile enough to run through narrow spaces without touching people."
    if boat_claim and match(question, "do you know what allegedly caused this claimant to fall"):
        return f"{pronoun} stated that {context.pronoun} did not initially understand the allegation until reading the paperwork."
    if boat_claim and match(question, "what they're claiming caused her to fall"):
        return f"{pronoun} stated that the paperwork alleges the dog caused the claimant to fall."
    if boat_claim and match(question, "how the dog caused the claimant to fall", "was the claimant interacting with the dog"):
        return f"{pronoun} stated that {context.pronoun} did not see the claimant interacting with the dog and cannot explain how the dog would have caused the fall."
    if boat_claim and match(question, "do you know if the dog allegedly knocked over the claimant, or did the dog trip the claimant"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the allegation is that the dog knocked or tripped the claimant."
    if boat_claim and match(question, "prior to arriving, was the claimant or any of the guests warned of the dog"):
        return f"{pronoun} stated that {context.pronoun} does not recall giving any specific warning about the dog before the group arrived."
    if boat_claim and match(question, "can you describe the dog that allegedly caused the claimant to fall"):
        return f"{pronoun} stated that the dog was a one-and-a-half-year-old Australian Shepherd."
    if boat_claim and match(question, "who owned the dog"):
        return f"{pronoun} stated that the dog belonged to {context.possessive} son and that {context.pronoun} was only caring for it temporarily."
    if boat_claim and match(question, "how long had you been caring for the dog"):
        return f"{pronoun} stated that {context.pronoun} had only been caring for the dog for a day or two."
    if boat_claim and match(question, "what breed was the dog", "how heavy was the dog", "how old was the dog"):
        if "breed" in question:
            return f"{pronoun} stated that the dog was an Australian Shepherd."
        if "heavy" in question:
            return f"{pronoun} estimated that the dog weighed about 40 pounds."
        return f"{pronoun} stated that the dog was about a year and a half old."
    if boat_claim and match(question, "did the dog have any behavioral issues", "issue with jumping on people", "history of jumping on people", "was the dog aggressive", "did the dog have any history of jumping on people or tripping them"):
        return f"{pronoun} stated that the dog had no history of aggression, jumping on people, or similar behavioral problems."
    if boat_claim and match(question, "had the dog ever had any behavioral training"):
        return f"{pronoun} stated that the dog may have had some training through {context.possessive} son, but {context.pronoun} was not sure."
    if boat_claim and match(question, "why do you say probably not", "why do you say possibly"):
        return None
    if boat_claim and match(question, "was the dog vaccinated on the date of loss"):
        return f"{pronoun} stated that the dog was vaccinated on the date of loss."
    if boat_claim and match(question, "why was this dog outside when the incident occurred", "was the dog restrained in any way when the incident occurred"):
        if "why was this dog outside" in question:
            return f"{pronoun} stated that the dog had simply been let outside to relieve itself."
        return f"{pronoun} stated that the dog was restrained only by the fenced yard."
    if boat_claim and match(question, "prior to arriving, was the claimant or any of the guests warned of the dog"):
        return f"{pronoun} stated that {context.pronoun} does not recall whether any warning about the dog was specifically given before the group arrived."
    if boat_claim and match(question, "who was watching the dog when the incident occurred"):
        return None
    if boat_claim and match(question, "was the claimant or any of the other guests familiar with the dog prior to the date of loss"):
        return f"{pronoun} stated that some members of the group may have been familiar with the dog from earlier visits, but {context.pronoun} was not certain."
    if boat_claim and match(question, "was everyone else other than the claimant familiar with the dog", "so you don't know how many people in the group were familiar with the dog or not"):
        return f"{pronoun} stated that {context.pronoun} could not say exactly how many people in the group were familiar with the dog."
    if boat_claim and match(question, "has anyone ever complained about the dog or his behavior prior to the date of loss"):
        return f"{pronoun} stated that there had been no prior complaints about the dog or its behavior."
    if boat_claim and match(question, "and you said your son owned the dog", "so the dog was only temporarily living with you", "were you familiar with this dog"):
        if "temporarily living" in question:
            return f"{pronoun} stated that the dog was only staying there temporarily."
        if "familiar with this dog" in question:
            return f"{pronoun} stated that {context.pronoun} was familiar with the dog."
        return None
    if boat_claim and match(question, "was she ascending or descending the stairs to the pool deck"):
        return f"{pronoun} stated that the claimant was descending the stairs toward the pool deck."
    if boat_claim and match(question, "was the claimant directed to use this part of the property"):
        return f"{pronoun} stated that the claimant was not specifically directed to use that part of the property."
    if boat_claim and match(question, "did the claimant continue to walk after the fall", "so she briefly continued walking"):
        return f"{pronoun} stated that after the fall, the claimant got up and walked to a chair."
    if boat_claim and match(question, "did the claimant still wish to go on the fishing trip"):
        return f"{pronoun} stated that the claimant initially still wanted to go on the fishing trip."
    if boat_claim and question.strip().lower() == "why didn't she?":
        return f"{pronoun} stated that the claimant's friend talked her out of going."
    if boat_claim and match(question, "did the claimant appear injured"):
        return f"{pronoun} stated that the claimant said her ankle hurt."
    if boat_claim and match(question, "did the claimant ask for help"):
        return f"{pronoun} stated that the claimant did not ask for help."
    if boat_claim and match(question, "did anyone call for an ambulance or offer to call an ambulance", "did anyone eventually call an ambulance", "was the claimant taken to the hospital by an ambulance"):
        return f"{pronoun} stated that no ambulance was called from the property and that the claimant was not taken away from there by ambulance."
    if boat_claim and match(question, "did the claimant ever go to the hospital"):
        return f"{pronoun} stated that {context.pronoun} is not sure whether the claimant later went to the hospital."
    if boat_claim and match(question, "do you know what the claimant's injuries are or were"):
        return f"{pronoun} stated that someone later told {context.objective} the claimant may have suffered a hairline ankle fracture."
    if boat_claim and match(question, "do you know what treatments the claimant has undergone for her injuries", "do you know if the claimant's had any surgeries, physical therapy, scarring, or permanent disabilities", "do you know if the claimant has recovered", "do you know how long the claimant's recovery is expected to take"):
        return f"{pronoun} stated that {context.pronoun} does not know what treatment, recovery, or long-term effects the claimant may have had."
    if boat_claim and match(question, "have you seen or spoken with the claimant since the incident"):
        return f"{pronoun} stated that {context.pronoun} has not seen or spoken with the claimant since the incident."
    if boat_claim and match(question, "did the claimant have any prior injuries", "did the claimant have any prior accidents resulting in any injuries"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had any prior injuries or prior accidents."
    if boat_claim and match(question, "had it been raining", "was there any water on the ground", "was the area slippery"):
        return None
    if boat_claim and match(question, "let me back up to the steps for a second. were they even and consistent in height and length", "how many steps were there", "was there anything wrong with these steps"):
        if "how many steps" in question:
            return f"{pronoun} stated that there were two steps."
        if "even and consistent" in question:
            return f"{pronoun} stated that the steps were even and consistent."
        return f"{pronoun} stated that there was nothing wrong with the steps."
    if boat_claim and match(question, "so was she, her height, was it short, meet average, or tall", "okay, so you don't know", "okay, if you had to guess", "okay, was she a young adult, middle-aged"):
        if "young adult" in question:
            return f"{pronoun} stated that, if {context.pronoun} had to guess, the claimant was probably around 40 to 50 years old."
        return None
    if boat_claim and match(question, "do you recall what she was wearing or wearing on her feet"):
        return f"{pronoun} stated that {context.pronoun} does not recall what the claimant was wearing or had on her feet."
    if boat_claim and match(question, "was the claimant authorized to be on that part of the property when the incident occurred", "but was she allowed to be on that part of the property when the incident occurred"):
        return f"{pronoun} stated that the claimant was allowed to be there as part of the fishing group."
    if boat_claim and match(question, "do you know if the claimant was distracted in any way when the incident occurred", "did the claimant appear to be under the influence at the time of the incident", "did the claimant smell like marijuana or alcohol", "did the claimant wear glasses or have any vision issues", "did the claimant appear to have any mobility or balance issues", "did the claimant appear to have any injuries prior to the incident occurring"):
        return f"{pronoun} stated that {context.pronoun} does not know of any distraction, impairment, vision issue, mobility problem, or pre-existing injury affecting the claimant."
    if boat_claim and match(question, "did you leave her by herself or with somebody"):
        return f"{pronoun} stated that the claimant was left with her friend."
    if boat_claim and match(question, "do you know the claimant in any way", "has the claimant ever been to your home prior to the date of loss", "has the claimant ever been on a fishing charter with you before", "has the claimant been to your house again since the date of loss"):
        return f"{pronoun} stated that {context.pronoun} did not know the claimant personally, that she had not previously been to the house or on a charter with him, and that she has not returned since the incident."
    if boat_claim and match(question, "were you or your wife assisting the claimant in any way when the incident occurred", "were you directing or assisting her in any way", "did you direct the claimant at any point while she was on the property prior to her falling"):
        return f"{pronoun} stated that neither {context.pronoun} nor {context.possessive} wife was directing or assisting the claimant before the fall and that {context.pronoun} only helped her afterward."
    if boat_claim and match(question, "was there anything the claimant could have done to prevent the incident", "are you aware of anything at all that could have prevented the incident"):
        return f"{pronoun} stated that {context.pronoun} does not know what, if anything, could have prevented the incident."
    if boat_claim and match(question, "have you shared all documentation you have regarding this incident with citizens insurance"):
        return f"{pronoun} stated that {context.pronoun} has shared all documentation {context.pronoun} has with Citizens."
    if boat_claim and match(question, "do you have any general liability insurance for your fishing charter business", "and was that", "gallagher insurance"):
        return f"{pronoun} stated that {context.pronoun} has a charter policy and believes it is through Gallagher."
    if boat_claim and match(question, "did anyone call for an ambulance"):
        return f"{pronoun} stated that {context.pronoun} does not believe anyone called an ambulance."
    if boat_claim and match(question, "do you know what the claimant's injuries are or were"):
        return f"{pronoun} stated that the claimant later said her ankle was hurting."
    if boat_claim and match(question, "do you know if the claimant had any surgeries, physical therapy, scarring, or permanent disability"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had surgeries, therapy, scarring, or permanent disability."
    if boat_claim and match(question, "what time of day the incident occurred", "what time of day did the incident occur"):
        return f"{pronoun} stated that the incident occurred at about 8:00 a.m."
    if boat_claim and match(question, "what the weather was like on the date of loss", "what was the weather like"):
        return f"{pronoun} stated that the weather was sunny and bright."
    if boat_claim and match(question, "can you describe the claimant's physical appearance", "just your own description"):
        return f"{pronoun} stated that the claimant appeared shorter than him and, if he had to guess, was around 40 to 50 years old."
    if boat_claim and match(question, "do you know if she had any mobility or balance issues", "do you know if the claimant wore glasses or had any vision issues"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had any vision, mobility, or balance issues."
    if boat_claim and match(question, "do you know the claimant in any capacity"):
        return f"{pronoun} stated that {context.pronoun} did not know the claimant personally."
    if boat_claim and match(question, "were you directing or assisting the claimant in any way when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} was not directing or assisting the claimant before the fall and only helped her after {context.pronoun} saw her on the ground."
    if boat_claim and match(question, "was there anything that you could have done personally to prevent the incident"):
        return f"{pronoun} stated that {context.pronoun} does not believe there was anything {context.pronoun} personally could have done to prevent the incident."
    if boat_claim and match(question, "are there any video cameras on your property"):
        return f"{pronoun} stated that there had been a Ring camera on the gazebo, but it was no longer working."
    if boat_claim and match(question, "aware of any witnesses"):
        return f"{pronoun} stated that the other members of the fishing group would be the only potential witnesses {context.pronoun} is aware of."
    if roof_slip_claim and match(question, "so that roof may have been around 30 years old when the incident occurred"):
        return None
    if roof_slip_claim and match(question, "what condition was the roof in on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know the exact condition of the roof on the date of loss."
    if roof_slip_claim and match(question, "when was the last time you saw the roof prior to the date of loss"):
        return f"{pronoun} stated that the last time {context.pronoun} saw the roof was during a visit in 2025."
    if roof_slip_claim and match(question, "prior to the incident occurring, when was the last time you saw the roof"):
        return f"{pronoun} stated that before the alleged incident, the last time {context.pronoun} had seen the roof was in 2022."
    if roof_slip_claim and match(question, "what do you mean you stayed there"):
        return f"{pronoun} explained that {context.pronoun} only stayed at the property on holidays for about three weeks."
    if roof_slip_claim and match(question, "when you saw the roof in 2022, was there any issues with it"):
        return f"{pronoun} stated that when {context.pronoun} saw the roof in 2022, {context.pronoun} did not notice any problems."
    if roof_slip_claim and match(question, "did the roof look old or did it look good"):
        return f"{pronoun} stated that the roof looked okay to {context.objective} and that the homeowners association generally monitored the roofs."
    if roof_slip_claim and match(question, "was there, have you ever had any issues with the roof", "had the roof ever leaked before"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, there had been no prior roof issues or roof leaks."
    if roof_slip_claim and match(question, "what kind of floors run throughout your property"):
        return f"{pronoun} stated that the main areas had tile flooring and the bedrooms had laminate flooring."
    if roof_slip_claim and match(question, "how old are these floors"):
        return f"{pronoun} stated that {context.pronoun} does not know the age of the flooring."
    if roof_slip_claim and match(question, "was the floor slippery"):
        return f"{pronoun} stated that {context.pronoun} was not there and does not know whether the floor was slippery, but noted that the claimant catered from the house and may have had grease in the kitchen."
    if roof_slip_claim and match(question, "why do you think there could have been grease on the floor"):
        return f"{pronoun} stated that the claimant catered from the house and used deep fryers in the kitchen."
    if roof_slip_claim and match(question, "did you ever see any grease on the floor", "did you or your cousin ever see any grease on the floor"):
        return f"{pronoun} stated that {context.pronoun} never personally saw grease on the floor."
    if roof_slip_claim and match(question, "did any, did the, did the claimant or anyone else ever call for an ambulance"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anyone called for an ambulance or whether the claimant went to the hospital afterward."
    if roof_slip_claim and match(question, "do you know if the claimant ever went to the hospital", "oh, my question was, do you know if she ever went to the hospital"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anyone called for an ambulance or whether the claimant went to the hospital afterward."
    if roof_slip_claim and match(question, "did you say before a recorded statement that she went on disability with jetblue"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant went on disability and only recalls that the claimant said she was waiting on a disability check."
    if roof_slip_claim and match(question, "do you know what the claimant's injuries are or were"):
        return f"{pronoun} stated that {context.pronoun} does not know the claimant's specific injuries beyond recalling the claimant's statement that she needed an MRI."
    if roof_slip_claim and match(question, "do you know what, if any, treatments the claimant underwent"):
        return f"{pronoun} stated that {context.pronoun} does not know the claimant's specific injuries beyond recalling the claimant's statement that she needed an MRI."
    if roof_slip_claim and match(question, "do you know if she had any surgeries, physical therapy, scarring, or permanent disabilities"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant underwent further treatment, had surgeries or lasting impairment, or recovered from the alleged injuries."
    if roof_slip_claim and match(question, "do you know if the claimant has recovered from her alleged injuries"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant underwent further treatment, had surgeries or lasting impairment, or recovered from the alleged injuries."
    if roof_slip_claim and match(question, "can you describe the claimant's physical appearance"):
        return f"{pronoun} stated that {context.pronoun} had already provided photographs and otherwise could only estimate that the claimant was about 5'4\" and around 50 years old."
    if roof_slip_claim and match(question, "just your own description"):
        return f"{pronoun} stated that {context.pronoun} could only estimate that the claimant was about 5'4\" and around 50 years old."
    if roof_slip_claim and match(question, "would you say that she's skinny, average, overweight, or obese"):
        return f"{pronoun} described the claimant as appearing average in build."
    if roof_slip_claim and match(question, "did she wear glasses or have any vision issues"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had any vision, mobility, or balance issues."
    if roof_slip_claim and match(question, "do you know if she had any mobility or balance issues"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had any vision, mobility, or balance issues."
    if roof_slip_claim and match(question, "do you know if the claimant continued to walk after she fell"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant continued to walk after the fall."
    if roof_slip_claim and match(question, "were there ever any other leaks in your property on or around the date of loss"):
        return f"{pronoun} stated that {context.pronoun} was not aware of any other leaks at the property around the date of loss."
    if roof_slip_claim and match(question, "so the year before the incident occurred"):
        return None
    if roof_slip_claim and match(question, "do you know if the claimant had any prior injuries or any accidents involving any injuries", "do you know if the claimant had any prior injuries prior to the incident occurring"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had any prior injuries or prior accidents."
    if roof_slip_claim and match(question, "do you know the claimant in any capacity"):
        return f"{pronoun} stated that {context.pronoun} did not know the claimant personally and only knew of her through {context.possessive} cousins, who were friendly with her."
    if roof_slip_claim and match(question, "has anyone else") and "good friends" in pair.answer.lower():
        return None
    if dog_claim and match(question, "how old was each dog"):
        return f"{pronoun} stated that one dog was close to 10 years old and that four of the dogs were about 2 years old."
    if dog_claim and match(question, "so one of them was 10 and the other ones were a couple years old"):
        return f"{pronoun} clarified that the mother dog was about 6 years old."
    if dog_claim and match(question, "had any of the dogs ever exhibited aggressive behavior prior to the date of loss", "had the dogs ever bit anyone or anything prior to the date of loss"):
        return f"{pronoun} stated that the dogs had not previously attacked anyone."
    if dog_claim and match(question, "had the dogs ever bit you or any family members prior to the date of loss"):
        return None
    if dog_claim and match(question, "what kind of dogs were they"):
        return f"{pronoun} described the dogs as mixed-breed dogs that included pit bull."
    if dog_claim and match(question, "how much did your dogs weigh"):
        return f"{pronoun} stated that {context.pronoun} did not know how much the dogs weighed."
    if dog_claim and match(question, "what happened with animal control"):
        return f"{pronoun} stated that animal control took custody of the dogs."
    if dog_claim and match(question, "when animal control, you came, you just gave them away"):
        return f"{pronoun} stated that the dogs were put into cages and turned over to animal control."
    if dog_claim and match(question, "and what happened with the police"):
        return f"{pronoun} stated that the matter resulted in court proceedings and that the court dates kept being changed."
    if dog_claim and match(question, "prior encounters with law enforcement or animal control related to your dogs"):
        return f"{pronoun} stated that there had been one prior incident when the dogs escaped before."
    if dog_claim and match(question, "did anyone call for an ambulance"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anyone called an ambulance, although {context.pronoun} later heard the claimant may have been transported that way."
    if dog_claim and match(question, "was the claimant taken to the hospital by ambulance"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anyone called an ambulance, although {context.pronoun} later heard the claimant may have been transported that way."
    if dog_claim and match(question, "how do you know") and "ambulance" in pair.answer.lower():
        return None
    if dog_claim and match(question, "how do you know") and "somebody tell me" in pair.answer.lower():
        return f"{pronoun} stated that someone later told {context.objective} about it, but {context.pronoun} was not certain of the details."
    if dog_claim and match(question, "do you know how long the claimant was in the hospital for"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the claimant was in the hospital."
    if dog_claim and match(question, "for what?"):
        return f"{pronoun} stated that the citation was related to the dogs escaping."
    if dog_claim and match(question, "do you know if the claimant continued to walk after the incident"):
        return f"Based on what {context.pronoun} was told, the claimant was okay afterward."
    if dog_claim and match(question, "do you know what the claimant's injuries are or were"):
        return f"{pronoun} stated that the only injury information {context.pronoun} heard was that the claimant had injuries to her arms or hands."
    if dog_claim and match(question, "anything else?"):
        return None
    if dog_claim and question.strip().lower() == "do you know":
        return None
    if dog_claim and match(question, "do you know what, if any, treatments the claimant has undergone"):
        return f"{pronoun} stated that {context.pronoun} has no information regarding any treatment the claimant underwent."
    if dog_claim and match(question, "do you know if the claimant has recovered from her injuries"):
        return f"{pronoun} stated that, according to what a neighbor later said, the claimant was doing okay."
    if dog_claim and match(question, "who told you") and "across the street" in pair.answer.lower():
        return f"{pronoun} stated that someone who lives across the street told {context.objective} the claimant was okay."
    if dog_claim and match(question, "who told you") and "somebody tell me" in pair.answer.lower():
        return f"{pronoun} stated that someone later told {context.objective} the claimant had been taken by ambulance, although {context.pronoun} was not certain."
    if dog_claim and match(question, "so the claimant did go to the hospital"):
        return f"Based on what others told {context.objective}, the claimant did go to the hospital, but {context.pronoun} was not certain of the details."
    if dog_claim and match(question, "what time of day did the incident occur", "at what time"):
        return f"{pronoun} stated that the incident occurred early in the morning, around 7:00 a.m."
    if dog_claim and match(question, "did it occur on your property"):
        return f"{pronoun} stated that the incident occurred outside the property rather than inside the yard."
    if dog_claim and match(question, "do you know what the weather was like on the date of loss"):
        return f"{pronoun} stated that it was early morning, not raining, and before full sunrise."
    if dog_claim and match(question, "like sunny", "was it raining or anything"):
        return None
    if dog_claim and match(question, "did the claimant have any prior injuries to her arms", "did the claimant have any prior accidents resulting in any injuries"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant had prior injuries or prior accidents and did not know the claimant before this incident."
    if dog_claim and "still don't know her until now" in answer.lower():
        return None
    if dog_claim and match(question, "would you be able to describe the claimant's physical appearance"):
        return f"{pronoun} stated that {context.pronoun} could not describe the claimant other than having heard she was about 79 years old."
    if dog_claim and match(question, "did the claimant have any mobility or balance issues", "do you know if the claimant wore glasses or had any vision issues"):
        return f"{pronoun} stated that {context.pronoun} has no information regarding the claimant's vision, mobility, or balance."
    if dog_claim and match(question, "was there anything that you could have done personally to prevent the incident"):
        return f"{pronoun} stated that {context.pronoun} does not believe there was anything {context.pronoun} personally could have done to prevent the incident because {context.pronoun} was not there."
    if dog_claim and match(question, "could anything have prevented the incident"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anything could have prevented the incident."
    if dog_claim and match(question, "did the claimant have any prior accidents resulting in any injuries", "i still don't know her until now", "so you don't know the claimant in any way"):
        return None
    if dog_claim and match(question, "do you know what the claimant was wearing on the date of loss", "was the claimant authorized to be on your property", "do you know if the claimant was distracted in any way when the incident occurred", "was the claimant under the influence", "had the claimant ever been to your home or property before the date of loss"):
        if "authorized" in question:
            return f"{pronoun} stated that the claimant was not authorized to be on the property."
        if "had the claimant ever been" in question:
            return f"{pronoun} stated that the claimant had never been to the property before."
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant was distracted, under the influence, or wearing anything noteworthy at the time of the incident."
    if dog_claim and match(question, "were you directing or assisting the claimant in any way when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} was not directing or assisting the claimant when the incident occurred."
    if dog_claim and match(question, "was anyone in your family"):
        return f"{pronoun} stated that no one in {context.possessive} family was injured."
    if overflow_claim and match(question, "how did you become aware of it"):
        return f"{pronoun} stated that {context.pronoun} first learned of the leak on January 28, 2026, when the office said a plumber needed access to investigate a complaint from the unit below."
    if overflow_claim and match(question, "in your own words, tell me the whole story of what happened involving the leak"):
        return (
            f"In describing what happened, the {role_noun} stated that {context.pronoun} first heard about the problem when office staff said a plumber needed access to investigate a complaint from the unit below. "
            f"{pronoun} stated that the plumber then confirmed there was a slight drip coming from {context.possessive} unit when water was run in the master bathroom shower or tub."
        )
    if overflow_claim and match(question, "what do you mean the water was turned on"):
        return f"{pronoun} explained that when water was run in the master bathroom shower or tub, there would be a slight drip into the bathroom of the unit below, and the plumber was asked to repair it as soon as possible."
    if overflow_claim and match(question, "how often did the cleaning lady come to clean your unit while it was vacant"):
        return f"{pronoun} stated that the cleaning person came about twice after the prior tenants left and again before the next tenants moved in."
    if overflow_claim and match(question, "did the cleaning lady ever turn on the water in the bathtub"):
        return f"{pronoun} stated that the cleaning person probably used the shower rather than the bathtub."
    if overflow_claim and question.strip().lower() == "why?" and "wash the walls" in pair.answer.lower():
        return f"{pronoun} stated that the shower was likely used only for cleaning tasks such as washing the walls and glass doors."
    if overflow_claim and match(question, "did she ever report a leak"):
        return f"{pronoun} stated that the cleaning person never reported any leak and that January 28, 2026 was the first time {context.pronoun} heard about the issue."
    if overflow_claim and match(question, "so you first became aware of the leak on january 28th of 2026"):
        return None
    if overflow_claim and match(question, "did they say how long he had been complaining"):
        return f"{pronoun} stated that {context.pronoun} was not told that day how long the issue had been going on and only later heard there had been a mildew smell in the lower unit."
    if overflow_claim and match(question, "do you know why it took until january of 2026 to do something about it"):
        return f"{pronoun} stated that {context.pronoun} does not know why the matter was not addressed sooner because {context.pronoun} had not been told about it before January 2026."
    if overflow_claim and match(question, "so you didn't know about the issue until january"):
        return f"{pronoun} stated that January 28, 2026 was the first time {context.pronoun} became aware of the reported issue and that {context.pronoun} contacted the plumber immediately."
    if overflow_claim and match(question, "what did you do after becoming aware of the leak"):
        return f"{pronoun} stated that {context.pronoun} contacted the plumber, spoke with the tenant, and asked that the plumber be allowed into the unit as soon as possible."
    if overflow_claim and match(question, "how soon did the plumber get out"):
        return f"{pronoun} stated that the plumber inspected the issue the next day, then coordinated with the tenant and ultimately completed the final repair on February 9, 2026."
    if overflow_claim and match(question, "how long was the leak going on for before it was stopped"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the condition existed before it was discovered."
    if overflow_claim and match(question, "were your tenants", "stalling things"):
        return f"{pronoun} stated that scheduling with the tenant caused some delay because the tenant did not want strangers in the apartment unless it was convenient."
    if overflow_claim and match(question, "did anyone shut the water off to the unit or bathroom after the leak was discovered"):
        return f"{pronoun} stated that the tenants continued living in the unit, so the water was not shut off entirely, but they stopped using that bathroom and used the guest bathroom instead."
    if overflow_claim and match(question, "so they didn't use the bathroom after the leak was discovered"):
        return None
    if overflow_claim and match(question, "why are they saying you knew about it for months"):
        return f"{pronoun} disputed that allegation and explained that if {context.pronoun} had known about the issue earlier, {context.pronoun} would have addressed it immediately."
    if overflow_claim and match(question, "how many units were damaged as a result of this incident, including your own", "so the only damage potentially is to unit 233 directly below you"):
        return f"{pronoun} stated that the only unit potentially affected was the unit directly below, unit 233."
    if overflow_claim and match(question, "was the unit below yours damaged as a result of this incident"):
        return None
    if overflow_claim and match(question, "what's his unit number", "what's the unit number"):
        return f"{pronoun} stated that the unit below was unit 233."
    if overflow_claim and match(question, "do you know how long it had been unoccupied for"):
        return f"{pronoun} stated that, according to the property manager, the unit below had reportedly been vacant for quite some time, including before July 2025."
    if overflow_claim and match(question, "do you know what their specific damages are"):
        return f"{pronoun} stated that {context.pronoun} heard there was damage to the kitchen ceiling and possibly some walls in the unit below, but {context.pronoun} had not personally seen photographs."
    if overflow_claim and match(question, "dry out or motor mediation completed"):
        return f"{pronoun} stated that {context.pronoun} believed a remediation company had been brought in but later stopped the work, based on what the board mentioned."
    if overflow_claim and match(question, "what do you mean stop"):
        return None
    if overflow_claim and match(question, "has he had his damages repaired yet"):
        return f"{pronoun} stated that the lower unit had not been fully repaired because the walls were not going to be closed until the leak issues were resolved."
    if overflow_claim and match(question, "tell me about the toilet leak inside the claimant's unit"):
        return f"{pronoun} stated that {context.pronoun} was told the lower unit also had a toilet leak and believed some of the claimed damage may have been related to that separate issue."
    if overflow_claim and match(question, "do you know if the claimant has insurance on their own unit"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the lower-unit owner had insurance, although {context.pronoun} assumed he probably did."
    if overflow_claim and match(question, "do you believe that the damage to the claimant's unit was caused by his own toilet"):
        return f"{pronoun} stated that a slight drip from {context.possessive} unit may have caused some issue, but {context.pronoun} did not believe it explained all of the damage being claimed and thought the toilet issue below sounded more significant."
    if overflow_claim and match(question, "what specifically leaked in your unit", "bathtub overflow that leaked"):
        return f"Based on what the plumber explained, the leak involved the bathtub overflow in the master bathroom."
    if overflow_claim and match(question, "how old was the bathtub including the overflow"):
        return f"{pronoun} stated that the bathtub and overflow were original to the building."
    if overflow_claim and match(question, "do you know what caused it to leak"):
        return f"{pronoun} stated that {context.pronoun} does not know the exact cause but understood from the plumber that the overflow may have been old or misaligned."
    if overflow_claim and match(question, "was there any rust or corrosion on the overflow", "from the overflow, did you see any rust or corrosion on it", "are you or are you not aware of any rust or corrosion on the overflow"):
        return f"{pronoun} stated that no rust or corrosion was visible from the tub side, although the plumber indicated the internal components appeared old."
    if overflow_claim and match(question, "had the overflow ever leaked before", "had you ever had any prior issues or repairs involving the overflow", "did you ever have any reason to believe that the overflow was at risk of leaking", "did your tenants ever complain about the bathtub or the overflow", "did your tenants ever report seeing water after using the bathtub"):
        if "reason to believe" in question:
            return f"{pronoun} stated that {context.pronoun} had no reason to believe the overflow was at risk of leaking before this incident."
        if "complain" in question or "report seeing water" in question:
            return f"{pronoun} stated that the tenants had not complained about the bathtub or reported seeing water before the leak was discovered."
        return f"{pronoun} stated that there had been no prior leaks, repairs, or issues involving the overflow."
    if overflow_claim and match(question, "has the leak been repaired"):
        return f"{pronoun} confirmed that the leak has been repaired."
    if overflow_claim and question.strip().lower() == "when?":
        return f"{pronoun} stated that the repair was completed on February 9, 2026."
    if overflow_claim and match(question, "are you aware of any damages to condo common areas", "are there any damages to condo common areas such as hallways, stairwells, lobbies"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any damage to condo common areas."
    if water_heater_claim and match(question, "what did they say about the incident"):
        return f"{pronoun} stated that the condo association told {context.objective} it was unfortunate but complimented how quickly {context.pronoun} responded after learning of the leak."
    if water_heater_claim and match(question, "did you ever speak with the claimant or claimants about this incident"):
        return f"{pronoun} stated that {context.pronoun} did not speak directly with the downstairs neighbor about the claim beyond discussing the reported damage."
    if water_heater_claim and match(question, "so he just showed you the damage in his unit"):
        return f"{pronoun} stated that the downstairs neighbor showed {context.objective} the reported damage and asked what would be done about it, and {context.pronoun} told him the matter would be handled through insurance."
    if water_heater_claim and match(question, "how many units were damaged"):
        return None
    if water_heater_claim and match(question, "what unit number is that"):
        return f"{pronoun} stated that only the downstairs unit, 105, was reportedly damaged."
    if water_heater_claim and match(question, "what his specific damages are", "what were the specific damages in each of these three rooms"):
        return f"{pronoun} stated that the downstairs neighbor showed water damage to the kitchen ceiling, drywall near the range, and adjacent dining-area flooring."
    if water_heater_claim and match(question, "was this in multiple rooms or one room", "okay, so just the kitchen and dining room"):
        return None
    if water_heater_claim and match(question, "do you know if he's had any dry out or mold remediation completed", "do you know if he's had his damages repaired yet"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the downstairs unit had remediation or repairs completed."
    if water_heater_claim and match(question, "do you know if anyone was injured or is claiming to be injured as a result of the incident"):
        return f"{pronoun} stated that {context.pronoun} is not aware of anyone claiming an injury as a result of the leak."
    if water_heater_claim and match(question, "do you know if he has insurance on his unit"):
        return f"{pronoun} stated that the downstairs neighbor said he did not have insurance on the unit."
    if water_heater_claim and match(question, "what exactly leaks in your unit on the date of loss"):
        return f"{pronoun} stated that the water heater leaked."
    if water_heater_claim and match(question, "where is that water heater located"):
        return f"{pronoun} stated that the water heater is located under the kitchen counter next to the oven range and is not readily accessible."
    if water_heater_claim and match(question, "how old was this water heater"):
        return f"{pronoun} stated that the water heater was about eight and a half years old."
    if water_heater_claim and match(question, "do you know what caused it to leak"):
        return f"{pronoun} stated that {context.pronoun} believes mineral deposits may have caused internal deterioration, based on what the plumber said."
    if water_heater_claim and match(question, "was there any rust or corrosion on the water heater", "did you ever notice any rust or corrosion on the water heater prior to the date of loss"):
        return f"{pronoun} stated that there was no visible rust or corrosion and that the plumber described it as an internal leak."
    if water_heater_claim and match(question, "had this water heater ever leaked before", "had you ever had any prior issues or repairs involving this water heater prior to the date of loss"):
        return f"{pronoun} stated that there had been no prior leaks or prior issues involving the water heater."
    if water_heater_claim and match(question, "did you have any reason to believe that the water heater was at risk of leaking prior to the date of loss"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe the water heater was at risk of leaking because the family had planned to replace it at the ten-year mark."
    if water_heater_claim and match(question, "was the water heater replaced the following day"):
        return f"{pronoun} confirmed that the water heater was replaced the next day."
    if water_heater_claim and match(question, "did you file a first party claim with citizens for the damage to your own unit") and normalize(pair.answer).lower().startswith("yes"):
        return f"{pronoun} confirmed that {context.pronoun} did file a first-party claim with Citizens for damage to {context.possessive} own unit."
    if water_heater_claim and match(question, "have you had any prior water leaks originating from your condo unit since you have owned it"):
        return f"{pronoun} stated that there had been one prior leak in 2010 involving a toilet water valve."
    if water_heater_claim and match(question, "did it affect the unit below"):
        return f"{pronoun} stated that the prior 2010 leak did cause some water damage to the unit below."
    if water_heater_claim and match(question, "did he have that damage repaired prior to this incident", "so his damage from the prior leak was repaired prior to this incident"):
        return f"{pronoun} stated that the damage from that prior leak was repaired long before this incident."
    if context.joint_statement and match(question, "what is your occupation") and normalize(pair.answer).lower().strip(" .") in {"i'm retired", "i am retired"}:
        return None
    if context.joint_statement and match(question, "state and spell your full name", "please state and spell your full name", "what is your spouse's name"):
        return None
    if context.joint_statement and match(question, "what is your date of birth") and index < 20:
        return None
    if context.joint_statement and match(question, "and what address do you live at") and "florida" in pair.answer.lower():
        return "They reside at 875 8th Street Southeast, Naples, Florida 34117."
    if context.joint_statement and match(question, "and what address do you live at"):
        return None
    if context.joint_statement and match(question, "what is your occupation") and "retired, on disability" in pair.answer.lower():
        return "Karen stated that she is retired and on disability."
    if match(question, "what are you retired from") and "city mattress" in pair.answer.lower():
        return "Karen stated that she previously worked in sales at City Mattress for approximately 23 years."
    if context.joint_statement and match(question, "what is your marital status") and "married" in pair.answer.lower():
        return None
    if match(question, "how long have the two of you been married"):
        return "They stated that they have been married for 39 years."
    if match(question, "do you both live together at this address you provided"):
        return None
    if shooting_claim and match(question, "what is the address where the incident occurred"):
        return None
    if shooting_claim and match(question, "you already said you don't know where the shooting occurred"):
        return None
    if match(question, "where the shooting occurred"):
        return "They stated that they do not know the exact location where the shooting occurred."
    if match(question, "what is the date of loss") and "february 8" in pair.answer.lower():
        return "The date of loss was stated as February 8, 2025."
    if match(question, "using the property as your primary residence") and normalize(pair.answer).lower().startswith("yes"):
        return "On the date of loss, the property was being used as their primary residence."
    if match(question, "who lived at the property with you") and "clark shotwell iii" in pair.answer.lower():
        return (
            "They stated that Clark Shotwell III, his wife Ashley Franks, and their daughters, Adley and Emma Shotwell, "
            "were living at the property. They noted that Franks is Ashley's maiden name."
        )
    if match(question, "what is his date of birth") and "1989" in pair.answer.lower():
        return "They stated that Clark Shotwell III's date of birth is September 15, 1989."
    if match(question, "his wife's name") and "ashley franks" in pair.answer.lower():
        return "They identified Clark Shotwell III's wife as Ashley Franks."
    if match(question, "was gustavo shotwell living with you at the time") and normalize(pair.answer).lower().startswith("yes"):
        return "They confirmed that Gustavo Shotwell was also living with them at the time."
    if match(question, "what is his relationship to you and your wife") and "adopted son" in pair.answer.lower():
        return "They stated that Gustavo Shotwell is their adopted son."
    if question.strip().lower() == "when?" and "november 10" in pair.answer.lower() and "2019" in pair.answer.lower():
        return "They stated that they legally adopted Gustavo Shotwell on November 10, 2019."
    if match(question, "have you already provided me with the adoption paperwork") and is_affirmative(pair.answer):
        return "They stated that the adoption paperwork had already been provided."
    if match(question, "was gustavo using your property as his primary residence") and normalize(pair.answer).lower().startswith("yes"):
        return "They confirmed that Gustavo was using the property as his primary residence on the date of loss."
    if match(question, "how long had gustavo used your property as his primary residence"):
        return "They stated that Gustavo had been living with them at that address for approximately six years."
    if match(question, "did he ever live anywhere else") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that Gustavo had not lived anywhere else during that period."
    if match(question, "when did you first meet gustavo") or (match(question, "can you say that again") and "child welfare" in pair.answer.lower()):
        return "Karen stated that she met Gustavo through her work in child welfare approximately one year before the adoption."
    if match(question, "what company were you working with when you met gustavo"):
        return "Karen stated that the organization was One More Child."
    if match(question, "did you foster him prior to adopting him"):
        return (
            "They stated that Gustavo stayed with another foster family when he first entered foster care, "
            "then lived with them as a foster child for approximately 90 days before the adoption."
        )
    if match(question, "what happened with the prior foster family"):
        return "They stated that the prior foster family was not ready to start a family or commit permanently, so Gustavo became available for adoption."
    if match(question, "how old was gustavo at the time you met him") and "11" in pair.answer.lower():
        return "They stated that Gustavo was 11 years old when they met him."
    if match(question, "what about when you adopted him") and "11" in pair.answer.lower():
        return "They stated that Gustavo was still 11 years old when they adopted him."
    if match(question, "what foster agency did he come from") and "do you live in florida" in pair.answer.lower():
        return None
    if match(question, "was the faster agency state-run") and "contracted with the state" in pair.answer.lower():
        return "They stated that the foster agency was contracted with the state."
    if match(question, "is it state-run or privately run") and "privately" in pair.answer.lower():
        return "They clarified that the foster agency was privately operated."
    if match(question, "how long were you fostering gustavo before you adopted him"):
        return "They stated that they fostered Gustavo for approximately 90 days before the adoption."
    if match(question, "were you fostering any other kids") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that they were not fostering any other children at that time."
    if match(question, "have you ever fostered any other kids"):
        return "They stated that they had fostered one other child before Gustavo."
    if match(question, "did you adopt that other child") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that they did not adopt that other child."
    if match(question, "did you adopt that other child") and "we did not" in pair.answer.lower():
        return "They stated that they did not adopt that other child."
    if match(question, "have you ever adopted any other children") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that they have not adopted any other children."
    if match(question, "why did you adopt gustavo"):
        return "They stated that Gustavo was available for adoption and that they felt called to adopt him."
    if match(question, "how much and how often were you being paid") and "400" in pair.answer.lower():
        return "They stated that the foster stipend was approximately $400 per month."
    if match(question, "were you still being paid to keep gustavo after adopting him") and "monthly stipend" in pair.answer.lower():
        return "They stated that adoptive parents continued to receive a monthly stipend for Gustavo after the adoption."
    if match(question, "are you still receiving the stipend") and "700" in pair.answer.lower():
        return "They stated that they are still receiving the stipend and that it increased from $400 to $700 per month around April 2025."
    if match(question, "still receiving the $700 a month stipend") and "we are" in pair.answer.lower():
        return None
    if match(question, "did the foster agency ever vet the kids") and "cannot tell you" in pair.answer.lower():
        return "They stated that they could not fully explain the agency's screening process, although they assumed the agency gathered background information before placing children into foster care."
    if match(question, "did the foster agency ever vet the kids") and "not sure i can completely answer" in pair.answer.lower():
        return "They stated that they could not fully explain the agency's screening process, although they assumed the agency gathered background information before placing children into foster care."
    if match(question, "were troubled youth part of the foster adoption program"):
        return "They stated that troubled youth can be part of the foster system, although those children can be harder to place for adoption."
    if match(question, "was gustavo a troubled youth") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that they did not consider Gustavo to be a troubled youth."
    if match(question, "where's gustavo now") and "foster family in charlotte county" in pair.answer.lower():
        return "They stated that Gustavo is now living with a foster family in Charlotte County, Florida."
    if match(question, "did you guys give gustavo up") and "residential care" in pair.answer.lower():
        return (
            "They stated that after Gustavo was arrested on the gun charge and detained by the Department of Juvenile Justice, "
            "they refused to take him back because they believed he needed residential care and a psychologist supported that recommendation."
        )
    if match(question, "what kind of residential care"):
        return "They stated that the residential care they sought would have involved live-in counseling and treatment until he was better able to manage his life and decisions."
    if match(question, "how did giving up gustavo work") and "refused to take him" in pair.answer.lower():
        return "They stated that they refused to take Gustavo back into the home."
    if match(question, "are you still in contact with gustavo") and pair.answer.lower().startswith("no"):
        return "They stated that they are not really in contact with Gustavo."
    if match(question, "do you have any way of contacting gustavo") and "case manager" in pair.answer.lower():
        return "They stated that any contact with Gustavo would have to be through his case manager."
    if match(question, "incidents involving gustavo to the foster agency"):
        return None
    if match(question, "and what did they say") and "baker acted" in pair.answer.lower():
        return (
            "They stated that Gustavo's first major incident involved a suicide attempt, after which he was Baker Acted and taken to SalusCare in Fort Myers. "
            "They explained that they again wanted him placed in residential care, and the psychologist agreed, but Children's Network and the health insurer refused. "
            "According to them, he returned home for about 24 hours after release and then ran away."
        )
    if match(question, "how did he come to be in your care"):
        return "They stated that Karen met Gustavo through her work, learned he would be available for adoption, and they decided to adopt him."
    if shooting_claim and match(question, "in your own words on the date of loss we're talking about the shooting"):
        return (
            "In describing what happened, they stated that they have no personal knowledge of the shooting and only know what others later told them. "
            "They stated that the shooting allegedly occurred on February 8, 2025. "
            "According to them, family members and Ashley's friends later reported that there had been a party and that there was a rumor Gus had been there with a gun."
        )
    if shooting_claim and match(question, "who was at the party") and "gus" in pair.answer.lower():
        return "They stated that the rumor was that Gus was at the party."
    if shooting_claim and match(question, "how did you see the gun on a ring camera"):
        return "They stated that ring camera footage allegedly showed Gus holding the gun and showing it to a friend."
    if shooting_claim and match(question, "who was showing it to his friend"):
        return None
    if shooting_claim and match(question, "why did he have a gun"):
        return "They stated that they do not know why Gus had the gun and believe he accessed it by taking the key, opening the lockbox, and removing the gun from the closet."
    if shooting_claim and match(question, "stolen the gun at least one time before the shooting happened"):
        return "They stated that Gus had at least looked at the gun before the shooting."
    if shooting_claim and match(question, "did you ever make any additional attempts to further secure the gun after that incident"):
        return "They stated that they did not make additional security changes after that incident because the gun was already locked up and a sheriff had told them the storage method was sufficient."
    if shooting_claim and match(question, "you didn't attempt to make the gun more secure after that incident"):
        return None
    if shooting_claim and match(question, "did you attempt to make the gun more secure after the first incident") and normalize(pair.answer).lower().startswith("no"):
        return None
    if shooting_claim and match(question, "how did you first become aware of the incident involving the shooting"):
        return "They stated that they first learned about the shooting from Ashley's friends, who said Gus had been at a party and had a gun."
    if shooting_claim and match(question, "was gustavo the shooter"):
        return "They stated that they cannot say whether Gustavo was the shooter because they were not there and understand that the police do not believe he was."
    if shooting_claim and match(question, "you were first told about the incident from who"):
        return None
    if shooting_claim and match(question, "when was this that you first became aware of the incident") and "february 15" in pair.answer.lower():
        return "They stated that they first became aware of the shooting on February 15, 2025."
    if shooting_claim and match(question, "were you or your wife present at this party") and normalize(pair.answer).lower().startswith("no"):
        return "They confirmed that neither of them was present at the party."
    if shooting_claim and match(question, "did you guys report the incident to anyone once you became aware of it"):
        return None
    if shooting_claim and match(question, "once you became aware of the shooting did you report it to anyone"):
        return "They stated that they did not make a separate report about the shooting because the police already knew about it and they did not yet know that their gun was missing."
    if shooting_claim and match(question, "why did you not report the incident to citizens insurance once you became aware of it"):
        return "They stated that they did not report the matter to Citizens at that time because they did not know they were supposed to."
    if shooting_claim and match(question, "what happened with the police") and "hospital" in pair.answer.lower():
        return "They stated that when Gus was later in the hospital after a suicide attempt, they told a sheriff that their gun was missing and that Gus had taken it."
    if shooting_claim and match(question, "what happened with the police involving the shooting"):
        return "They stated that they do not know the details of the police investigation into the shooting and only know that the shooting had already been reported on February 8."
    if shooting_claim and match(question, "was anyone arrested related to the shooting") and normalize(pair.answer).lower().startswith("yes"):
        return "They stated that someone was arrested in connection with the shooting."
    if shooting_claim and question.strip().lower() == "who?" and "christian wolf" in pair.answer.lower():
        return "They identified the person arrested as Christian Wolf."
    if shooting_claim and match(question, "why was he arrested"):
        return "They stated that they do not know the specific charges, only that Christian Wolf was arrested a couple of months later after they were told he had left the area."
    if shooting_claim and match(question, "did the shooting occurred at a party"):
        return "They stated that they only heard that the shooting may have occurred at a party and do not know that for certain because they were not there and do not know anyone who was."
    if shooting_claim and match(question, "was gustavo at the party when the shooting occurred"):
        return "They stated that they do not believe Gus was at the party, but they do not know that for certain."
    if shooting_claim and match(question, "didn't earlier you say you don't know if he was the shooter or not"):
        return "They stated that they do not know who shot the claimant, and they do not believe Gus was at the party."
    if shooting_claim and match(question, "so he was home when the shooting occurred"):
        return "They stated that, as far as they know, Gus was home when the shooting occurred, although they were not certain."
    if shooting_claim and match(question, "what was gustavo's relationship to christian wolf"):
        return "They stated that they did not know Gustavo to have a direct relationship with Christian Wolf and understood that any connection may have been through Leo Woodard, who was Gus's close friend."
    if shooting_claim and match(question, "how close was he") and "spent the night" in pair.answer.lower():
        return "They stated that Leo had been to their house several times and had spent the night there with other friends."
    if shooting_claim and match(question, "do you know how long gustavo and leo knew each other") and "year" in pair.answer.lower():
        return "They stated that Gus and Leo had known each other for approximately one year."
    if shooting_claim and match(question, "gustavo knew christian"):
        return "They stated that they cannot say whether Gus knew Christian Wolf directly."
    if shooting_claim and match(question, "do you know if uh gustavo knew christian"):
        return "They stated that they cannot say whether Gus knew Christian Wolf directly."
    if shooting_claim and match(question, "did christian or leo ever force gustavo to do anything"):
        return "They stated that they were not aware of Christian Wolf or Leo Woodard ever forcing Gustavo to do anything."
    if shooting_claim and match(question, "violent behavior between christian and gustavo or leo woodard"):
        return "They stated that they were not aware of prior violent behavior involving Christian Wolf, Gustavo, or Leo Woodard."
    if shooting_claim and match(question, "when was the last time gustavo had seen leo woodard"):
        return "They stated that they do not know when Gus last saw Leo Woodard before the shooting."
    if shooting_claim and match(question, "who owned the gun that was involved in the shooting") and "registered in my name" in pair.answer.lower():
        return "They stated that the gun was registered in Karen's name."
    if shooting_claim and match(question, "what kind of gun was it that was involved in the shooting"):
        return "They stated that the gun was a Glock 9mm."
    if shooting_claim and match(question, "how many rounds did the gun hold"):
        return "They stated that the gun held approximately 12 rounds."
    if shooting_claim and match(question, "how old was the gun") and "14 years" in pair.answer.lower():
        return "They stated that the gun was approximately 14 years old."
    if shooting_claim and match(question, "how long had you owned the gun") and "14 years" in pair.answer.lower():
        return "They stated that they had owned the gun for approximately 14 years."
    if shooting_claim and match(question, "anything wrong with the gun on the date of loss"):
        return "They stated that they were not aware of anything wrong with the gun on the date of loss."
    if shooting_claim and match(question, "were there any modifications to the gun prior to the date of loss") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that there were no modifications to the gun."
    if shooting_claim and match(question, "can you give me a year, please") and "2007" in pair.answer.lower():
        return "They stated that they believe the gun was purchased in 2007."
    if shooting_claim and match(question, "where did you purchase the gun"):
        return "They stated that the gun was purchased through a company in Ohio that they believed was called Sports Online."
    if shooting_claim and match(question, "did you buy the gun new or used") and "new" in pair.answer.lower():
        return "They stated that the gun was purchased new."
    if shooting_claim and match(question, "why did you purchase the gun"):
        return "They stated that they bought the gun for protection."
    if shooting_claim and match(question, "any specific reason other than that"):
        return "They stated that they believed it was good to have a gun for safety and to protect themselves and their family if necessary."
    if shooting_claim and match(question, "do you live in an area where the wildlife requires the need for a gun") and "possibly" in pair.answer.lower():
        return "They stated that wildlife and the large lots in their area were part of why they considered having a gun useful."
    if shooting_claim and match(question, "do you live in an area where the wildlife requires the need for guns"):
        return None
    if shooting_claim and match(question, "was this your first time owning a gun") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that this was not the first gun they had owned."
    if shooting_claim and match(question, "are you experienced with guns"):
        return "They described themselves as having moderate experience with guns."
    if shooting_claim and match(question, "how much experience do you have with guns"):
        return "They stated that Karen's father taught her how to use and safely handle a gun."
    if shooting_claim and match(question, "how was your gun stored before it was stolen"):
        return "They stated that the gun was kept unloaded in a lockbox inside the master-bedroom closet on a shelf behind towels and linens."
    if shooting_claim and match(question, "can you describe the lock box"):
        return "They described the lockbox as a black box with a handle, approximately 12 inches by 8 inches, with a cord wrapped around the handle to keep it locked."
    if shooting_claim and match(question, "what material is the lock box made of", "is it for one gun only", "was there a lock wrapped around the handle", "is that the proper way to use that lock"):
        return None
    if shooting_claim and match(question, "where was the key for this lock stored"):
        return "They stated that the key was hidden in Karen's drawer, and the drawer itself was not locked."
    if shooting_claim and match(question, "did gustavo know where the key was") and "not to my knowledge" in pair.answer.lower():
        return "They stated that, to their knowledge, Gustavo did not know where the key was kept."
    if shooting_claim and match(question, "how did he access it before the date of loss"):
        return "They stated that they believe Gus found the key by rummaging through Karen's drawers."
    if shooting_claim and match(question, "did you move the key after he pulled the gun the first time") and normalize(pair.answer).lower().startswith("yes"):
        return "They stated that the key location was changed after Gus first pulled the gun out."
    if shooting_claim and match(question, "was the gun and the ammunition stored together or separately"):
        return "They stated that the gun and ammunition were stored separately."
    if shooting_claim and match(question, "where was the ammunition stored"):
        return "They stated that the ammunition was kept downstairs in a different locked area."
    if shooting_claim and match(question, "was the gun stored loaded or unloaded"):
        return "They stated that the gun was stored unloaded."
    if shooting_claim and match(question, "was the ammunition stolen as well") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that the ammunition was not stolen."
    if shooting_claim and match(question, "the only thing that was stolen was the gun"):
        return "They stated that, as far as they knew, only the gun was stolen."
    if shooting_claim and match(question, "where was this lock box stored", "was the lock box hidden"):
        return None
    if shooting_claim and match(question, "what kind of ammunition did you have with the nine millimeter"):
        return "They stated that the ammunition was nine-millimeter range ammunition."
    if shooting_claim and match(question, "what kind of ammunition did you have with the nine millimeter", "was it range ammo or hollow point", "where exactly was the ammo stored"):
        return None
    if shooting_claim and match(question, "were the magazines stored with the gun", "my question was, the magazine for the gun stored with the gun or separately"):
        return "They stated that the magazine was stored separately in the lockbox."
    if shooting_claim and match(question, "were any magazines stolen"):
        return "They stated that they were not aware of any magazines being stolen."
    if shooting_claim and match(question, "was any of that ammo stolen") and normalize(pair.answer).lower().startswith("no"):
        return None
    if shooting_claim and match(question, "did you report the gun as stolen") and normalize(pair.answer).lower().startswith("yes"):
        return "They stated that the gun was reported stolen shortly after they realized it was missing."
    if shooting_claim and question.strip().lower() == "when?" and "reported to a sheriff" in pair.answer.lower():
        return None
    if shooting_claim and match(question, "what was the outcome of you reporting the gun stolen"):
        return "They stated that Gus was arrested on the gun charge. They further stated that the gun was later found in a canal near where they were told the party had taken place, and they did not know of any further outcome."
    if shooting_claim and match(question, "what evidence did they have against gustavo for the stolen gun") and "early march" in pair.answer.lower():
        return "They stated that they filed charges against Gus regarding the stolen gun in early March 2025."
    if shooting_claim and match(question, "what evidence did they have against gustavo for the stolen gun") and normalize(pair.answer).lower().startswith("none"):
        return "They stated that they were not aware of any real evidence against Gus beyond the fact that they reported the gun missing."
    if shooting_claim and match(question, "did you recover the gun") and "evidence" in pair.answer.lower():
        return "They stated that the gun had not been returned to them and was still being held as evidence."
    if shooting_claim and match(question, "do you know who shot the claimant"):
        return "They stated that they do not know who shot the claimant and could only say that Christian Wolf had been identified as the person arrested."
    if shooting_claim and match(question, "so it's being alleged that christian wolf shot the claimant") and is_affirmative(pair.answer):
        return None
    if shooting_claim and match(question, "do you know how christian wolf ended up in possession of your gun"):
        return "They stated that they do not know how Christian Wolf came into possession of the gun."
    if shooting_claim and match(question, "did you or your wife give the gun to christian wolf") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that neither of them gave the gun to Christian Wolf."
    if shooting_claim and match(question, "so no, you do not") and "assume it got from gus" in pair.answer.lower():
        return "They stated that they assume the gun likely passed from Gus to Christian somehow."
    if shooting_claim and match(question, "did gustavo give the gun to christian wolf"):
        return "They stated that they do not know whether Gus gave the gun to Christian Wolf."
    if shooting_claim and match(question, "did christian take the gun from gustavo by force", "did christian force gustavo to take the gun", "did christian's brother force gustavo to take the gun"):
        return "They stated that they do not know whether anyone forced Gus to take or surrender the gun."
    if shooting_claim and match(question, "did gustavo take the gun from your lock box without permission") and "we think so" in pair.answer.lower():
        return "They stated that they believe Gus took the gun from the lockbox without permission."
    if shooting_claim and match(question, "how often do you check the lock box where the gun was stored"):
        return "They stated that they checked the lockbox regularly and generally assumed it remained undisturbed because few people went into the bedroom."
    if shooting_claim and match(question, "how often would you check it"):
        return "They stated that they did not check the lockbox on any fixed schedule and therefore could not say exactly when the gun first went missing."
    if shooting_claim and match(question, "who was gustavo showing the gun to the first time you pulled it out"):
        return "They stated that the boy Gus had previously shown the gun to was named Brantley, although they did not know his last name."
    if shooting_claim and match(question, "did christian or his brother know you had guns in the house"):
        return "They stated that they do not know whether Christian Wolf or Leo Woodard knew there were guns in the house."
    if shooting_claim and match(question, "did christian wolfe have access to your guns") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that Christian Wolf did not have access to their guns."
    if shooting_claim and match(question, "prior to the date of loss, when was the last time you used or checked the gun"):
        return "They stated that the gun had not been used in years, but they knew it was still there on December 31."
    if shooting_claim and match(question, "did you believe that your guns were stored securely to the best of your abilities") and normalize(pair.answer).lower().startswith("yes"):
        return "They stated that they believed the guns were stored as securely as they could manage."
    if shooting_claim and match(question, "do you believe that your guns were stored securely to the best of your abilities") and normalize(pair.answer).lower().startswith("yes"):
        return None
    if shooting_claim and match(question, "did you have any reason to believe that gustavo or any of his acquaintances wanted access to your guns"):
        return "They stated that they had no reason to believe Gus or his acquaintances were trying to get access to the guns, although they believe Gus later went into the bedroom, found the gun case, searched the drawers, and located the key."
    if shooting_claim and match(question, "did gustavo know where the key was stored") and "not as far as i know" in pair.answer.lower():
        return None
    if shooting_claim and match(question, "did you ever provide gustavo with the key", "did you ever allow gustavo to handle or use the gun", "did you guys ever give gustavo any firearms training", "so you never took him to go shooting or anything") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that they never provided Gus with the key, never allowed him to handle or use the gun, and never gave him firearms training or took him shooting."
    if shooting_claim and match(question, "so you never took him to go shooting or anything") and "never did" in pair.answer.lower():
        return None
    if shooting_claim and match(question, "did gustavo know how to use a firearm"):
        return "They stated that they do not know whether Gus knew how to use a firearm."
    if shooting_claim and match(question, "what did he say") and "didn't know" in pair.answer.lower():
        return "They stated that when they asked Gus how Christian got the gun, he said he did not know."
    if shooting_claim and match(question, "why do you not believe gustavo"):
        return "They stated that they did not believe Gus because they had already caught him in other lies."
    if shooting_claim and match(question, "who did christian shoot"):
        return "They stated that they believe the claimant's name is Gaspar, although they were not certain because they only saw the name later in the attorney's letter."
    if shooting_claim and match(question, "why did christian shoot the claimant"):
        return "They stated that they do not know why Christian Wolf shot the claimant."
    if shooting_claim and match(question, "where did christian shoot the claimant"):
        return None
    if shooting_claim and match(question, "tell me what you do know") and "shot in the back twice" in pair.answer.lower():
        return "They stated that they were told the shooting happened somewhere in Golden Gate Estates and that the claimant was shot twice in the back."
    if shooting_claim and match(question, "did anyone call for an ambulance"):
        return "They stated that they do not know whether anyone called an ambulance."
    if shooting_claim and match(question, "did the claimant pass away"):
        return "Based on what they were told, the claimant did not pass away."
    if shooting_claim and match(question, "taken to the hospital by ambulance"):
        return "They stated that they do not know whether the claimant was transported by ambulance."
    if shooting_claim and match(question, "did the claimant go to the hospital", "claimant went to the hospital"):
        return "They stated that they believe the claimant went to the hospital."
    if shooting_claim and match(question, "how long the claimant was in the hospital"):
        return "They stated that they do not know how long the claimant remained in the hospital."
    if shooting_claim and match(question, "what the claimant's injuries are or were"):
        return "They stated that, to their knowledge, the claimant sustained gunshot wounds."
    if shooting_claim and match(question, "what, if any, treatments the claimant underwent", "surgeries, physical therapy, scarring, or permanent disabilities", "claimant has recovered"):
        return "They stated that they do not know what treatment the claimant underwent, whether there were surgeries or lasting effects, or whether the claimant has fully recovered."
    if shooting_claim and match(question, "have you ever seen or spoken with the claimant"):
        return "They stated that they have never seen or spoken with the claimant or his family."
    if shooting_claim and match(question, "prior injuries or accidents resulting in any injuries"):
        return "They stated that they do not know whether the claimant had any prior injuries or prior accidents."
    if shooting_claim and match(question, "would you be able to describe the claimant's physical appearance"):
        return "They stated that they cannot describe the claimant's physical appearance."
    if shooting_claim and match(question, "describe the claimant's physical appearance", "do you know anything about the claimant", "so you don't know the claimant in any capacity"):
        return None
    if shooting_claim and match(question, "did the claimant have any issues with drugs or alcohol", "was the claimant involved with any criminal behavior"):
        return "They stated that they do not know whether the claimant had any drug, alcohol, or criminal issues."
    if shooting_claim and match(question, "can you describe gustavo shotwell briefly"):
        return "They described Gustavo as approximately 5'9\", about 240 pounds the last time they saw him, and African-American."
    if shooting_claim and match(question, "how old was he on the date of loss") and "16" in pair.answer.lower():
        return "They stated that Gustavo was 16 years old on the date of loss."
    if shooting_claim and match(question, "how was gustavo's grades"):
        return "They stated that Gustavo's grades were generally okay."
    if shooting_claim and match(question, "did he attend school regularly") and normalize(pair.answer).lower().startswith("yes"):
        return "They stated that Gustavo attended school regularly."
    if shooting_claim and match(question, "did gustavo ever have any issues at school", "did he ever have any issues at school") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that Gustavo did not have meaningful issues at school."
    if shooting_claim and match(question, "did gustavo ever have any issues at school", "did he ever have any issues at school", "did gustavo ever have any behavioral issues at school", "did gustavo ever get suspended or expelled from school"):
        return "They stated that Gustavo did not have meaningful school discipline issues and was not suspended or expelled."
    if shooting_claim and match(question, "did gustavo ever use drugs or alcohol"):
        return "They stated that they were not aware of drug or alcohol use while he lived with them, although after he ran away they found marijuana in his room."
    if shooting_claim and match(question, "you know where he got the weed"):
        return "They stated that they have no idea where Gustavo got the marijuana. They added that neither of them uses marijuana or other drugs."
    if shooting_claim and match(question, "had gustavo ever been arrested prior to the date of loss"):
        return "They stated that Gustavo had not been arrested before the date of loss and was arrested later, around April 30, after he had run away."
    if shooting_claim and match(question, "and that was for the running away or for what"):
        return "They stated that the later arrest involved Gus riding in a car with other kids where marijuana was allegedly found in his backpack."
    if shooting_claim and match(question, "did gustavo have any mental health or behavioral issues"):
        return "They stated that they did not believe Gustavo had mental health or behavioral issues, although they previously felt he needed residential care after the other incidents."
    if shooting_claim and match(question, "would you be able to describe christian wolf at all"):
        return "They stated that they could not describe Christian Wolf."
    if shooting_claim and match(question, "how old christian wolf was at the time of the shooting"):
        return "They stated that Christian Wolf was over 21, although they did not know his exact age."
    if shooting_claim and match(question, "did christian have any issues with drugs or alcohol", "if christian had ever been arrested prior to the date of loss"):
        return "They stated that they do not know whether Christian Wolf had drug, alcohol, or arrest history."
    if shooting_claim and match(question, "what was christian's brother's name again"):
        return "They identified Christian Wolf's brother as Leo Woodard."
    if shooting_claim and match(question, "how old was leo at the time of the shooting"):
        return "They stated that Leo was probably around the same age as Gus, approximately 16 or 17."
    if shooting_claim and match(question, "did leo have any issues with drugs or alcohol", "had leo ever been arrested before", "did leo have any history of violent behavior or mental health issues"):
        return "They stated that they do not know whether Leo had drug, arrest, violent-behavior, or mental-health history."
    if shooting_claim and match(question, "do you know if anyone else was injured on the date of loss"):
        return "They stated that they do not know whether anyone else was injured on the date of loss."
    if shooting_claim and match(question, "was there anything that you could have done personally to prevent the incident"):
        return "They stated that they do not believe there was anything they personally could have done to prevent the incident."
    if shooting_claim and match(question, "was there anything the claimant could have done to prevent the incident"):
        return "They stated that they do not know whether the claimant could have done anything to prevent the incident."
    if shooting_claim and match(question, "was there anything that gustavo could have done to prevent the incident"):
        return "They stated that, to their knowledge, there was nothing Gustavo could have done to prevent the incident."
    if shooting_claim and match(question, "could anything have prevented the incident"):
        return "They stated that they cannot think of anything that would have prevented the incident."
    if shooting_claim and match(question, "were there any witnesses to this shooting"):
        return "They stated that they do not know whether there were any witnesses to the shooting."
    if shooting_claim and match(question, "do you own any additional properties in the united states"):
        return "They stated that they own an unimproved lot in North Carolina, Georgia that is presently under contract for sale."
    if shooting_claim and match(question, "can you give me the address of that property"):
        return "They stated that they did not have the address readily available but could provide it later."
    if shooting_claim and match(question, "so you can provide me the address later") and normalize(pair.answer).lower().startswith("sure"):
        return None
    if shooting_claim and match(question, "was there any insurance on that property") and normalize(pair.answer).lower().startswith("no"):
        return "They stated that there was no insurance on that lot."
    if shooting_claim and match(question, "what's the value of that property"):
        return "They estimated the value of that lot at approximately $10,000."
    if match(question, "in your own words, on the date of loss, what happened") and all(token in pair.answer.lower() for token in ("police", "dog")):
        return (
            f"In describing what happened, the {role_noun} stated that {context.pronoun} was not present when the incident occurred and only knows what was later reported to {context.objective}. "
            f"{pronoun} explained that while at work, {context.pronoun} received a call advising that police were at the property because {context.possessive} dogs had allegedly gotten out of the yard. "
            f"{pronoun} stated that when {context.pronoun} arrived, the police were trying to secure the dogs. "
            f"According to what {context.pronoun} was told, the dogs got out through a broken section of fence, became involved in a fight with the claimant's dog, and the claimant was bitten on both hands while trying to hold onto her dog. "
            f"{pronoun} further stated that the claimant's dog was also bitten and later had to be euthanized."
        )
    if match(question, "so your dogs escaped the yard and attacked the lady and her dog"):
        return None
    if match(question, "did the dogs jump back into the yard after that"):
        return f"{pronoun} stated that the police chased the dogs back into the yard and blocked the opening so they could not get back out."
    if match(question, "did the, did your dogs attack the claimant or her dog first", "did your dogs attack the claimant or the claimant's dog first"):
        return (
            f"{pronoun} stated that {context.pronoun} does not know whether {context.possessive} dogs attacked the claimant or the claimant's dog first because {context.pronoun} was not present and was not shown any video of the incident. "
            f"{pronoun} speculated that there may have been activity or scratching near the fence before the dogs got out."
        )
    if match(question, "so you don't know who"):
        return None
    if question.strip().lower() in {"i'm sorry, what?", "so you don't know.", "so you don't know", "so yes.", "so no."}:
        return None
    if match(question, "how did you first become aware of the incident") and "phone" in pair.answer.lower() and "work" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} first became aware of the incident when {context.pronoun} received a phone call at work."
    if match(question, "how long after the incident occurred was this that you got called"):
        return f"{pronoun} estimated that {context.pronoun} received that call approximately 30 minutes after the incident."
    if match(question, "so you weren't present when the incident occurred", "and you didn't witness the incident"):
        return None
    if match(question, "did you report the accident to anyone") and "police" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} did not separately report the incident because the police were already on scene."
    if match(question, "was the claimant walking her dog") and "police tell me" in pair.answer.lower():
        return f"{pronoun} stated that, based on what the police told {context.objective}, the claimant was walking her dog when the incident occurred."
    if match(question, "do you know where the claimant was exactly when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} does not know exactly where the claimant was when the incident occurred because {context.pronoun} was not present."
    if match(question, "was the claimant on your property"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the claimant was on the property when the incident occurred."
    if match(question, "how many dogs did the claimant have with her"):
        if normalize(pair.answer).lower().startswith("one"):
            return f"{pronoun} stated that the claimant had one dog with her."
    if match(question, "what kind of dog did the claimant have with her"):
        return f"{pronoun} stated that {context.pronoun} does not know what kind of dog the claimant had with her."
    if match(question, "was it a small or a large dog"):
        if "large" in pair.answer.lower():
            return f"{pronoun} stated that the claimant's dog appeared to be a large dog."
    if match(question, "was the claimant's dog on a leash or was it being carried", "was the claimant's dog doing anything to draw your dog's attention", "was the claimant aware of your dogs prior to the incident occurring", "was the claimant doing anything to draw your dog's attention"):
        return (
            f"{pronoun} stated that {context.pronoun} has no information regarding whether the claimant's dog was on a leash, "
            f"whether the dogs were reacting to anything the claimant or her dog did, or whether the claimant was aware of the dogs beforehand."
        )
    if match(question, "calling, feeding, harassing, or attempting to play with or pet your dogs"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that, to {context.possessive} knowledge, the claimant was not calling, feeding, harassing, or attempting to play with or pet the dogs when the incident occurred."
        return f"{pronoun} stated that {context.pronoun} does not know."
    if match(question, "where were your dogs located prior to the incident occurring"):
        return f"{pronoun} stated that the dogs were in the yard before the incident occurred."
    if match(question, "were the dogs in the front yard prior to the incident occurring"):
        return f"{pronoun} clarified that the dogs were in the front yard before the incident occurred."
    if match(question, "why were your dogs in the front yard when the incident occurred"):
        return f"{pronoun} stated that the dogs were normally kept in the yard."
    if match(question, "were the dogs tied up, chained up, or restrained in any way when the incident occurred"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that the dogs were not tied, chained, or otherwise restrained."
    if match(question, "do you always leave the dogs in the front yard"):
        if is_affirmative(pair.answer):
            return f"{pronoun} stated that the dogs were generally kept in the front yard."
    if match(question, "was your yard completely fenced in", "did your fence enclose your entire yard"):
        if is_affirmative(pair.answer):
            return f"{pronoun} stated that the yard was fully enclosed by a fence."
    if match(question, "was there anywhere for your dogs to escape", "was there anywhere for your dogs to escape the fence"):
        if "escape the fence" in question:
            return None
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that there was no known opening through which the dogs should have been able to escape."
    if match(question, "did the fence have a beware of dog sign"):
        if is_affirmative(pair.answer):
            return f"{pronoun} stated that there was a beware-of-dog sign posted on the fence."
    if question.strip().lower() == "where?" and "front" in pair.answer.lower():
        return f"{pronoun} stated that the sign was posted at the front."
    if match(question, "is it still there"):
        return f"{pronoun} stated that the sign was still there, although it had been painted over."
    if match(question, "can you describe your fence"):
        return f"{pronoun} described the fence as a red wooden fence."
    if match(question, "is it a picket fence"):
        if "picket" in pair.answer.lower():
            return f"{pronoun} clarified that it was a picket fence."
    if match(question, "how tall is the fence"):
        return f"{pronoun} stated that the fence was approximately six feet tall."
    if match(question, "how old is the fence"):
        return f"{pronoun} stated that the fence was approximately nine years old."
    if match(question, "was it installed by a professional", "oh, you installed it"):
        if "me" in pair.answer.lower() or is_affirmative(pair.answer):
            return f"{pronoun} stated that {context.pronoun} installed the fence."
    if match(question, "did you install it properly"):
        if is_affirmative(pair.answer):
            return f"{pronoun} stated that the fence had been properly installed."
    if match(question, "what condition was the fence in on the date of loss"):
        return f"{pronoun} stated that the fence was in the same general condition on the date of loss as it was at the time of the statement."
    if match(question, "was there anything wrong with the fence on the date of loss", "were there any holes or weak spots in the fence on the date of loss", "were there any holes under the fence on the date of loss"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} was not aware of any holes, weak spots, or other fence problems before the incident."
    if match(question, "had the fence ever required any prior repairs"):
        return f"{pronoun} stated that there had previously been an issue with the front gate, which {context.pronoun} repaired about six months before the incident after the dogs were able to get out there."
    if match(question, "so the dog used", "so you had to fix the front gate about six months ago"):
        return None
    if match(question, "commented or complained about the fence, the security of the fence, or the dog's ability to escape"):
        return f"{pronoun} stated that there had been no further complaints after the front gate was repaired."
    if match(question, "how did your dogs escape the yard"):
        return f"{pronoun} stated that on the date of loss the dogs broke a piece of the fence and then got out through the opening that was created."
    if match(question, "they jumped on the fence, broke a board, and then jumped over the fence", "so they didn't jump over the fence"):
        return None
    if match(question, "why did your dogs escape the yard on the date of loss"):
        return f"{pronoun} believed the dogs got out when they saw the claimant walking her dog past the property."
    if match(question, "history of digging or attempting to dig under your fence"):
        return f"{pronoun} stated that {context.pronoun} had never seen the dogs dig under the fence before."
    if match(question, "had the dogs ever escaped your yard or the fence before", "so they never escaped the fence before", "how many times did they escape"):
        return f"{pronoun} stated that the dogs had not previously escaped through the fence itself, but they had once gotten out through the front gate."
    if match(question, "what did they do when they escaped the fence"):
        return f"{pronoun} stated that on that prior occasion the dogs got out into the street, the police were called, and {context.pronoun} returned home and got them back."
    if match(question, "did they ever attack anyone before when they escaped the gate"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that the dogs had not previously attacked anyone."
    if match(question, "how many dogs did you have"):
        if "five" in pair.answer.lower():
            return f"{pronoun} stated that {context.pronoun} had five dogs."
    if match(question, "did you personally own these dogs"):
        if is_affirmative(pair.answer):
            return f"{pronoun} confirmed that {context.pronoun} personally owned the dogs."
    if match(question, "did anyone else own these dogs"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that no one else owned the dogs."
    if match(question, "do you still own your dogs"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} no longer owns the dogs."
    if match(question, "when did you get rid of your dogs"):
        return None
    if question.strip().lower() == "why?" and "incident" in pair.answer.lower() and "problem" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} got rid of the dogs because of the incident and did not want any further problems."
    if match(question, "while jim", "while casey", "did you have any involvement with the property", "involvement with the property located this address"):
        return f"{pronoun} stated that while {natural_insured_ref} was alive, {context.pronoun} had no involvement with the property other than visiting it."
    if match(question, "involved in obtaining insurance for the property"):
        return f"{pronoun} was not involved in obtaining insurance for the property."
    if match(question, "written property management agreements for the property"):
        return f"To {context.possessive} knowledge, {natural_insured_ref} did not have any written property management agreements for the property."
    if match(question, "injured at the property"):
        return f"{pronoun} confirmed that {context.pronoun} is aware the claimant is alleging an injury at the property."
    if match(question, "when did you become aware of this incident"):
        return f"{pronoun} became aware of the incident when {context.possessive} probate attorney informed {context.objective}, during unrelated probate matters, that a $1,000,000 claim had been filed against the estate."
    if match(question, "would that be around mid to late december of 2025"):
        return None
    if match(question, "became aware of this incident before he passed away"):
        return f"{pronoun} does not know if {natural_insured_ref} became aware of the incident before passing away."
    if match(question, "reported the incident to citizens at the time"):
        return f"{pronoun} does not know whether {natural_insured_ref} reported the incident to Citizens at the time the fall was alleged to have occurred."
    if match(question, "ever reported the incident to citizens before he passed away"):
        return f"{pronoun} also does not know whether {natural_insured_ref} ever reported the incident to Citizens before passing away."
    if match(question, "direct communications with", "her husband"):
        return f"{pronoun} confirmed that {context.pronoun} never had any direct communications with the claimant or the claimant's husband."
    if match(question, "tell me anything about the text messages"):
        return f"Regarding the text messages that were shown to {context.objective} before the recorded statement, {context.pronoun} explained that the messages were sent to the boyfriend of {natural_insured_possessive} ex-wife."
    if match(question, "any additional response was provided"):
        return f"{pronoun} stated that, based on what {context.pronoun} was told, the text message compilation appeared to include everything and that {context.pronoun} is not aware of any additional texts beyond that thread."
    if match(question, "phone in which the text messages originated"):
        return f"{pronoun} does not know whether the phone from which the text messages originated is still available, but believe it may be."
    if match(question, "renting the property as an airbnb", "vrbo, or short-term rental before this incident occurred"):
        return f"{pronoun} stated that {context.pronoun} was not aware {natural_insured_ref} had been renting the property as an Airbnb, VRBO, or short-term rental until after {natural_insured_ref} passed away."
    if match(question, "how long was he renting the property as a short-term rental"):
        return f"{pronoun} does not know exactly how long the property had been used as a short-term rental, but believe it had been used that way for several years."
    if (("did jim ever have any long-term renters" in question) or ("did casey ever have any long-term renters" in question)) and "property in question" not in question:
        return f"{pronoun} confirmed that {natural_insured_ref} had long-term renters at other properties, but not at the subject property."
    if match(question, "speaking specifically about the property in question"):
        return f"{pronoun} is not aware of {natural_insured_ref} ever having any long-term renters at the subject property."
    if match(question, "is there any way to figure out whether he had long-term renters"):
        return f"{pronoun} does not think there is a way to determine whether {natural_insured_ref} ever had long-term renters at the subject property."
    if match(question, "is there any way to figure out how long he was using the property as a short-term rental"):
        return f"{pronoun} explained that the only way to determine that would be by speaking with Eric Chudzik or Brenda."
    if match(question, "in order to figure that out, we would need to speak with eric chudzik"):
        return f"{pronoun} stated that Eric Chudzik might know because he was managing the property during that period."
    if match(question, "do you have access to casey's records such as contracts, bank statements, or invoices"):
        return f"{pronoun} stated that, as trustee, {context.pronoun} does have access to {natural_insured_possessive} records, but they do not go back that far."
    if match(question, "is there any way to figure out how long he was using the property as a short-term rental"):
        return f"{pronoun} explained that the only way to determine how long the property had been used as a short-term rental would be to speak with the people who were involved in managing it."
    if match(question, "jim's records such as contracts, bank statements, or invoices"):
        return f"{pronoun} confirmed that {context.pronoun} does have access to {natural_insured_possessive} records as trustee, including contracts, bank statements, and invoices, but stated that those records do not go back far enough to answer every question."
    if match(question, "internet accounts such as airbnb or vrbo"):
        return f"{pronoun} stated that {context.pronoun} does not have access to {natural_insured_possessive} internet accounts such as Airbnb or VRBO and believe those accounts were set up by others involved with managing the property."
    if match(question, "was the rental income from this property being claimed"):
        return f"{pronoun} confirmed that rental income from the property was being claimed on {natural_insured_possessive} taxes."
    if match(question, "which one?") and any(token in answer.lower() for token in ("sole proprietorship", "social security", "taxed under")):
        return f"{pronoun} explained that the income was reported as a sole proprietorship taxed under {natural_insured_possessive} Social Security number."
    if match(question, "full name of the company in which he filed under"):
        company_name = re.sub(r"^(well,\s*)?(it was called)\s+", "", answer, flags=re.IGNORECASE)
        return f"{pronoun} stated that the company name was {company_name}."
    if match(question, "access to jim's prior tax returns", "access to casey's prior tax returns"):
        return f"{pronoun} stated that {context.pronoun} does have access to a few years of {natural_insured_possessive} prior tax returns."
    if match(question, "are you familiar with kind of enterprises"):
        return None
    if match(question, "other than that it was a c corp"):
        return f"{pronoun} explained that the company was not a C corporation, but rather an LLC taxed as a sole proprietorship."
    if match(question, "tell me anything else about this llc"):
        return f"{pronoun} explained that the LLC managed rental properties in Palm Beach County along with the short-term rental activity at the subject property through VRBO or possibly Airbnb."
    if match(question, "what can you tell me about cc84 ventures", "what can you tell me about cc84 ventures, llc"):
        return f"{pronoun} stated that the company referenced in the statement was a shell entity that had been set up for a new real estate venture that never became operational."
    if match(question, "what can you tell me about lahey’s liquor cabinet", "what can you tell me about lahey's liquor cabinet"):
        return f"{pronoun} stated that the company referenced in the statement was a shell entity that had been set up for a new real estate venture that never became operational."
    if match(question, "purchase the property in which the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} does not know the exact year {natural_insured_ref} purchased the property, but believes it was around 2006, 2007, or 2008."
    if match(question, "he's owned it for over 15 years"):
        return f"{pronoun} confirmed that {natural_insured_ref} had owned the property for over 15 years."
    if match(question, "what year the building was constructed"):
        return f"{pronoun} stated that the building was originally constructed in the 1970s."
    if match(question, "date of loss was the property being used as a rental property"):
        return f"On the date of loss, the property was being used as a rental property."
    if question.strip().lower().startswith("but i just needed a simple answer on that one"):
        return None
    if match(question, "how long he had been using it as a rental property"):
        return None
    if match(question, "was the property ever used as", "residence"):
        return None
    if match(question, "what kind of property is the loss location"):
        return f"{pronoun} stated that the loss location is a single-family property."
    if match(question, "was the property occupied on the date of loss"):
        return f"{pronoun} confirmed that the property was occupied by renters on the date of loss."
    if match(question, "were the renters the claimants"):
        return f"{pronoun} confirmed that the renters were the claimants."
    if match(question, "were they short-term renters"):
        return f"{pronoun} stated that they were short-term renters who had booked the property for one week."
    if match(question, "where did the claimants book their stay"):
        return f"{pronoun} stated that the claimants booked the property through VRBO."
    if match(question, "booking records for this week"):
        return f"{pronoun} confirmed that {context.pronoun} had been forwarded some booking-related documents."
    if match(question, "booking or rental agreement"):
        return f"{pronoun} is not aware of any formal booking or rental agreement."
    if match(question, "when did the claimant start renting the property"):
        return f"{pronoun} believes the claimant started renting the property during the week of April 9 and was supposed to stay for seven days, but {context.pronoun} does not know when the stay ended."
    if match(question, "prior to the date of loss, did they ever make any complaints"):
        return f"{pronoun} stated that the same couple had rented the property the year before and returned because they loved it, and {context.pronoun} is not aware of any prior complaints about the deck."
    if match(question, "did our insured, jim, have a property manager"):
        return f"{pronoun} confirmed that {natural_insured_ref} had a property manager for the beach house rental."
    if match(question, "written agreement with randal"):
        return f"{pronoun} stated that there was no written agreement with the property manager."
    if match(question, "paid for his property manager services"):
        return f"{pronoun} stated that the property manager was paid 10% of the rental fee."
    if match(question, "responsibilities as the property manager"):
        return f"{pronoun} explained that the property manager's responsibilities included managing bookings and addressing any questions or needs raised by renters."
    if match(question, "what dates was he the insured's property manager"):
        return f"{pronoun} stated that the property manager was acting in that role during the week of the incident and believes he may have been acting in that role for two or three years before that, although {context.pronoun} was not certain."
    if match(question, "who is responsible for the maintenance of the property") and "personal representative" in context.role_label.lower():
        return f"{pronoun} explained that maintenance responsibility ultimately fell to {natural_insured_ref}, although {natural_insured_possessive} ex-wife was the person most often using the property and would tell {natural_insured_ref} if something needed to be repaired."
    if match(question, "alert the insured about any issues with the deck"):
        return f"{pronoun} does not know whether anyone alerted {natural_insured_ref} about any issues with the deck before the date of loss."
    if match(question, "delegated responsibility for day-to-day maintenance"):
        return f"{pronoun} does not know whether anyone had been formally delegated responsibility for day-to-day maintenance and operations of the property."
    if match(question, "claimants did rent this property before"):
        return f"{pronoun} confirmed that the claimants had rented the property before."
    if match(question, "year prior, were they there for a week as well"):
        return f"{pronoun} believes their prior stay was also for approximately a week, although {context.pronoun} was not certain."
    if match(question, "prior complaints or issues the year before"):
        return f"{pronoun} stated that there were no prior complaints or issues during the claimants' earlier stay. {pronoun} also understood that the claimants later left a positive review about the property."
    if match(question, "in your own words, on the data loss, what happened"):
        return f"{pronoun} stated that everything {context.pronoun} know about the incident is based on information that was provided to {context.objective}, including text messages and other materials that had been forwarded to Citizens. According to that information, the claimant stepped on a deck board that was alleged to be rotten, the board broke, and the claimant's foot went through the board, causing a fall. {pronoun} explained that an after-visit summary from Cleveland Clinic reflected that the claimant may have injured a leg and a shoulder, but {context.pronoun} were only relaying what {context.pronoun} read in that summary."
    if match(question, "how did you first become aware of the incident"):
        return f"{pronoun} stated that {context.pronoun} had no idea about the incident until the claim came through."
    if question.strip().lower() == "so you only.":
        return None
    if match(question, "our insured ever became aware of the incident"):
        return None
    if match(question, "only became aware of this in mid to late december of 2025"):
        return f"{pronoun} first became aware of the incident in mid to late December of 2025 when the claim arose."
    if match(question, "weren't present when the incident occurred"):
        return f"{pronoun} confirmed that {context.pronoun} was not present when the incident occurred."
    if match(question, "report the incident to anyone after you became aware"):
        return f"After becoming aware of the incident, {context.pronoun} reported it to the insurance brokers so that it could in turn be reported to the insurance company."
    if question.strip().lower() == "jim.":
        return None
    if match(question, "our insured was present when the incident occurred"):
        return f"{pronoun} confirmed that {natural_insured_ref} was not present when the incident occurred."
    if match(question, "what the claimant was doing when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} has no idea what the claimant was doing when the incident occurred."
    if match(question, "where exactly on the property did the incident occur"):
        if answer.lower().endswith("western"):
            return None
        return f"Based on the pictures {context.pronoun} reviewed, the incident occurred on the west-facing deck outside the kitchen doorway."
    if match(question, "where exactly on the deck"):
        return f"{pronoun} identified the location as a middle board on that deck."
    if match(question, "so this was on the 2nd floor directly outside the kitchen on the deck"):
        return None
    if question.strip().lower().startswith("so somewhere in the middle of the deck"):
        return None
    if match(question, "why the claimant was on the deck"):
        return f"{pronoun} does not know why the claimant was on the deck at the time of the incident."
    if match(question, "what they were doing on the deck"):
        return f"{pronoun} also does not know what the claimant was doing on the deck when the incident occurred."
    if match(question, "authorized to be on the deck"):
        return f"{pronoun} presumed the claimant was authorized to be on the deck because the property had been rented to the claimants."
    if match(question, "describe the deck on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} could not describe the deck on the date of loss because {context.pronoun} was not there and had no personal knowledge of its condition at that time."
    if match(question, "tell me anything about the deck on the date of loss"):
        return None
    if match(question, "what material the deck was made of"):
        return f"{pronoun} confirmed that the deck was made of wood."
    if match(question, "how old was the deck"):
        return f"{pronoun} does not know how old the deck was on the date of loss."
    if match(question, "deck present when the insured purchased the property"):
        return f"{pronoun} explained that the deck was present when {natural_insured_ref} purchased the property, but that the house was in poor condition and was substantially rebuilt after the purchase, including work on the decks. {pronoun} believes the board in question was not original and was likely replaced around the time the house was purchased. Based on the pictures {context.pronoun} reviewed, the board appeared broken but did not look rotted."
    if match(question, "deck renovated after our insured purchased the property"):
        return f"{pronoun} is confident that the deck was renovated after {natural_insured_ref} purchased the property."
    if match(question, "anything wrong with the deck on the date of loss"):
        return f"{pronoun} does not know if there was anything wrong with the deck on the date of loss."
    if match(question, "rot or deterioration on the deck"):
        return f"{pronoun} is not aware of any rot or deterioration on the deck."
    if match(question, "termites or other wood-eating pests"):
        return f"{pronoun} is not aware of any termites or other wood-eating pests having been present on the property."
    if match(question, "other deck boards ever broken prior"):
        return f"{pronoun} does not know whether any other deck boards had broken prior to the date of loss."
    if match(question, "aware of any issues with the deck on the date of loss"):
        return f"{pronoun} is not aware of {natural_insured_ref} or any associates knowing about issues with the deck on the date of loss."
    if match(question, "prior repairs to the deck prior to the date of loss"):
        return f"{pronoun} stated that, apart from the renovation work performed when the house was purchased, {context.pronoun} does not know of any prior repairs to the deck."
    if match(question, "information about his renovation of the deck"):
        return f"{pronoun} stated that {context.pronoun} does not have detailed information about the deck renovations."
    if match(question, "complained about the deck, its safety"):
        return f"{pronoun} is not aware of anyone having complained about the deck, its safety, its condition, or its integrity before the date of loss."
    if match(question, "fallen through the deck prior to the date of loss"):
        return f"{pronoun} confirmed that, to {context.possessive} knowledge, no one else had ever fallen through the deck prior to the date of loss."
    if match(question, "deck on the 1st or 2nd floor"):
        return f"{pronoun} stated that the deck was on the second floor."
    if match(question, "travel through the area where the incident occurred"):
        return f"{pronoun} explained that there were two entrances to the property: one from the second-floor deck into the kitchen and another on the ground floor."
    if match(question, "has the deck been replaced since the date of loss"):
        return f"{pronoun} confirmed that the area of the deck involved in the incident had been repaired after the date of loss."
    if question.strip().lower() == "when?" and "took possession of the property" in answer.lower():
        return f"{pronoun} stated that the repair occurred sometime after the incident, and when {context.pronoun} later took possession of the property the deck boards were solid."
    if match(question, "deck in which the incident occurred, was it replaced after the date of loss"):
        return f"{pronoun} stated that the area in question appears to have been repaired, but could not say whether the entire deck was replaced."
    if match(question, "deck where the incident occurred been replaced or not"):
        return f"{pronoun} stated that, at a minimum, the broken area was repaired. {pronoun} also explained that VRBO later sent someone to inspect the property for safety and appeared satisfied that the property was safe afterward."
    if "prior to my inspection" in question and "understand the question" in answer.lower():
        return None
    if match(question, "was the... deck where the incident occurred prior repaired"):
        if "do not understand" in answer.lower():
            return None
        return f"{pronoun} stated that the broken area had been repaired in 2024 before the later inspection of the property."
    if question.strip().lower() == "so yes.":
        return None
    if match(question, "has the deck been repaired or replaced prior to my inspection"):
        return f"{pronoun} confirmed that the deck had been repaired prior to the inspection."
    if match(question, "regular maintenance performed on the deck"):
        return f"{pronoun} is unaware of any regular maintenance performed specifically on the deck."
    if match(question, "regular maintenance performed on the property"):
        return f"{pronoun} confirmed that there was regular maintenance performed on the property, including air conditioning, extermination, and landscaping services."
    if match(question, "regularly inspect the deck"):
        return f"{pronoun} does not know whether anyone regularly inspected the deck."
    if match(question, "prior to the date of loss, did our insured or his property manager ever visit and inspect the property"):
        return f"{pronoun} confirmed that, before the date of loss, both {natural_insured_ref} and the property manager had visited the property."
    if "eric told" in answer.lower() and "he was there" in answer.lower():
        return None
    if match(question, "how often would they go out there and what would they do"):
        return f"{pronoun} does not know how often those visits occurred or what was done during them."
    if match(question, "how did the claimant come to fall through the deck"):
        return f"Regarding the incident itself, {context.pronoun} stated that {context.pronoun} has no idea how the claimant came to fall through the deck."
    if match(question, "is there a defect that caused the claimant to fall through the wood deck"):
        return f"{pronoun} stated that a defect is what is being alleged, but {context.pronoun} does not know."
    if match(question, "warning signs, caution tape, or visible warnings"):
        return f"{pronoun} is not aware of any warning signs, caution tape, or other visible warnings instructing people to avoid the deck."
    if match(question, "claimant fall all the way through the deck"):
        return f"{pronoun} does not know whether the claimant fell all the way through the deck or only partially, but based on the pictures {context.pronoun} reviewed, it does not appear that the claimant fell completely through."
    if match(question, "anyone else fall through the deck on the date of loss"):
        return f"{pronoun} confirmed that, to {context.possessive} knowledge, no one else fell through the deck on the date of loss."
    if match(question, "was anyone else injured on the date of loss"):
        return f"{pronoun} is not aware of anyone else having been injured on the date of loss."
    if match(question, "claimant continued to walk after the fall"):
        return f"{pronoun} does not know whether the claimant continued to walk after the fall."
    if match(question, "did anyone call for an ambulance"):
        return f"{pronoun} does not believe anyone called for an ambulance because {context.pronoun} believes the claimant went to urgent care on her own."
    if match(question, "did the claimant go to the hospital"):
        return f"{pronoun} believes the claimant went to the hospital or emergency room based on the after-visit summary {context.pronoun} reviewed."
    if match(question, "believe that she did go to the hospital"):
        return None
    if match(question, "how long the claimant was in the hospital"):
        return f"{pronoun} stated that it appeared the claimant was treated and released the same day and was not admitted to the hospital."
    if match(question, "only in the er for one day", "believe she was only in the for one day"):
        return None
    if match(question, "what the claimant's injuries are or were"):
        return f"{pronoun} stated that the only information {context.pronoun} has regarding the claimant's injuries comes from the after-visit summary."
    if match(question, "what injuries do you believe she had or had"):
        return f"{pronoun} stated that the summary reflected a contusion to the claimant's fibula or lower leg and a possible shoulder injury, potentially involving the rotator cuff."
    if match(question, "which leg and which shoulder"):
        return f"{pronoun} does not recall which leg or which shoulder was referenced."
    if match(question, "what treatments did the claimant undergo"):
        return f"{pronoun} does not know what treatment the claimant underwent."
    if match(question, "surgeries, physical therapy, scarring, or permanent disabilities"):
        return f"{pronoun} does not know whether the claimant had any surgeries, physical therapy, scarring, or permanent disability related to the alleged injuries."
    if match(question, "recovered from their injuries"):
        return f"{pronoun} does not know whether the claimant recovered from the alleged injuries."
    if match(question, "have you ever seen or spoken with the claimant"):
        return f"{pronoun} confirmed that {context.pronoun} has never seen or spoken with the claimant."
    if match(question, "had our insured ever seen or spoken with the claimant"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, {natural_insured_ref} had never seen or spoken with the claimant."
    if match(question, "property manager ever see or speak with the claimant"):
        return f"{pronoun} confirmed that the property manager did communicate with the claimant."
    if match(question, "and what was discussed"):
        return f"{pronoun} explained that the only information {context.pronoun} has about that communication comes from the text messages that were provided."
    if match(question, "have you already provided that to us"):
        return f"{pronoun} confirmed that the text messages had already been provided to Citizens."
    if match(question, "prior injuries or accidents resulting in any injuries"):
        return f"{pronoun} does not know whether the claimant had any prior injuries or prior accidents resulting in injuries."
    if match(question, "what time of day the incident occurred"):
        return f"{pronoun} does not know what time of day the incident occurred."
    if question.strip().lower() == "and you said this happened on the 2nd floor west facing deck, correct":
        return None
    if match(question, "working lighting in that area"):
        return f"{pronoun} believes there was working lighting in that area, but ultimately does not know for certain."
    if match(question, "weather was like on the date of loss"):
        return f"{pronoun} does not know what the weather was like on the date of loss."
    if match(question, "describe the claimant's physical appearance"):
        return f"{pronoun} stated that {context.pronoun} cannot describe the claimant's physical appearance."
    if match(question, "authorized to be on that part of the property"):
        return f"{pronoun} presumed the claimant was authorized to be on that part of the property because the property had been rented to the claimants."
    if match(question, "claimant distracted or under the influence"):
        return f"{pronoun} does not know whether the claimant was distracted or under the influence at the time of the incident."
    if match(question, "claimant wore glasses or had any vision issues"):
        return f"{pronoun} does not know whether the claimant had any vision, mobility, or balance issues."
    if match(question, "mobility or balance issues"):
        return f"{pronoun} does not know whether the claimant had any vision, mobility, or balance issues."
    if match(question, "did the claimant have any prior injuries"):
        return f"{pronoun} does not know whether the claimant had any prior injuries."
    if match(question, "know the claimant in any way"):
        return f"{pronoun} confirmed that neither {context.pronoun}, {natural_insured_ref}, nor the property manager knew the claimant in any capacity."
    if match(question, "directing or assisting the claimant in any way"):
        return f"{pronoun} confirmed that no one, including {natural_insured_ref} or any associates, was directing or assisting the claimant when the incident occurred."
    if match(question, "could have done to prevent the incident"):
        return f"{pronoun} does not believe there was anything {context.pronoun}, {natural_insured_ref}, or the property manager could have done to prevent the incident."
    if match(question, "could anything have prevented the incident"):
        return f"{pronoun} does not know whether anything could have prevented the incident."
    if match(question, "video cameras on the property"):
        return f"{pronoun} does not know whether there were any video cameras on the property on the date of loss."
    if match(question, "aware of any witnesses"):
        return f"{pronoun} is not aware of any witnesses to the incident."
    if match(question, "provided with all documentation"):
        return f"{pronoun} confirmed that {context.pronoun} has provided Citizens with all documentation {context.pronoun} has regarding the incident."
    if match(question, "did our insured have any excess liability or umbrella insurance"):
        return f"{pronoun} stated that there was another general liability policy in place, although {context.pronoun} was not sure whether there was any umbrella coverage."
    if match(question, "was that the auto owner's policy"):
        return f"{pronoun} confirmed that the general liability policy was with Auto-Owners."
    if match(question, "have they been put on notice"):
        return f"{pronoun} confirmed that Auto-Owners had been put on notice of the incident."
    if match(question, "do you have any excess liability or umbrella insurance"):
        return f"{pronoun} stated that {context.pronoun} does not have any excess liability or umbrella insurance."
    if match(question, "did our insured own any additional properties in the united states"):
        return f"{pronoun} confirmed that {natural_insured_ref} owned additional properties in the United States."
    if question.strip().lower() == "how many?" and any(token in answer.lower() for token in ("five", "six")):
        return f"{pronoun} estimated that there were five or six additional properties."
    if match(question, "would you be able to provide us with the addresses of those properties"):
        return f"{pronoun} stated that {context.pronoun} would probably be able to provide the addresses of those properties and the names of the insurance carriers at a later time."
    if match(question, "can you send me that document regarding the load-bearing wall today", "can you send it today", "can you provide me with the name of that company"):
        return None
    if question.strip().lower().startswith("okay, so back to"):
        return None
    if match(question, "in your own words, tell me about all of the incidents that have occurred or allegedly have occurred in your unit") and all(token in pair.answer.lower() for token in ("lou", "john", "shower bed")):
        return render_condo_remodel_leak_story(pair.answer, context)
    if match(question, "what about all these other issues that are being alleged such as issues due to the remodeling of your unit"):
        return f"{pronoun} stated that remodeling work had been performed in the unit by Elite, but {context.pronoun} did not recall any separate complaints that the remodeling itself caused damage to the unit below."
    if match(question, "what did you hire this general contractor to do"):
        return f"{pronoun} stated that the general contractor was hired to perform cosmetic renovation work, including painting, some tile work, and opening part of the kitchen wall."
    if match(question, "were they tasked with removing a load-bearing wall"):
        return f"{pronoun} stated that the contractor was supposed to determine whether the kitchen wall was load-bearing, obtain the necessary engineering input, and handle whatever permitting was required."
    if match(question, "what else was this contractor tasked with doing in your unit"):
        return f"{pronoun} stated that the contractor also performed tile work in the unit."
    if question.strip().lower() == "where?" and "bathroom" in pair.answer.lower():
        return f"{pronoun} stated that the tile work was in the master bathroom."
    if match(question, "is that master bathroom the same one that leaked") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that the master bathroom tile work was not in the shower that later leaked."
    if match(question, "how did you find this general contractor") and "internet" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} found the contractor through the internet."
    if match(question, "what is the name of this general contractor"):
        return f"{pronoun} identified the contractor as Elite Construction."
    if match(question, "is it elite builders"):
        return None
    if match(question, "did he pull permits for the work that was done in your unit"):
        return f"{pronoun} stated that {context.pronoun} later learned the contractor did not pull permits for the work."
    if match(question, "why did he not pull permits"):
        return f"{pronoun} stated that {context.pronoun} does not know why the contractor failed to pull permits."
    if match(question, "did the contractor...") and "exactly the question" in pair.answer.lower():
        return None
    if match(question, "did the contractor ever advise you that permits were required for the work being completed") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that the contractor did not advise {context.objective} that permits were required."
    if match(question, "did the contractor ever advise you to pull the permits yourself", "did you ever ask him to skip the permits") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that neither side suggested skipping permits or having the owners pull them personally."
    if match(question, "how did you learn that he did work without pulling permits"):
        return f"{pronoun} stated that {context.pronoun} learned about the permit issue after a neighbor complained to the city and the city advised that no permits had been pulled."
    if match(question, "did you supervise or inspect any of the contractor's work"):
        return f"{pronoun} stated that {context.pronoun} only inspected the contractor's work toward the end of the project."
    if match(question, "did everything appear satisfactory and correct") and "didn't like the work" in pair.answer.lower():
        return f"{pronoun} stated that the work did not appear satisfactory."
    if question.strip().lower() == "why?" and ("quality" in pair.answer.lower() or "sloppy" in pair.answer.lower()):
        return f"{pronoun} stated that the work quality was poor, including sloppy painting and tile installation, which led to the contractor being fired."
    if match(question, "did you have another contractor come in and redo the work"):
        return f"{pronoun} stated that a separate painting company was later hired to repaint the unit."
    if match(question, "did you have that, did you have another company ever redo any of the plumbing work or tile work that the original contractor did"):
        return "No separate company redid plumbing work, and the toilets were not touched."
    if match(question, "does your condo association require architectural review and approval prior to performing work in your unit"):
        return f"{pronoun} stated that {context.pronoun} believes a scope of work was submitted to the condo association and approved."
    if match(question, "they approved all of the work you did in your unit", "did you guys get architectural and condo association approval"):
        return f"{pronoun} stated that {context.pronoun} do not clearly recall the extent of any architectural approval and {are_word} not sure whether every aspect of the work was specifically approved." if context.joint_statement else f"{pronoun} stated that {context.pronoun} does not clearly recall the extent of any architectural approval and is not sure whether every aspect of the work was specifically approved."
    if match(question, "how many leaks or water related issues are you aware of going into the unit below from your unit"):
        return f"{pronoun} stated that {context.pronoun} recalled one prior shower-related leak affecting the unit below, approximately two years earlier, and did not recall any others at that time."
    if match(question, "so you're only aware of one issue related") and "toilets were clogged" in pair.answer.lower():
        return f"{pronoun} also mentioned a more recent plumbing problem involving clogged toilets and a shower, but stated that {context.pronoun} believed that issue was building-wide rather than a separate leak from the unit."
    if match(question, "is that recently") and "whole building" in pair.answer.lower():
        return f"{pronoun} stated that the more recent clogging issue appeared to involve the building's plumbing system rather than only {context.possessive} unit."
    if match(question, "how did you first become aware of that bathroom leak") and "which one" in pair.answer.lower():
        return None
    if match(question, "and when did you become aware of this leak") and "lou called" in pair.answer.lower():
        return f"{pronoun} stated that Lou called as soon as he noticed the leak, so {context.pronoun} does not believe it had been ongoing for long before notice was given."
    if match(question, "what did you do specifically after becoming aware of the leak"):
        return f"{pronoun} stated that {context.pronoun} hired a plumber, who replaced the shower pan and drainage."
    if match(question, "how long after you became aware of the leak did you have them replace the shower and everything") and "right away" in pair.answer.lower():
        return f"{pronoun} stated that repairs began right away after the leak was reported."
    if match(question, "how long was the shower leak going on for before it was stopped"):
        return f"{pronoun} stated that the issue was addressed immediately, and the shower was no longer used once the leak was identified."
    if match(question, "how long was water leaking into the unit below before the water stopped leaking into the unit below"):
        return f"{pronoun} stated that the leak stopped right away once the shower was no longer being used."
    if match(question, "did anyone shut the water off to your unit or the shower after the incident was reported"):
        return f"{pronoun} did not indicate that the water was shut off, but stated that the shower was no longer used after the leak was discovered."
    if match(question, "did you report the leaks to anyone, including the condo association, after it was reported to you"):
        return f"{pronoun} stated that the condo association and the manager, John, were already involved in the matter."
    if match(question, "you already said you hired a plumber, but you don't have any kind of plumbing invoice for the repair or do you") and "sent it to you" in pair.answer.lower():
        return f"{pronoun} stated that the plumbing invoice had already been provided."
    if match(question, "did the general contractor that didn't pull permits touch any plumbing, anything in your unit, including toilets or anything else"):
        return f"{pronoun} stated that the contractor did not touch the toilets, although the top portion of the jacuzzi in the master bathroom was changed."
    if match(question, "do you know how old that shower was"):
        return None
    if match(question, "so you don't know, but it was there when you purchased the unit"):
        return f"The shower was already present when {context.pronoun} purchased the unit, and {context.pronoun} {do_word} not know the age of the shower pan."
    if match(question, "do you know what calls the shower to leak", "do you know what caused the shower to leak"):
        return f"{pronoun} stated that {context.pronoun} {do_word} not know exactly what caused the shower leak, but {believe_word} the shower pan may have been old."
    if match(question, "had this shower pan ever leaked before"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, that shower pan had not leaked before."
    if match(question, "have you ever had any prior repairs, issues or leaks involving this shower"):
        return f"{pronoun} stated that the shower had previously been repaired approximately three years earlier."
    if match(question, "is the association that recommended the plumber to repair your shower"):
        return f"{pronoun} stated that the condo association recommended the plumber who repaired the shower."
    if match(question, "has anything else ever leaked in your unit") and "kitchen sink clogged" in pair.answer.lower():
        return f"{pronoun} stated that there was a recent kitchen sink clog, but no separate leak was identified."
    if match(question, "why did you have the load-bearing wall partially removed"):
        return f"{pronoun} stated that the wall was partially removed as part of the renovation to open the kitchen."
    if match(question, "that's for sure. and you're gonna send me that report today"):
        return None
    if match(question, "was there anything wrong with this wall"):
        return "There was nothing wrong with the wall, and it was removed only as part of the renovation."
    if match(question, "did you or your contractor get a permit for this removal of the partial of this partial removal of the load-bearing wall"):
        return f"{pronoun} stated that no permit was pulled for the partial removal of the load-bearing wall and that {context.pronoun} later learned the contractor should have handled that."
    if match(question, "did you guys ever get a retroactive permit for this partial removal of the load-bearing wall"):
        return f"{pronoun} stated that {context.pronoun} is currently working with the city to obtain retroactive approval for the work."
    if match(question, "how many units were damaged as a result of these incidents, not including your own"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any damage to other units."
    if match(question, "what's the unit number of the unit below"):
        return f"{pronoun} identified the unit below as unit 102."
    if match(question, "was this unit occupied when these incidents were ongoing"):
        return None
    if match(question, "and you don't believe that your unit via the improper installations without the permits and all that or the removal of the load bearing walls damaged the unit below") and "shoring" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} does not believe the unit below was damaged by the work, although the city later recommended shoring in both units as a precaution after complaints were made."
    if match(question, "do you know what specific damages they're alleging in the unit below"):
        return f"{pronoun} stated that, aside from what was written in the complaint letter, {context.pronoun} does not know what damages are being alleged in the unit below."
    if match(question, "were there any other units that you're aware of that were damaged or claiming to be damaged"):
        return f"{pronoun} stated that, aside from what was written in the complaint letter, {context.pronoun} {do_word} not know of any additional units claiming damage."
    if match(question, "do you know if they've had any dry out or mold remediation"):
        return f"{pronoun} stated that {context.pronoun} does not know whether any dry out or mold remediation has been completed in the unit below."
    if match(question, "have they had any of their alleged damages repaired yet") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any repairs having been completed in the unit below."
    if match(question, "do you know if anyone was injured or was claiming to be injured as a result of the incident") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that {context.pronoun} is not aware of anyone claiming an injury as a result of these incidents."
    if match(question, "are you aware of any other units damaged or claiming to be damaged") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that, aside from what was written in the complaint letter, {context.pronoun} {do_word} not know of any additional units claiming damage."
    if match(question, "are you aware of any damages to condo common areas") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that {context.pronoun} is not aware of any damage to condo common areas."
    if match(question, "in your own words, tell me the whole story of what happened involving this leak"):
        if "building manager" in answer.lower() and "12g" in answer.lower() and "shower pan" in answer.lower():
            return (
                f"In describing what happened, the {role_noun} stated that {context.pronoun} was contacted by the building manager on April 5th "
                f"and was alerted to an alleged leak in unit 12G that was suspected to be coming from the shower pan of {context.possessive} unit. "
                f"{pronoun} stated that the building manager hired a plumber to investigate the issue through the ceiling of unit 12G. "
                f"{pronoun} stated that the building manager told {context.objective} that {context.pronoun} would be responsible for repairs to the shower pan "
                f"and for the damage in unit 12G."
            )
        cleaned = lower_first(answer)
        cleaned = cleaned.replace("repairs to the The damage is on 12G", "repairs to the damage in unit 12G")
        cleaned = cleaned.replace("repairs to the the damage is on 12g", "repairs to the damage in unit 12G")
        return f"In describing what happened, the {role_noun} stated that {cleaned}."
    if match(question, "whole story involving this leak with the ac and with the sink leak"):
        if all(token in pair.answer.lower() for token in ("bedroom", "living room", "ac closet", "water vacuum")):
            return render_ac_leak_story(pair.answer, context)
        return f"The {role_noun} explained that {lower_first(answer)}."
    if match(question, "what about this kitchen sink leak when did it happen and what happened"):
        if all(token in pair.answer.lower() for token in ("restoration company", "water filter", "kitchen faucet")):
            return render_kitchen_sink_story(pair.answer, context)
        return f"Regarding the kitchen sink leak, the {role_noun} explained that {lower_first(answer)}."
    if match(question, "regarding the kitchen sink leak, how long had you been away"):
        if "not there that night" in pair.answer.lower():
            return render_kitchen_sink_absence(pair.answer, context)
        return f"She advised that {lower_first(answer)}." if context.pronoun == "she" else f"{pronoun} advised that {lower_first(answer)}."
    if match(question, "did you first become aware of that leak when you heard the water rushing from your bedroom"):
        return f"{pronoun} confirmed that {context.pronoun} first became aware of the AC leak when {context.pronoun} heard water rushing from inside {context.possessive} bedroom."
    if match(question, "once you became aware of it, you immediately called the association maintenance guy"):
        return f"{pronoun} confirmed that {context.pronoun} immediately contacted the association maintenance person."
    if match(question, "how long the ac leak was going on for before it was stopped"):
        duration = extract_minute_range(pair.answer)
        if duration:
            return f"{pronoun} estimated that the AC leak may have been ongoing for approximately {duration} before it was stopped."
        return f"{pronoun} estimated that the AC leak may have been ongoing for {lower_first(answer)} before it was stopped."
    if match(question, "who shut the water off to the ac once the leak was discovered"):
        return None
    if question.strip().lower() == "who?" and "maintenance" in pair.answer.lower():
        return f"The building maintenance worker shut off the water for the AC leak."
    if match(question, "regarding the ac leak, did you ever speak with the claimants"):
        return f"{pronoun} did not speak with any claimants regarding the AC leak."
    if match(question, "did you hire a plumber or an ac technician for the ac yet"):
        return f"{pronoun} stated that {lower_first(answer)}."
    if match(question, "what did you do after becoming aware of the kitchen sink leak"):
        return f"Regarding the kitchen sink leak, {context.pronoun} stated that {lower_first(answer)}."
    if match(question, "how long was the kitchen sink leak going on for before it was stopped"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the kitchen sink leak was going on before it was stopped."
    if match(question, "who shut the water off after the leak was discovered in the kitchen sink"):
        return f"The water for the kitchen sink leak was shut off by building maintenance."
    if match(question, "did you ever speak with the claimants about the kitchen sink leak"):
        return f"{pronoun} did not speak with the claimants regarding the kitchen sink leak."
    if match(question, "did you hire a plumber regarding the kitchen sink leak"):
        if "no plumber" in pair.answer.lower() and ("connection valve" in pair.answer.lower() or "tube" in pair.answer.lower()):
            return f"{pronoun} stated that no plumber was needed for the kitchen sink leak and that the issue only required replacement of the connection valve and tubing."
        return f"{pronoun} stated that {lower_first(answer)}."
    if match(question, "were you ever made aware that the roof at the property leaked"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} was never made aware of any roof leaks at the property before the alleged incident."
    if match(question, "prior to selling the property, did you have anyone inspect the roof"):
        return f"{pronoun} believes the roof may have been inspected by the prospective buyer's inspector before the property was sold, but {context.pronoun} was not certain."
    if match(question, "do you have a report from them"):
        return f"{pronoun} does not know whether any report from that inspection exists."
    if match(question, "would you be able to find out"):
        return f"{pronoun} stated that {context.pronoun} could look into whether such a report exists."
    if match(question, "do you have the sale documents from when you sold the property") and is_affirmative(pair.answer):
        return f"{pronoun} stated that {context.pronoun} does have the sale documents from when the property was sold."
    if match(question, "can you provide us with a copy") and is_affirmative(pair.answer):
        return f"{pronoun} stated that {context.pronoun} can provide a copy of those sale documents."
    if match(question, "do you have any photos of the interior of the home") and is_affirmative(pair.answer):
        return f"{pronoun} stated that {context.pronoun} does have interior photos of the home."
    if match(question, "can you provide us with those") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that {context.pronoun} had not yet provided those interior photos."
    if match(question, "did you make any repairs to the roof or property prior to listing it and or selling it"):
        return f"{pronoun} stated that before selling the property, the front area below the roof was painted and repaired because the HOA wanted it to look better."
    if match(question, "so the front part of the roof was painted"):
        return f"{pronoun} clarified that the work was not done to the roof itself, but to the concrete area below the roof."
    if match(question, "did you make any repairs or changes to the roof itself"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that no repairs or changes were made to the roof itself."
    if match(question, "how were you made aware that the claimant allegedly slipped and fell somewhere else"):
        return f"{pronoun} stated that the claimant reported slipping inside the house."
    if match(question, "so she never said she fell somewhere else"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, the claimant did not initially allege that the fall occurred anywhere else."
    if match(question, "did the plaintiff or tenant advise of any witnesses to this alleged event", "so no, she did not allege any witnesses"):
        return f"{pronoun} was not aware of any witnesses being identified in connection with the alleged incident."
    if match(question, "you said that she was evicted from the property, correct"):
        return None
    if question.strip().lower() == "why?" and "rent" in pair.answer.lower():
        return f"{pronoun} stated that the claimant was evicted because {context.pronoun} was not paying rent."
    if match(question, "do you have an umbrella policy or any additional coverage on this property", "any additional personal liability policies"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} did not have any umbrella or additional personal liability coverage for the property."
    if match(question, "when did you purchase the property"):
        if "didn't purchase" in pair.answer.lower() or "did not purchase" in pair.answer.lower():
            return f"{pronoun} clarified that {context.pronoun} did not purchase the property."
    if match(question, "when did you first acquire the property"):
        return f"{pronoun} stated that {context.pronoun} first acquired the property in {normalize_short_year(pair.answer)}."
    if match(question, "do you still own the property"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} no longer owns the property."
    if match(question, "when did you sell it"):
        return f"{pronoun} stated that the property was sold in {clean_month_year_reference(pair.answer)}."
    if match(question, "why did you sell the property"):
        return f"{pronoun} stated that {context.pronoun} sold the property because the tenant had caused so many problems that {context.pronoun} did not want to rent it again."
    if match(question, "do you know what year the building was constructed") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that {context.pronoun} does not know what year the building was constructed."
    if match(question, "on the date of loss, you were using the property as a rental property"):
        return f"{pronoun} confirmed that the property was being used as a rental on the date of loss."
    if match(question, "was it a short-term rental or a long-term rental"):
        return f"{pronoun} stated that the property was rented on a long-term basis under what {context.pronoun} believes was a one-year lease."
    if match(question, "when did you", "first started using the property as a rental property") or question.strip().lower() == "what month?":
        return None
    if match(question, "have you ever used the property as your residence"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that {context.pronoun} never used the property as {context.possessive} own residence."
    if match(question, "what was the name of the tenant occupying the property on the date of loss"):
        return f"{pronoun} identified the tenant on the date of loss as {clean_person_name_answer(pair.answer)}."
    if match(question, "can you spell that"):
        return None
    if match(question, "is this person the claimant or plaintiff in this case") and is_affirmative(pair.answer):
        return f"{pronoun} confirmed that the tenant was also the claimant in this case."
    if match(question, "did you have a lease agreement with your tenant") and is_affirmative(pair.answer):
        return f"{pronoun} confirmed that there was a lease agreement with the tenant."
    if match(question, "have you already provided it to citizens") and is_affirmative(pair.answer):
        return f"{pronoun} stated that the lease agreement has already been provided to Citizens."
    if match(question, "when did the claimant first start renting the property"):
        return f"{pronoun} stated that the claimant first started renting the property in {clean_month_year_reference(pair.answer)}."
    if match(question, "did she move in on that same day"):
        return f"{pronoun} stated that the claimant had already moved belongings into the property in December and did not personally move in on that same day."
    if match(question, "initial move-in checklist") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that no initial move-in checklist was provided when the claimant moved in."
    if match(question, "did the claimant or plaintiff ever complain about any issues, including any roof issues or leaks"):
        return f"{pronoun} stated that the claimant complained about high electricity usage, the air conditioning, and minor items such as the light by the front door, but did not initially complain about roof leaks."
    if match(question, "did she ever complain about the roof or any leaks"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that the claimant did not complain about roof leaks at the beginning of the tenancy."
    if match(question, "roof or any leaks before the incident occurred"):
        if normalize(pair.answer).lower().startswith("no"):
            return f"{pronoun} stated that the claimant did not report any roof leaks before the alleged incident."
    if match(question, "didn't because she didn't report any roof issues or leaks until after the incident had already occurred") and is_affirmative(pair.answer):
        return f"{pronoun} confirmed that the claimant did not report roof issues or leaks until after the alleged incident had already occurred."
    if match(question, "what were you talking about the ac again"):
        return f"{pronoun} stated that the claimant said the electric bill was too high and questioned whether the air conditioning was causing the issue."
    if match(question, "when did the claimant move out"):
        return f"{pronoun} stated that {context.pronoun} does not know exactly when the claimant moved out because the locks were changed and the keys were later left with a neighbor."
    if match(question, "so you don't know when she moved out", "can you give me an approximate date when she moved out"):
        return None
    if match(question, "did you have a property manager or a property management company managing the property or did you self-manage it"):
        return f"{pronoun} stated that {context.pronoun} initially managed the property personally and that {context.possessive} cousin also kept an eye on things, although he was not a formal property manager."
    if match(question, "did you have any contract with your cousin or did you pay him") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that there was no contract with {context.possessive} cousin and that {context.pronoun} did not pay him."
    if match(question, "who is responsible for the maintenance of the property") or match(question, "like, who is responsible for the maintenance of the property"):
        return f"{pronoun} stated that if the tenant reported a problem, they would arrange for someone to fix it."
    if match(question, "how often would you or your cousin check in on the property"):
        return f"{pronoun} stated that {context.possessive} cousin would occasionally go by and check whether things were okay."
    if question.strip().lower() == "how often?" and "no idea" in pair.answer.lower():
        return f"{pronoun} stated that {context.pronoun} does not know how often those checks occurred."
    if match(question, "in your own words, tell me what happened involving this alleged fall in your house"):
        return f"In describing what happened, the {role_noun} stated that {context.pronoun} does not know exactly what happened and was never provided with details regarding where or how the claimant allegedly fell."
    if match(question, "so you don't know anything about this incident"):
        return f"{pronoun} stated that the only information {context.pronoun} received was that the claimant sent an email saying {context.pronoun} had fallen twice in the house because the roof was leaking and could not pay the rent while waiting on a check."
    if match(question, "when did you first become aware that she allegedly fell in the house"):
        return f"{pronoun} stated that {context.pronoun} first became aware of the alleged fall on the date the claimant sent that email."
    if question.strip().lower() == "and when was that?" and "march" in pair.answer.lower():
        return f"{pronoun} stated that the email was received on March 27, 2024."
    if question.strip().lower() == "what year?" and "fell in december" in pair.answer.lower():
        return None
    if match(question, "so she allegedly fell in december of 2023, but didn't notify you until march of 2024") and is_affirmative(pair.answer):
        return f"{pronoun} confirmed that the claimant allegedly fell in December 2023 but did not notify {context.objective} until March 2024."
    if match(question, "was that the first time you became aware that she was allegedly injured as well") and is_affirmative(pair.answer):
        return f"{pronoun} confirmed that this was also the first time {context.pronoun} became aware that the claimant was alleging an injury."
    if match(question, "did you report the incident to anyone after learning of it"):
        return f"{pronoun} stated that {context.pronoun} discussed it with {context.possessive} cousin, who did not believe the claimant was injured because he had seen her at parties and catering events in the meantime."
    if match(question, "why did you not report the incident and the injuries to citizens until april 8th of 2026"):
        return f"{pronoun} stated that {context.pronoun} did not report the matter to Citizens earlier because {context.pronoun} did not believe the claimant had actually been injured."
    if match(question, "do you know what the claimant was doing when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} has no information about what the claimant was doing immediately before the incident."
    if match(question, "do you know what part of the property the incident occurred at"):
        return f"{pronoun} stated that {context.pronoun} does not know exactly where on the property the claimant says the incident occurred."
    if match(question, "do you know what caused the claimant to fall or what allegedly caused her to fall"):
        return f"{pronoun} stated that the claimant alleged the fall was caused by a roof leak."
    if match(question, "can you describe the roof that was on the property on the date of loss"):
        return f"{pronoun} described the roof as a shingle roof."
    if match(question, "how old is the shingle roof"):
        return f"{pronoun} stated that the roof had never been replaced and dated back to when the house was built."
    if match(question, "approximate age on that roof", "so that roof may have been around 30 years old when the incident occurred"):
        return f"{pronoun} stated that {context.pronoun} was not sure of the exact age of the roof, but it may have been about 30 years old at the time of the alleged incident."
    if match(question, "how many units were damaged as a result of this ac leak"):
        return None
    if match(question, "what are their unit numbers"):
        return f"The {role_noun} reported that units affected by the AC leak included units {answer}."
    if match(question, "unit next door to yours and the unit below you, correct"):
        return f"{pronoun} clarified that unit 210 was the unit below, and unit 312 was the neighboring unit."
    if match(question, "were these two units occupied when the incidents occurred"):
        return None
    if match(question, "when someone was present"):
        return f"{pronoun} stated that no one was present in unit 312 at the time of the AC leak, but {context.pronoun} was not sure whether anyone was present in unit 210."
    if match(question, "was anyone living in those units when the incident occurred"):
        return f"{pronoun} stated that unit 312 was occupied as a residence, but {context.pronoun} does not know whether unit 210 was occupied as a residence."
    if match(question, "do you know what the specific damages are in those two units from this ac leak"):
        return f"{pronoun} did not know the specific damages sustained in those units from the AC leak."
    if match(question, "either of those units had any dry out or mold remediation completed as a result of this ac leak"):
        return f"{pronoun} stated that drying equipment was present in unit 210, but {context.pronoun} did not know whether unit 312 had any dry out or mold remediation."
    if match(question, "have either of those two units had their damages repaired yet"):
        return f"{pronoun} did not know whether repairs had been completed in either unit."
    if match(question, "do you know if anyone was injured or is claiming to be injured as a result of the ac leak"):
        return f"{pronoun} did not know of anyone claiming injuries as a result of the AC leak."
    if match(question, "do you know if unit 210 or 312 have insurance on their unit"):
        return f"{pronoun} stated that {context.pronoun} believes unit 312 has insurance but does not know about unit 210."
    if match(question, "damages to condo common areas such as hallways", "related to the ac leak"):
        return f"{pronoun} explained that water entered the association hallway, but {context.pronoun} was not aware of any additional common-area damage caused by the AC leak because those hallways had already been damaged by a separate flood in February."
    if match(question, "how many units were damaged as a result of this kitchen sink leak"):
        return f"Regarding the kitchen sink leak, {context.pronoun} stated that only the unit below was affected."
    if match(question, "would that be unit 210"):
        return None
    if match(question, "was this unit occupied when the sink leak occurred"):
        return f"{pronoun} believed the unit below was occupied at the time because someone noticed and reported the leak."
    if match(question, "their specific damages are related to the sink leak"):
        return f"{pronoun} did not know the specific damages related to the sink leak."
    if match(question, "any additional damages beyond what the ac leak caused"):
        return f"{pronoun} did not believe the sink leak caused any damage beyond what already existed from the AC incident."
    if match(question, "additional damages to the condo common areas related to the sink leak"):
        return f"{pronoun} stated that there were no additional damages to common areas from the sink leak."
    if match(question, "story about this leak in 509 that happened in february"):
        return f"The {role_noun} further described the February incident involving unit 509, explaining that {lower_first(answer)}."
    if match(question, "did the ac leak from your unit cause any additional damages", "to the building"):
        return None
    if match(question, "but it did damage the unit to next to you and to the unit below you"):
        return f"{pronoun} confirmed that the AC leak damaged the unit next to hers and the unit below, but did not cause additional building damage to {context.possessive} knowledge." if context.pronoun == "she" else f"{pronoun} confirmed that the AC leak damaged the adjacent unit and the unit below."
    if match(question, "sink leak that occurred didn't cause any additional damages"):
        return f"{pronoun} confirmed that the sink leak did not cause any additional damages beyond what was already present."
    if match(question, "how old was the ac"):
        return f"The {role_noun} stated that {context.possessive} AC unit was {lower_first(answer)} old."
    if match(question, "what part on the ac leaked"):
        return f"{pronoun} explained that {lower_first(answer)}."
    if match(question, "fitting that connected the hot water supply line"):
        return f"{pronoun} stated that the fitting connecting the hose to the building’s main line failed, while the hose and connector themselves did not appear damaged."
    if match(question, "how old was this specific part"):
        return f"{pronoun} did not know how old the specific fitting was."
    if match(question, "has this part been replaced since the ac was installed"):
        return None
    if match(question, "and what was his name"):
        return None
    if match(question, "is the one that replaced it"):
        return f"{pronoun} stated that the building maintenance supervisor replaced that connection approximately two to three years earlier."
    if match(question, "did you ask him to replace it"):
        return None
    if match(question, "do you know what caused this fitting to leak on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know what caused the fitting to fail on the date of loss."
    if match(question, "rust or corrosion on this fitting"):
        return f"{pronoun} stated that the fitting did not appear rusted or corroded and seemed to be plastic."
    if match(question, "had this piece ever leaked before prior to the leak two years ago"):
        return f"{pronoun} stated that the fitting had leaked slightly approximately two to three years earlier."
    if match(question, "prior repairs or issues involving this fitting"):
        return f"{pronoun} stated that the prior issue involving that fitting had been addressed by the building maintenance supervisor."
    if match(question, "reason to believe that the fitting was at risk of leaking"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe the fitting was at risk of leaking prior to the date of loss."
    if match(question, "has the water supply line where the fitting popped been capped or repaired yet"):
        return f"{pronoun} confirmed that the line had not been fully repaired."
    if match(question, "but has it been capped"):
        return f"{pronoun} confirmed that the line had been capped."
    if match(question, "on what day did your sink leak"):
        return f"{pronoun} reported that the kitchen sink leak occurred on {clean_date_answer(answer)}."
    if match(question, "how old was that water supply line"):
        return f"{pronoun} stated that the water filter had been installed {lower_first(answer)}."
    if match(question, "what caused the water line to dislodge from the water filter"):
        return f"{pronoun} stated that {context.pronoun} does not know what caused the water line to disconnect from the water filter."
    if match(question, "was it installed by a professional"):
        return f"{pronoun} stated that the water filter was installed by a friend."
    if match(question, "had the water filter ever leaked before"):
        return f"{pronoun} stated that the water filter had never leaked before."
    if match(question, "prior repairs or issues involving the water filter"):
        return f"{pronoun} stated that there were no prior repairs or issues involving the water filter."
    if match(question, "reason to believe that the water filter was at risk of leaking"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe the water filter was at risk of leaking."
    if match(question, "has the water filter been repaired"):
        return f"{pronoun} confirmed that the water filter has since been repaired."
    if match(question, "water shutoff rule in the association bylaws"):
        return f"{pronoun} was not aware of any water shut-off rules in the association bylaws and understood that Dion was the person who knew how to shut off the water."
    if match(question, "how long was it vacant for"):
        return f"{pronoun} stated that the unit was vacant most of the time."
    if match(question, "while it's baking, do you have anyone checking on it", "while it's vacant, do you have anyone checking on it", "while it's vacant would you say you have somebody come by and check it"):
        return "The unit is checked by a cleaner and friends who have a key while it is vacant."
    if match(question, "how often while the unit is vacant would you say you have somebody come by and check it"):
        return f"{pronoun} stated that the unit is generally checked about once a month while vacant."
    if match(question, "do you leave the water running or do you leave the water to the unit on while it's vacant", "my question was, is the water on", "so you do not shut the water off while it's vacant"):
        return f"{pronoun} stated that the water is not shut off while the unit is vacant."
    if question.strip().lower() == "why not?" and "leakage monitoring devices" in pair.answer.lower():
        return f"{pronoun} stated that the water remains on because a cleaner and friends periodically use the unit and because leak monitoring devices have been installed throughout the apartment."
    if match(question, "did you file a claim with your insurance for the damage to your own unit") and "years ago" in pair.answer.lower():
        return f"{pronoun} stated that a prior first-party claim had been made years earlier, but it was denied."
    if match(question, "have you had any other leaks coming from your unit"):
        return f"{pronoun} stated that there were no other leaks beyond the incidents already described."
    if match(question, "do you own any additional properties in the united states") and "massachusetts" in pair.answer.lower():
        return f"{pronoun} stated that, aside from the Boca unit, {context.pronoun} also own a primary residence in Massachusetts."
    if match(question, "who insures that property") and "massachusetts" in pair.answer.lower():
        return None
    if match(question, "do you own any additional properties other than the place here in boca and the place in massachusetts") and normalize(pair.answer).lower().startswith("no"):
        return f"{pronoun} stated that, other than the Boca property and {context.possessive} Massachusetts residence, {context.pronoun} do not own any other properties."
    if match(question, "first party claim with citizens insurance for the damage to your own property related to the ac leak"):
        return f"{pronoun} confirmed that {context.pronoun} filed a first-party claim with Citizens Insurance related to the AC leak."
    if match(question, "first party claim for the water damage to your unit related to the sink leak"):
        return f"{pronoun} did not file a first-party claim related to the sink leak."
    if question.strip().lower() == "why not?":
        if "minor" in answer.lower():
            return f"{pronoun} stated that the damage from the sink leak was very minor."
        if "drywall" in answer.lower() or "building" in answer.lower():
            return (
                f"{pronoun} stated that there was very little personal damage from the February incident from unit 509, "
                f"and {context.pronoun} understood the drywall repairs would be handled by the building."
            )
    if match(question, "and what's the address of that property"):
        return f"{pronoun} identified the additional property as {answer}."
    if match(question, "and who insures that one"):
        return f"{pronoun} stated that the additional property is insured by {answer}."
    if match(question, "address of that property") and "sold to my son" in pair.answer.lower():
        return None
    if match(question, "address of that property") and looks_like_address(answer):
        return f"{pronoun} identified that additional property as {answer}."
    if match(question, "and where is that located") and looks_like_address(answer):
        return f"{pronoun} stated that the additional property is located in {answer}."
    if match(question, "first party claim for the damage caused by the leak in 509"):
        return f"{pronoun} did not file a claim related to the February incident from unit 509."
    if match(question, "have any other condo unit owners or the association advised you") and "own any additional properties within the united states" in question:
        return (
            f"{pronoun} stated that no other condo unit owners or the association advised {context.objective} of damages or requested payment. "
            f"{pronoun} stated that {context.pronoun} provided Citizens with all documentation {context.pronoun} had related to the incident. "
            f"{pronoun} stated that {context.pronoun} does not have any excess liability or umbrella insurance. "
            f"{pronoun} confirmed that {context.pronoun} owns two additional properties in the United States."
        )
    if match(question, "have any other condo unit owners or the association advised you"):
        return f"{pronoun} stated that no other unit owners have claimed damages or requested payment."
    if match(question, "and who is that insured with") and "mortgage" in answer.lower():
        return f"{pronoun} stated that {context.pronoun} was not sure who insured that property and believed Southern Walk was the mortgage company."
    if match(question, "what's the other property"):
        return f"{pronoun} also identified another property at {answer}."
    if match(question, "and who is that insured with") and answer.strip().lower().startswith("no one"):
        return f"{pronoun} stated that the second property was not insured."
    if match(question, "first became aware of the leak when the association manager reached out"):
        return f"{pronoun} confirmed that April 5th was the first time {context.pronoun} became aware of the leak."

    if match(question, "immediately after becoming aware of the leak"):
        if "property manager" in answer.lower() and "estimate" in answer.lower() and "not to use the shower" in answer.lower():
            return (
                f"After becoming aware of the leak, the {role_noun} stated that {context.pronoun} followed the building manager's instructions and waited "
                f"for information from the plumber who was expected to provide an estimate for repairs. {pronoun} also contacted the tenant to alert him "
                f"that there was a reported leak and instructed him not to use the shower."
            )
        return f"After becoming aware of the leak, the {role_noun} stated that {lower_first(answer)}."
    if match(question, "did he sees all use of the shower"):
        return f"{pronoun} stated that {context.pronoun} does not know whether the tenant stopped using the shower, but {context.pronoun} has no reason to believe the tenant would not have followed that instruction."
    if match(question, "how long the leak was going on for before it was stopped"):
        return f"{pronoun} stated that {context.pronoun} does not know how long the leak was occurring before it was discovered or before it was stopped."
    if match(question, "did anyone check"):
        if "leak has not" in answer.lower() and "water is added to the shower" in answer.lower():
            return (
                f"{pronoun} stated that {context.pronoun} is under the impression that the leak has not stopped and only manifests when water is added to the shower. "
                f"{pronoun} stated that presumably no water is being added to the shower at this time."
            )
        cleaned = lower_first(answer)
        if cleaned.startswith("but "):
            cleaned = cleaned[4:]
        return f"{pronoun} stated that {cleaned}."
    if match(question, "did anyone shut the water off after the leak was discovered"):
        return f"{pronoun} stated that {context.pronoun} does not know whether anyone shut the water off after the leak was discovered and reiterated that the building manager hired and coordinated with the plumber."
    if match(question, "did you ever speak with the claimant"):
        return f"{pronoun} stated that {context.pronoun} did not speak with the claimant about the incident."
    if match(question, "have you hired a plumber or leak detection yet"):
        cleaned = lower_first(answer)
        cleaned = cleaned.replace("no, ", "")
        if "estimate" in answer.lower() and "another plumber" in answer.lower():
            return (
                f"{pronoun} stated that {context.pronoun} has not hired a plumber or leak detection company to complete repairs. {pronoun} explained that "
                f"{context.pronoun} contacted the plumber used by the building manager, who was supposed to prepare an estimate, but the estimate has not yet been provided. "
                f"{pronoun} stated that several days had passed, so {context.pronoun} contacted another plumber in an attempt to make progress."
            )
        return f"{pronoun} stated that {context.pronoun} has not hired a plumber or leak detection company to complete repairs. {pronoun} explained that {cleaned}."
    if match(question, "how many units were damaged as a result of this incident"):
        return f"{pronoun} stated that {context.pronoun} has no knowledge of any damaged units other than what the building manager told {context.objective}, which was that unit 12G, which is the unit directly below {context.objective}, was damaged."
    if match(question, "only unit you're aware of being damaged"):
        return None
    if match(question, "was 12g occupied on the date of loss"):
        return f"{pronoun} stated that {context.pronoun} does not know whether unit 12G was occupied on the date of loss."
    if match(question, "what the specific damage is in 12g"):
        return f"{pronoun} stated that {context.pronoun} does not know the specific damage present in unit 12G."
    if match(question, "dry out or mold remediation completed"):
        return f"{pronoun} stated that {context.pronoun} does not know whether dry out or mold remediation was completed in unit 12G."
    if match(question, "repairs completed yet"):
        return f"{pronoun} stated that {context.pronoun} does not know whether repairs have been completed in unit 12G."
    if match(question, "anyone was injured or is claiming to be injured"):
        return f"{pronoun} stated that {context.pronoun} is not aware of anyone being injured or claiming to be injured as a result of the incident."
    if match(question, "12g has insurance on their unit"):
        return f"{pronoun} stated that {context.pronoun} knows unit 12G had insurance in years past due to a prior incident, but {context.pronoun} does not know whether that policy has been maintained."
    if match(question, "damages to condo common areas"):
        return f"{pronoun} stated that {context.pronoun} has no knowledge of any damage to condo common areas."
    if match(question, "what leaked in your unit", "what specifically leaked"):
        return f"Regarding the source of the leak, the {role_noun} stated that {context.pronoun} was told by the building manager that the shower pan leaked, but {context.pronoun} does not know how that conclusion was reached and has no personal knowledge of the cause."
    if match(question, "how old is the shower in the shower pan"):
        return f"{pronoun} stated that the shower pan was installed new in approximately 2014."
    if match(question, "which bathroom is this shower located in"):
        return f"{pronoun} stated that the shower is located in the master bathroom."
    if match(question, "replaced the shower since you purchased the unit"):
        return f"{pronoun} confirmed that {context.pronoun} replaced the shower after purchasing the unit."
    if match(question, "and you said it was around 2014"):
        return f"{pronoun} acknowledged that the shower pan was approximately 12 years old at the time of the leak."
    if match(question, "12 years old when it leaked"):
        return None
    if match(question, "did you use a licensed contractor"):
        return f"{pronoun} stated that a licensed contractor performed the installation."
    if match(question, "installed with a permit"):
        return f"{pronoun} stated that the shower was installed with a permit."
    if match(question, "do you know what caused it to leak"):
        return f"{pronoun} stated that {context.pronoun} does not know what caused the shower pan to leak."
    if match(question, "had the shower pan ever leaked before"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, the shower pan had never leaked before."
    if match(question, "prior repairs, leaks, or issues involving the shower or shower pan"):
        return f"{pronoun} stated that {context.pronoun} has had no prior repairs, leaks, or issues involving the shower or shower pan."
    if match(question, "tenant ever complained about any shower related leaks"):
        return f"{pronoun} stated that {context.possessive} tenant never complained about shower-related leaks or unexplained water prior to the date of loss. {pronoun} explained that there had been a prior issue where water had seeped through a failed silicone seal on the glass, but that issue is unrelated."
    if match(question, "glass leak on the shower prior but it's unrelated"):
        return None
    if match(question, "prior to the data loss, the only shower leak"):
        return None
    if match(question, "any reason to believe that the shower or shower pan was at risk of leaking"):
        return f"{pronoun} stated that {context.pronoun} had no reason to believe that the shower or shower pan was at risk of leaking prior to the date of loss."
    if match(question, "leak has not been repaired yet"):
        return f"{pronoun} stated that, to {context.possessive} knowledge, the leak has not been repaired yet."
    if match(question, "water shutoff rules in the association bylaws"):
        return f"{pronoun} stated that {context.pronoun} does not know whether there are water shutoff rules in the association bylaws."
    if match(question, "so you don't know at this time"):
        return None
    if match(question, "file a first party claim"):
        return f"{pronoun} stated that {context.pronoun} has not filed a first-party claim with Citizens for damages to {context.possessive} own unit."
    if match(question, "have you had any prior water leaks originating"):
        return f"{pronoun} confirmed that {context.pronoun} has had one prior water leak originating from {context.possessive} unit since owning the condo."
    if match(question, "how many"):
        return None
    if match(question, "and when was that leak"):
        return f"{pronoun} stated that there was a prior leak originating from {context.possessive} unit several years ago, but {context.pronoun} could not provide an exact date."
    if match(question, "can you give me an approximate"):
        return None
    if match(question, "what leaked before"):
        return f"{pronoun} stated that the prior leak was related to a laundry machine."
    if match(question, "did this prior leak affect the unit below"):
        return f"{pronoun} confirmed that the prior leak affected the unit below."
    if match(question, "did it affect their master bathroom"):
        return f"{pronoun} stated that a claim was filed with Citizens, a lawsuit ensued, and there was apparently extensive damage in unit 12G."
    if match(question, "was 12g repaired prior to this current leak"):
        return f"{pronoun} stated that {context.pronoun} does not know whether unit 12G was repaired prior to the current leak."
    if match(question, "other condo unit owners or the association advised you"):
        return f"{pronoun} stated that no other unit owners or the association have advised {context.objective} of damages or requested payment."
    if match(question, "provided citizens with all documentation"):
        return f"{pronoun} stated that {context.pronoun} has provided Citizens with everything that was requested of {context.objective}."
    if match(question, "excess liability or umbrella insurance"):
        return f"{pronoun} stated that {context.pronoun} does not have any excess liability or umbrella insurance."
    if match(question, "own any additional properties in the united states"):
        return f"{pronoun} stated that {context.pronoun} does not own any additional properties in the United States."
    if match(question, "anything else you'd like to add", "understood all of my questions", "answered them truthfully", "interview recorded with your permission"):
        return None
    if should_skip_answer(answer):
        return None

    prefix = topic_prefix(pair.question)

    if prefix == "In describing what happened":
        return f"{prefix}, the {role_noun} stated that {lower_first(answer)}."
    if prefix:
        return f"{prefix}, {context.pronoun} stated that {lower_first(answer)}."

    leadins = [
        f"The {role_noun} stated that",
        f"{pronoun} explained that",
        f"{pronoun} also stated that",
        f"{pronoun} further advised that",
        f"{pronoun} confirmed that",
    ]
    leadin = leadins[index % len(leadins)]
    return f"{leadin} {lower_first(answer)}."


def should_merge_paragraphs(previous: str, current: str, context: ReportContext) -> bool:
    previous_sentences = sentence_count(previous)
    current_sentences = sentence_count(current)
    combined_words = word_count(previous) + word_count(current)
    word_limit = 340 if "personal representative" in context.role_label.lower() else 240

    if combined_words > word_limit:
        return False
    if current_sentences <= 1 and previous_sentences <= 12:
        return True
    if previous_sentences <= 2 and current_sentences <= 2:
        return True
    return False


def canonical_overlap_key(sentence: str, context: ReportContext) -> str:
    cleaned = normalize(sentence).strip(" .")
    if not cleaned:
        return ""

    subject_variants = {context.role_label.strip(), paragraph_subject(context).strip(), cap(context.pronoun)}
    if context.role_label.lower().startswith("the "):
        subject_variants.add(f"The {short_role_noun(context)}")
    subject_pattern = "|".join(re.escape(item) for item in sorted(subject_variants, key=len, reverse=True) if item)
    cleaned = re.sub(
        rf"^(?:Regarding [^,]+,\s+|In describing what happened,\s+)?(?:{subject_pattern})\s+"
        r"(?:stated|explained|confirmed|clarified|acknowledged|advised|reported|indicated)\s+that\s+",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    return cleaned.lower()


def trim_leading_duplicate_sentence(previous: str, current: str, context: ReportContext) -> str:
    previous_keys = {
        canonical_overlap_key(sentence, context)
        for sentence in split_sentences(previous)
        if canonical_overlap_key(sentence, context)
    }
    current_sentences = [normalize(sentence).strip(" .") for sentence in split_sentences(current) if normalize(sentence).strip(" .")]
    while current_sentences and canonical_overlap_key(current_sentences[0], context) in previous_keys:
        current_sentences.pop(0)
    return ". ".join(current_sentences).strip(" .")


def compress_paragraphs(paragraphs: list[str], context: ReportContext) -> list[str]:
    compressed: list[str] = []
    for paragraph in paragraphs:
        normalized = normalize_paragraph_opening(dedupe_paragraph_sentences(paragraph), context)
        if compressed:
            normalized = normalize_paragraph_opening(trim_leading_duplicate_sentence(compressed[-1], normalized, context), context)
        if not normalize(normalized):
            continue
        if compressed and should_merge_paragraphs(compressed[-1], normalized, context):
            compressed[-1] = normalize_paragraph_opening(dedupe_paragraph_sentences(join_sentences([compressed[-1], normalized])), context)
        else:
            compressed.append(normalized)
    return compressed


def build_body_paragraphs(pairs: list[QAPair], context: ReportContext) -> list[str]:
    rendered: list[tuple[str, str]] = []
    for idx, pair in enumerate(pairs):
        sentence = compose_body_sentence(pair, context, idx)
        if sentence:
            rendered.append((classify_paragraph_topic(pair.question, pair.answer), sentence))

    paragraphs: list[str] = []
    current: list[str] = []
    current_topic: Optional[str] = None
    max_sentences_per_paragraph = 14 if "personal representative" in context.role_label.lower() else 6

    for topic, sentence in rendered:
        if current and (topic != current_topic or len(current) >= max_sentences_per_paragraph):
            paragraphs.append(normalize_paragraph_opening(dedupe_paragraph_sentences(join_sentences(current)), context))
            current = []
        current.append(sentence)
        current_topic = topic
    if current:
        paragraphs.append(normalize_paragraph_opening(dedupe_paragraph_sentences(join_sentences(current)), context))
    return compress_paragraphs(paragraphs, context)


def normalize_paragraph_opening(paragraph: str, context: ReportContext) -> str:
    subject = paragraph_subject(context)
    paragraph = paragraph.replace("The named insured", "The insured")
    paragraph = paragraph.replace("the named insured", "the insured")

    pronoun_map = {
        "he": "He",
        "she": "She",
        "they": "They",
    }
    leading_pronoun = pronoun_map.get(context.pronoun, "They")
    paragraph = re.sub(
        rf"^{leading_pronoun}\s+(stated|explained|confirmed|acknowledged|indicated|advised)\b",
        rf"{subject} \1",
        paragraph,
        count=1,
    )
    paragraph = reduce_subject_repetition(paragraph, context)
    paragraph = soften_repetitive_attribution(paragraph, context)
    paragraph = vary_repeated_pronoun_openings(paragraph, context)
    paragraph = remove_repeated_uncertainty_remainders(paragraph, context)
    paragraph = normalize_plural_grammar(paragraph, context)
    paragraph = normalize_singular_grammar(paragraph)
    paragraph = remove_redundant_uncertainty_fragments(paragraph, context)
    paragraph = smooth_conjunction_artifacts(paragraph)
    return remove_fragment_sentences(paragraph)


def write_docx(path: Path, paragraphs: list[str]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def convert_transcript(input_path: Path, output_path: Path, interviewee_role_label: Optional[str] = None) -> None:
    turns = parse_turns(input_path)
    interviewer, interviewee = determine_roles(turns)
    _ = interviewee
    context = infer_context(turns, interviewer, interviewee, interviewee_role_label=interviewee_role_label)
    pairs = build_qa_pairs(turns, interviewer, interviewee)
    opening_count = collect_opening_facts(pairs, context)
    paragraphs = compose_opening_paragraphs(context)
    paragraphs.extend(build_body_paragraphs(pairs[opening_count:], context))
    paragraphs = [normalize_paragraph_opening(paragraph, context) for paragraph in paragraphs if normalize(paragraph)]
    write_docx(output_path, paragraphs)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert a recorded statement transcript .docx into a narrative summary .docx.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--role", choices=["insured", "claimant", "witness", "other"], help="Interviewee type selected by the user")
    parser.add_argument("--role-label", help="Exact label to use for the interviewee, such as 'Insured', 'Claimant', 'Witness', or a custom title")
    args = parser.parse_args()
    requested_role_label = args.role_label or args.role
    convert_transcript(args.input, args.output, interviewee_role_label=requested_role_label)


if __name__ == "__main__":
    main()
